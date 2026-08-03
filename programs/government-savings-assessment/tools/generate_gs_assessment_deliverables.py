#!/usr/bin/env python3
"""Generate and validate GS assessment XLSX, DOCX, PDF from CSV sources."""

from __future__ import annotations

import csv
import re
import subprocess
import sys
from pathlib import Path

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except ImportError:
    raise SystemExit("Install openpyxl: pip install openpyxl")

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    raise SystemExit("Install python-docx: pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent
METRICS_CSV = ROOT / "03-analysis" / "verified-metrics-register.csv"
BUSINESS_CSV = ROOT / "03-analysis" / "government-savings-business-coverage-register.csv"
MATRIX_CSV = ROOT / "03-analysis" / "government-savings-coverage-matrix.csv"
CLAIMS_CSV = ROOT / "00-review" / "current-claim-register.csv"
PIPE_CSV = ROOT / "01-inventory" / "pipeline-job-inventory.csv"
EVIDENCE_CSV = ROOT / "04-leadership" / "evidence-register.csv" if (ROOT / "04-leadership" / "evidence-register.csv").exists() else None
XLSX_OUT = ROOT / "03-analysis" / "government-savings-coverage-matrix.xlsx"
DOCX_OUT = ROOT / "04-leadership" / "Government-Savings-Automation-Coverage-Assessment.docx"
PDF_OUT = ROOT / "04-leadership" / "Government-Savings-Automation-Coverage-Assessment.pdf"
REVIEW_DOCX = ROOT / "04-leadership" / "Government-Savings-Automation-Assessment-Review-Findings.docx"
REVIEW_PDF = ROOT / "04-leadership" / "Government-Savings-Automation-Assessment-Review-Findings.pdf"
REPORT_DATE = "July 20, 2026"
REBUILD_DATE = "July 21, 2026"
VERSION = "1.1"
PREPARED_BY = "Swapnil Patil"
ROLE = "QA Automation & Quality Program Lead, Government Savings"
CLASSIFICATION = "Confidential - Internal Use"

NAVY = "003057"
TEAL = "007A8C"
WHITE = "FFFFFF"
GREEN_BG = "C8E6C9"
AMBER_BG = "FFE0B2"
RED_BG = "FFCDD2"
GRAY_BG = "ECEFF1"


def load_csv(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    with path.open(encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        return list(reader.fieldnames or []), rows


def style_header(ws, row_idx: int = 1) -> None:
    fill = PatternFill("solid", fgColor=NAVY)
    font = Font(bold=True, color=WHITE, name="Calibri", size=10)
    for cell in ws[row_idx]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def auto_width(ws, max_width: int = 48) -> None:
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        length = max(len(str(c.value or "")) for c in col)
        ws.column_dimensions[letter].width = min(max(length + 2, 10), max_width)


def validate_metrics(metrics: list[dict[str, str]]) -> None:
    errors: list[str] = []
    for m in metrics:
        status = m.get("verification_status", "")
        safe = m.get("leadership_safe", "")
        if status == "Pending verification" and safe == "Yes":
            errors.append(f"{m['metric_id']}: pending metric marked leadership_safe=Yes")
        if status == "Verified" and "Rejected" in m.get("notes", ""):
            errors.append(f"{m['metric_id']}: inconsistent status/notes")
    rejected = []
    if CLAIMS_CSV.exists():
        _, claims = load_csv(CLAIMS_CSV)
        rejected = [c["claim_id"] for c in claims if c.get("verification_status") == "Rejected"]
    if errors:
        raise SystemExit("Validation failed:\n" + "\n".join(errors))
    if rejected:
        print(f"Validation note: {len(rejected)} rejected claims in register — ensure docs do not cite them.")


def status_fill(status: str) -> PatternFill:
    s = status.lower()
    if "verified" in s and "pending" not in s and "external" not in s:
        return PatternFill("solid", fgColor=GREEN_BG)
    if "pending" in s or "partial" in s or "conditional" in s:
        return PatternFill("solid", fgColor=AMBER_BG)
    if "rejected" in s or "gap" in s or "unknown" in s:
        return PatternFill("solid", fgColor=RED_BG)
    return PatternFill("solid", fgColor=GRAY_BG)


def build_xlsx(metrics: list[dict[str, str]]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Executive Summary"
    ws["A1"] = "Government Savings — Automation Coverage (Rebuilt)"
    ws["A1"].font = Font(bold=True, size=14, color=NAVY)
    ws["A2"] = f"Assessment: {REPORT_DATE}  |  Rebuild validation: {REBUILD_DATE}"
    ws.merge_cells("A1:G1")

    headers = ["Metric ID", "Label", "Result", "Numerator", "Denominator", "Status", "Leadership safe"]
    start = 4
    for c, h in enumerate(headers, 1):
        ws.cell(row=start, column=c, value=h)
    style_header(ws, start)

    leadership_rows = [m for m in metrics if m.get("leadership_safe") in ("Yes", "Conditional")]
    for r, m in enumerate(leadership_rows, start + 1):
        vals = [
            m["metric_id"], m["label"], m["formula"], m["numerator"], m["denominator"],
            m["verification_status"], m["leadership_safe"],
        ]
        for c, v in enumerate(vals, 1):
            cell = ws.cell(row=r, column=c, value=v)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if c == 6:
                cell.fill = status_fill(m["verification_status"])
    ws.freeze_panes = f"A{start + 1}"
    auto_width(ws)

    if BUSINESS_CSV.exists():
        bh, br = load_csv(BUSINESS_CSV)
        ws_b = wb.create_sheet("Business Coverage")
        for c, h in enumerate(bh, 1):
            ws_b.cell(row=1, column=c, value=h)
        style_header(ws_b, 1)
        for r, row in enumerate(br, 2):
            for c, h in enumerate(bh, 1):
                ws_b.cell(row=r, column=c, value=row.get(h, "")).alignment = Alignment(wrap_text=True)
        ws_b.freeze_panes = "A2"
        auto_width(ws_b)

    if MATRIX_CSV.exists():
        mh, mr = load_csv(MATRIX_CSV)
        ws_m = wb.create_sheet("Coverage Matrix")
        for c, h in enumerate(mh, 1):
            ws_m.cell(row=1, column=c, value=h)
        style_header(ws_m, 1)
        for r, row in enumerate(mr, 2):
            for c, h in enumerate(mh, 1):
                ws_m.cell(row=r, column=c, value=row.get(h, "")).alignment = Alignment(wrap_text=True)
        ws_m.freeze_panes = "A2"
        ws_m.auto_filter.ref = f"A1:{get_column_letter(len(mh))}{len(mr)+1}"
        auto_width(ws_m)

    if PIPE_CSV.exists():
        ph, pr = load_csv(PIPE_CSV)
        ws_p = wb.create_sheet("CI Jobs Gates")
        for c, h in enumerate(ph, 1):
            ws_p.cell(row=1, column=c, value=h)
        style_header(ws_p, 1)
        for r, row in enumerate(pr, 2):
            for c, h in enumerate(ph, 1):
                ws_p.cell(row=r, column=c, value=row.get(h, "")).alignment = Alignment(wrap_text=True)
        auto_width(ws_p)

    _, claims = load_csv(CLAIMS_CSV)
    unresolved = [c for c in claims if c.get("verification_status") in ("Rejected", "Pending verification", "Unknown")]
    ws_u = wb.create_sheet("Unresolved Verification")
    uh = ["claim_id", "claim_text", "verification_status", "action_required"]
    for c, h in enumerate(uh, 1):
        ws_u.cell(row=1, column=c, value=h)
    style_header(ws_u, 1)
    for r, row in enumerate(unresolved, 2):
        for c, h in enumerate(uh, 1):
            cell = ws_u.cell(row=r, column=c, value=row.get(h, ""))
            cell.alignment = Alignment(wrap_text=True)
            cell.fill = status_fill(row.get("verification_status", ""))
    auto_width(ws_u)

    wb.save(XLSX_OUT)
    print(f"Wrote {XLSX_OUT}")


def kill_word() -> None:
    if sys.platform == "win32":
        subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"], capture_output=True)


def set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level in range(1, 4):
        heading = doc.styles[f"Heading {level}"]
        heading.font.name = "Calibri"
        heading.font.color.rgb = RGBColor(0x00, 0x30, 0x57)


def set_docx_header_footer(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    hp = section.header.paragraphs[0]
    hp.text = title
    for run in hp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x57)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{CLASSIFICATION}  |  {REBUILD_DATE}  |  Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].bold = True
        sub.runs[0].font.name = "Calibri"
    meta = doc.add_paragraph(
        f"\nPrepared by: {PREPARED_BY}\n"
        f"{ROLE}\n\n"
        f"Assessment Date: {REPORT_DATE}\n"
        f"Rebuild validated: {REBUILD_DATE}\n"
        f"Version: {VERSION}\n"
        f"Classification: {CLASSIFICATION}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    doc.add_page_break()


def _shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def sanitize_text(text: str) -> str:
    return (
        text.replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2026", "...")
        .replace("\u2260", "is not")
    )


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(sanitize_text(str(text)))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_table(doc: Document, headers: list[str], data: list[list[str]], row_shades: list[str] | None = None) -> None:
    if not data:
        doc.add_paragraph("No items in this section.")
        doc.add_paragraph()
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, header, bold=True, size=10)
        _shade_cell(cell, NAVY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_idx, row_data in enumerate(data):
        cells = table.add_row().cells
        shade = row_shades[row_idx] if row_shades and row_idx < len(row_shades) else (GRAY_BG if row_idx % 2 else WHITE)
        for i, val in enumerate(row_data):
            _set_cell_text(cells[i], val, size=9)
            if shade != WHITE:
                _shade_cell(cells[i], shade)
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def shorten_evidence(path: str, limit: int = 72) -> str:
    cleaned = path.replace("\\", "/")
    for prefix in (
        "C:/Workspace/GitLab/",
        "c:/Workspace/GitLab/",
        "C:/Workspace/GitLab/qa-automation-kb/",
    ):
        if cleaned.lower().startswith(prefix.lower()):
            cleaned = cleaned[len(prefix):]
            break
    cleaned = re.sub(r"^programs/government-savings-assessment/", "", cleaned)
    if len(cleaned) > limit:
        return cleaned[: limit - 1] + "…"
    return cleaned


def parse_contradiction_ledger() -> list[list[str]]:
    ledger = ROOT / "00-review" / "contradiction-resolution-ledger.md"
    if not ledger.exists():
        return []
    rows: list[list[str]] = []
    for line in ledger.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| C-"):
            continue
        parts = [p.strip() for p in line.strip("|").split("|")]
        if len(parts) < 7:
            continue
        rows.append([
            parts[0],
            sanitize_text(parts[1]),
            sanitize_text(parts[4].replace("**", "")),
            sanitize_text(parts[5].replace("**", "")),
        ])
    return rows


def severity_shade(severity: str) -> str:
    s = severity.lower()
    if s == "high":
        return RED_BG
    if s == "medium":
        return AMBER_BG
    return GRAY_BG


def build_main_docx(metrics: list[dict[str, str]]) -> None:
    doc = Document()
    set_document_defaults(doc)
    set_docx_header_footer(doc, "Ascensus Government Savings  |  Automation Coverage Assessment")
    add_title_page(
        doc,
        "Government Savings",
        "Automation Coverage & CI Integration Assessment",
    )

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Government Savings automation is organized by business platform and operational activation - "
        "what is implemented, what runs on recurring schedules, and what awaits reactivation. "
        "A single GS-wide coverage percentage is not defensible. Inventory-share metrics (e.g. V2/V3 "
        "TestNG or qTest populations) appear in the technical appendix only, not as functional completeness."
    )

    if BUSINESS_CSV.exists():
        doc.add_heading("Business Platform Coverage", level=1)
        _, business = load_csv(BUSINESS_CSV)
        leadership_business = [
            b for b in business
            if b.get("percent_class") != "Verified inventory share only"
            and "technical appendix" not in b.get("notes", "").lower()
            and "NOT leadership headline metric" not in b.get("notes", "")
        ][:14]
        add_table(
            doc,
            ["Business area", "Sub-area", "Position", "Implementation", "Scheduled", "Confidence"],
            [
                [
                    b["business_area"],
                    b["sub_area"],
                    b["leadership_safe_wording"][:100] + ("…" if len(b["leadership_safe_wording"]) > 100 else ""),
                    b["implementation_status"],
                    b["scheduled_status"],
                    b["confidence"],
                ]
                for b in leadership_business
            ],
        )

    doc.add_heading("Verified Metrics Register", level=1)
    verified = [m for m in metrics if m.get("leadership_safe") == "Yes" and m.get("verification_status") == "Verified"]
    add_table(
        doc,
        ["Area", "Metric", "Result", "As of", "Evidence"],
        [
            [
                m["gs_domain"],
                m["label"],
                m["formula"],
                m["as_of_date"][:10],
                shorten_evidence(m["evidence_path"]),
            ]
            for m in verified
        ],
    )

    doc.add_heading("Conditional / Gap Metrics (do not quote as final)", level=1)
    doc.add_paragraph(
        "The following metrics are leadership-relevant but not fully verified execution or scheduling evidence."
    )
    pending = [
        m for m in metrics
        if m.get("verification_status") != "Verified"
        or m.get("leadership_safe") in ("Conditional",)
    ]
    add_table(
        doc,
        ["Area", "Metric", "Value", "Status", "Dependency"],
        [
            [
                m["gs_domain"],
                m["label"],
                m["formula"],
                m["verification_status"],
                m["notes"][:90] + ("…" if len(m["notes"]) > 90 else ""),
            ]
            for m in pending
        ],
    )

    doc.add_heading("CI/CD Integration", level=1)
    doc.add_paragraph(
        "GitLab: V3 scheduled_regression_job and metadataweb Stage1 - scheduled jobs exit non-zero on failure "
        "(not verified as MR merge gates). GitHub Actions: Mobile 2 Dashboard vertical slice externally validated; "
        "workflow repository not in audit clone. Jenkins: IDP performance scheduled; V2 UI recurring job not verified. "
        "Mobile 2 GitLab nightly not created (QA-1405)."
    )

    doc.add_heading("Leadership Narrative", level=1)
    doc.add_paragraph(
        "The team has meaningful automation across Government Savings. The current limitation is not the ability "
        "to build automation; it is establishing a consistent, traceable coverage model across Jira, qTest, "
        "repositories, and CI execution."
    )

    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [[VERSION, REBUILD_DATE, PREPARED_BY, "Formatting refresh; conditional metrics table; evidence path cleanup"]],
    )

    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


def build_review_docx() -> None:
    doc = Document()
    set_document_defaults(doc)
    set_docx_header_footer(doc, "Ascensus GS  |  Assessment Review Findings")
    add_title_page(
        doc,
        "Government Savings Automation Assessment",
        "Review Findings - Rebuild & Verification",
    )

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "Independent rebuild review of the Government Savings automation coverage assessment prior to "
        "2026-07-21 leadership distribution. Identifies contradictions, corrects headline metrics, and "
        "separates verified from pending evidence."
    )

    doc.add_heading("Executive Findings", level=1)
    findings = [
        ("F-01", "High", "Mobile 2 100% claim corrected", "Prior 24/24 superseded by leadership 22/25 (88%). Projected 24/25 (96%) pending sign-off."),
        ("F-02", "High", "Mobile 1 verified baseline retained", "1/27 (3.7%) verified; 6/27 code-implemented pending QC4/Stage1 evidence."),
        ("F-03", "Medium", "GitHub Actions wording corrected", "Dashboard slice externally validated; not equivalent to absent implementation."),
        ("F-04", "Medium", "CI gate terminology corrected", "Scheduled hard-fail is not a merge/deployment gate without pipeline rule evidence."),
        ("F-05", "High", "No GS-wide %", "Enterprise percentage remains undefined by design."),
        ("F-06", "Low", "COPACS scope unknown", "No automation repository identified in reviewed set."),
    ]
    add_table(
        doc,
        ["ID", "Severity", "Finding", "Resolution"],
        findings,
        row_shades=[severity_shade(row[1]) for row in findings],
    )

    doc.add_page_break()
    doc.add_heading("Contradictions Resolved", level=1)
    contradictions = parse_contradiction_ledger()
    if contradictions:
        add_table(
            doc,
            ["ID", "Topic", "Corrected claim", "Status"],
            contradictions,
        )
    else:
        doc.add_paragraph("Contradiction ledger not available.")

    doc.add_heading("Approved Leadership Narrative", level=1)
    doc.add_paragraph(
        "The team has meaningful automation across Government Savings. The current limitation is not the ability "
        "to build automation; it is establishing a consistent, traceable coverage model across Jira, qTest, "
        "repositories, and CI execution. We can report verified platform-specific metrics now, but we should not "
        "publish one GS-wide percentage until the denominators and pipeline evidence are governed."
    )

    doc.add_heading("Recommendations", level=1)
    recs = [
        "Retain 22/25 (88%) as verified Mobile 2 leadership baseline until sign-off path completes.",
        "Complete Mobile 2 GitLab nightly (QA-1405) before claiming execution or gate coverage.",
        "Report Mobile 1 as 1/27 verified with sprint progress tracked separately.",
        "Provision read-only qTest/Jira/GitLab API access for automated register.",
        "Do not regenerate DOCX headline numbers without updating verified-metrics-register.csv.",
    ]
    for rec in recs:
        doc.add_paragraph(rec, style="List Number")

    doc.add_heading("Artifacts Reviewed", level=1)
    doc.add_paragraph(
        "Primary assessment folder: programs/government-savings-assessment/ "
        "(verified-metrics-register.csv, business-coverage-register, contradiction-resolution-ledger.md, "
        "pipeline-job-inventory.csv, and linked repository evidence)."
    )

    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [[VERSION, REBUILD_DATE, PREPARED_BY, "Formatting refresh; contradiction table; severity shading"]],
    )

    doc.save(REVIEW_DOCX)
    print(f"Wrote {REVIEW_DOCX}")


def export_pdf(docx_path: Path, pdf_path: Path) -> str:
    kill_word()
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        pages = 0
        try:
            from pypdf import PdfReader
            pages = len(PdfReader(str(pdf_path)).pages)
        except Exception:
            pass
        return f"OK - {pdf_path} ({pages} pages)"
    except Exception as exc:
        return f"SKIPPED - {exc}"


def main() -> None:
    _, metrics = load_csv(METRICS_CSV)
    validate_metrics(metrics)
    build_xlsx(metrics)
    build_main_docx(metrics)
    build_review_docx()
    r1 = export_pdf(DOCX_OUT, PDF_OUT)
    r2 = export_pdf(REVIEW_DOCX, REVIEW_PDF)
    print(r1)
    print(r2)


if __name__ == "__main__":
    main()
