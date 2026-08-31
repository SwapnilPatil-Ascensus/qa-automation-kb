#!/usr/bin/env python3
"""Generate leadership DOCX, technical DOCX, and coverage matrix XLSX."""
from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

ROOT = Path(__file__).resolve().parents[1]
ANALYSIS = ROOT / "03-analysis"
LEADERSHIP = ROOT / "04-leadership"
ASSESSMENT_DATE = "20 July 2026"
NAVY = RGBColor(0x00, 0x33, 0x66)
TEAL = RGBColor(0x00, 0x80, 0x80)


def read_csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def add_header_footer(doc: Document, title: str, doc_type: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"Government Savings Automation | {title}"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = NAVY
    footer = section.footer.paragraphs[0]
    footer.text = f"Confidential | Assessment {ASSESSMENT_DATE} | {doc_type} | Page "
    footer.runs[0].font.size = Pt(8)


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)
    r2.font.color.rgb = TEAL
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Assessment Date: {ASSESSMENT_DATE}").font.size = Pt(11)
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_leadership_docx() -> Path:
    out = LEADERSHIP / "Government-Savings-Automation-Coverage-Assessment.docx"
    doc = Document()
    add_header_footer(doc, "Coverage Assessment", "Leadership")
    add_title_page(doc, "Government Savings", "Automation Coverage Assessment")

    summary = (ROOT / "04-leadership" / "leadership-summary.md").read_text(encoding="utf-8")
    for line in summary.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("|") and "---" not in line:
            pass  # tables handled below
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())

    metrics = read_csv(ANALYSIS / "verified-metrics-register.csv")
    rows = [[m["platform"], m["metric_name"], m["numerator"], m["denominator"], m["status_label"]] for m in metrics[:12]]
    doc.add_heading("Verified Metrics Snapshot", level=2)
    add_table(doc, ["Platform", "Metric", "Numerator", "Denominator", "Status"], rows)

    doc.save(out)
    return out


def build_technical_docx() -> Path:
    out = LEADERSHIP / "Government-Savings-Automation-Coverage-Assessment-Technical.docx"
    doc = Document()
    add_header_footer(doc, "Technical Appendix", "Technical")
    add_title_page(doc, "Government Savings Automation", "Technical Evidence Appendix")

    doc.add_heading("Mobile 2 Endpoint Inventory", level=1)
    m2 = read_csv(ROOT / "01-inventory" / "mobile2-endpoint-current-state.csv")
    rows = [[r["endpoint_id"], r["http_method"], r["path"], r["test_class"], r["status"], r["execution_model"]] for r in m2]
    add_table(doc, ["ID", "Method", "Path", "Class", "Status", "Execution Model"], rows)

    doc.add_page_break()
    doc.add_heading("Mobile 1 Endpoint Inventory", level=1)
    m1 = read_csv(ROOT / "01-inventory" / "mobile1-endpoint-current-state.csv")
    rows = [[r["endpoint_id"], r["http_method"], r["path"], r["test_class"], r["status"]] for r in m1]
    add_table(doc, ["ID", "Method", "Path", "Class", "Status"], rows)

    doc.add_page_break()
    doc.add_heading("Pipeline Inventory", level=1)
    pipes = read_csv(ROOT / "01-inventory" / "pipeline-job-inventory.csv")
    rows = [[p["job_id"], p["platform"], p["job_name"], p["gate_class"], p["validated_live"]] for p in pipes]
    add_table(doc, ["ID", "Platform", "Job", "Gate Class", "Live Validated"], rows)

    notes = (ANALYSIS / "coverage-calculation-notes.md").read_text(encoding="utf-8")
    doc.add_page_break()
    doc.add_heading("Coverage Calculation Notes", level=1)
    for line in notes.splitlines():
        if line.startswith("##"):
            doc.add_heading(line.lstrip("# ").strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(out)
    return out


def build_xlsx() -> Path:
    out = ANALYSIS / "government-savings-coverage-matrix.xlsx"
    wb = Workbook()
    header_fill = PatternFill("solid", fgColor="003366")
    header_font = Font(color="FFFFFF", bold=True)

    sheets = {
        "Verified Metrics": ANALYSIS / "verified-metrics-register.csv",
        "Mobile2 Endpoints": ROOT / "01-inventory" / "mobile2-endpoint-current-state.csv",
        "Mobile1 Endpoints": ROOT / "01-inventory" / "mobile1-endpoint-current-state.csv",
        "Coverage Matrix": ANALYSIS / "government-savings-coverage-matrix.csv",
        "Suite Placement": ANALYSIS / "suite-placement-register.csv",
        "Pipelines": ROOT / "01-inventory" / "pipeline-job-inventory.csv",
    }

    first = True
    for name, path in sheets.items():
        ws = wb.active if first else wb.create_sheet(name)
        if first:
            ws.title = name
            first = False
        else:
            ws.title = name[:31]
        rows = read_csv(path)
        if not rows:
            continue
        headers = list(rows[0].keys())
        ws.append(headers)
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(wrap_text=True)
        for row in rows:
            ws.append([row.get(h, "") for h in headers])
        for col in ws.columns:
            max_len = max(len(str(c.value or "")) for c in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

    wb.save(out)
    return out


def try_pdf(docx_path: Path) -> Path | None:
  pdf_path = docx_path.with_suffix(".pdf")
  try:
    from docx2pdf import convert
    convert(str(docx_path), str(pdf_path))
    return pdf_path
  except Exception:
    return None


def main() -> None:
    leadership = build_leadership_docx()
    technical = build_technical_docx()
    xlsx = build_xlsx()
    pdf_l = try_pdf(leadership)
    pdf_t = try_pdf(technical)
    print(f"Created: {leadership}")
    print(f"Created: {technical}")
    print(f"Created: {xlsx}")
    if pdf_l:
        print(f"Created: {pdf_l}")
    if pdf_t:
        print(f"Created: {pdf_t}")
    if not pdf_l:
        print("PDF: not generated (docx2pdf unavailable or Word not installed)")


if __name__ == "__main__":
    main()
