#!/usr/bin/env python3
"""Generate modern leadership briefing DOCX + classic-path PPTX (does not touch Executive/Detailed decks)."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from leadership_design import (
    CHARTS,
    DELIVERABLES,
    DOCX_CORAL,
    DOCX_CYAN,
    DOCX_INK,
    DOCX_INSIGHT,
    DOCX_MUTED,
    DOCX_NAVY,
    DOCX_SURFACE,
    DOCX_TEAL,
    DOCX_TEXT,
    DOCX_VIOLET,
    REFS,
    SUBTITLE,
    TITLE,
    add_footer,
    add_header,
    load_metrics,
    new_presentation,
    slide_bullets_modern,
    slide_chart_insight,
    slide_close,
    slide_hero,
    slide_kpi_dashboard,
    slide_section_modern,
)

DOCX_OUT = DELIVERABLES / "AM-Squad-Leadership-Briefing-Aug2026.docx"
PPTX_OUT = DELIVERABLES / "AM-Squad-Leadership-Update-Aug2026.pptx"

# Protected — never overwrite from this script
PROTECTED = {
    DELIVERABLES / "AM-Squad-Leadership-Executive-Glimpse-Aug2026.pptx",
    DELIVERABLES / "AM-Squad-Leadership-Detailed-Modern-Aug2026.pptx",
}

WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x64, 0x74, 0x8B)
NAVY = RGBColor(0x00, 0x32, 0x41)
TEAL = RGBColor(0x00, 0xB3, 0x88)


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def styled_run(paragraph, text: str, *, bold=False, size=11, color: RGBColor | None = None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Segoe UI"
    if color:
        run.font.color.rgb = color
    return run


def add_section_banner(doc: Document, number: str, title: str, subtitle: str = "") -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    left, right = table.rows[0].cells
    shade_cell(left, DOCX_CYAN)
    left.width = Cm(0.5)
    set_cell_margins(left, 40, 40, 40, 40)
    shade_cell(right, DOCX_INK)
    set_cell_margins(right)
    p = right.paragraphs[0]
    styled_run(p, f"{number}  ", bold=True, size=10, color=TEAL)
    styled_run(p, title, bold=True, size=16, color=WHITE)
    if subtitle:
        p2 = right.add_paragraph()
        styled_run(p2, subtitle, size=10, color=GRAY)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, bullets: list[str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    shade_cell(left, DOCX_TEAL)
    left.width = Cm(0.35)
    shade_cell(right, DOCX_INSIGHT)
    set_cell_margins(right)
    p = right.paragraphs[0]
    styled_run(p, title, bold=True, size=11, color=NAVY)
    for b in bullets:
        bp = right.add_paragraph()
        styled_run(bp, f"• {b}", size=10, color=RGBColor(0x1A, 0x23, 0x32))
    doc.add_paragraph()


def add_chart_section(
    doc: Document,
    heading: str,
    chart_name: str,
    caption: str,
    takeaways: list[str],
) -> None:
    """Standard executive layout: heading → narrative chart → caption → takeaways."""
    doc.add_heading(heading, level=2)
    path = CHARTS / chart_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        p.add_run().add_picture(str(path), width=Inches(6.4))
    else:
        styled_run(p, f"[Chart: {chart_name}]", italic=True, color=GRAY)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(cap, caption, italic=True, size=9, color=GRAY)
    doc.add_paragraph()
    styled_run(doc.add_paragraph(), "Key takeaways", bold=True, size=10, color=NAVY)
    for item in takeaways:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()


def add_exec_summary_table(doc: Document, m: dict) -> None:
    sc = m["scorecard"]
    j = m["jira"]["totals"]
    dc = m.get("data_confidence", {})
    perf_base = sum(a["labels"] for a in m["perf_inventory"]["areas"])
    rows = [
        ("GitLab merges to main", str(sc["gitlab_merges"]), "Apr–Aug 2026 · 3 repositories"),
        ("Jira story points delivered", f"{j['story_points']:.0f}", "AMSQUAD Sprints 26.04–26.12"),
        ("Automation bugs discovered", str(j["automation_bugs_logged"]), "Logged via regression triage"),
        ("V2 Stage1 nightly methods", str(sc["v2_nightly_methods"]), "Jenkins nightly · Aug 4 · ~12 mo build"),
        ("V3 Stage1 nightly methods", str(sc["v3_nightly_methods"]), "GitLab CI nightly · Aug 4 snapshot"),
        ("Performance test cases", str(sc["perf_test_cases"]), f"{perf_base} base flows × plan matrix"),
        ("MSC Mobile 2 / M1 core", f"{sc['msc_m2_endpoints']} · {sc['msc_m1_core']}", "api-test-automation"),
        ("Period delivery (Apr–Aug)", str(dc.get("period_delivery_estimate", "—")), "Estimated new coverage"),
        ("Release automation", f"~{sc['release_automation_pct']}%", "17 FTE → 2 FTE equivalent"),
    ]
    add_data_table(doc, ["Metric", "Value", "Context"], rows)


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = DOCX_NAVY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        shade_cell(c, header_fill)
        for run in c.paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def add_toc(doc: Document) -> None:
    doc.add_heading("Contents", 1)
    p = doc.add_paragraph()
    run = p.add_run()
    for el in (
        OxmlElement("w:fldChar"),
        OxmlElement("w:instrText"),
        OxmlElement("w:fldChar"),
        OxmlElement("w:fldChar"),
    ):
        pass
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    note = doc.add_paragraph()
    styled_run(note, "Right-click Contents → Update Field in Word to refresh page numbers.", italic=True, size=9, color=GRAY)
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    """Clean corporate cover — title block + metadata table."""
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styled_run(t, TITLE, bold=True, size=24, color=NAVY)
    s = doc.add_paragraph()
    styled_run(s, SUBTITLE, size=14, color=TEAL)
    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Prepared by", "QA Automation — AM Squad"),
        ("Reporting period", SUBTITLE),
        ("Version", datetime.now().strftime("%B %Y")),
        ("Classification", "Internal — Leadership distribution"),
    ]
    for i, (label, value) in enumerate(fields):
        shade_cell(meta.rows[i].cells[0], DOCX_SURFACE)
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = value
        for run in meta.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
        for run in meta.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
    doc.add_page_break()


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        styled_run(hp, f"{TITLE}  |  {SUBTITLE}", size=8, color=GRAY)
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(fp, f"QA Automation AM Squad  ·  Confidential  ·  Page ", size=8, color=GRAY)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r = fp.add_run()
        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_end)


def build_docx(m: dict) -> None:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)

    add_cover(doc)
    add_header_footer(doc)
    add_toc(doc)

    sc = m["scorecard"]
    j = m["jira"]
    jt = j["totals"]
    dc = m.get("data_confidence", {})
    ui = m.get("ui_inventory_scope", {})
    v3 = m["v3_snapshot"]
    perf_base = sum(a["labels"] for a in m["perf_inventory"]["areas"])

    # ── Executive Summary ─────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The QA Automation AM Squad delivers across six parallel channels — V2 Legacy UI, V3 Universal "
        "Platform, API/Unite MSC, performance testing, pipeline/CI, and department standards. This briefing "
        "covers Apr–Aug 2026 delivery within a ~12-month team build (squad formed Q2 2025)."
    )
    headlines = [
        f"{sc['gitlab_merges']} GitLab merges to main across three automation repositories (peak Jun–Jul MSC sprint)",
        f"{jt['story_points']:.0f} Jira story points · {jt['work_items_in_sprints']} work items · {jt['automation_bugs_logged']} automation bugs logged",
        f"Stage1 nightly inventory: V2 {sc['v2_nightly_methods']} + V3 {sc['v3_nightly_methods']} methods "
        f"(built since Q2 2025; excludes smoke/Stage 2/5/integrations); perf {sc['perf_test_cases']} plan-expanded cases",
        f"Unite MSC rescued — M2 {sc['msc_m2_endpoints']} endpoints, M1 {sc['msc_m1_core']} core (~50% ETA savings)",
        f"~{sc['release_automation_pct']}% release validation automated (17 FTE → 2 FTE equivalent)",
    ]
    for h in headlines:
        doc.add_paragraph(h, style="List Bullet")
    add_exec_summary_table(doc, m)
    scope_note = doc.add_paragraph()
    styled_run(
        scope_note,
        ui.get("scorecard_footnote") or dc.get("scorecard_footnote", ""),
        italic=True,
        size=9,
        color=GRAY,
    )
    doc.add_page_break()

    # ── 1 Delivery Velocity ───────────────────────────────────────────────────
    doc.add_heading("1. Delivery Velocity", level=1)
    doc.add_paragraph(
        "GitLab merge activity measures code delivery throughput. Monthly automation delivery measures "
        "estimated new test coverage closed in each month — these are separate metrics and should not be combined."
    )
    add_chart_section(
        doc, "GitLab Merges by Repository", "01-gitlab-mrs-by-month.png",
        f"Source: GitLab MR export · Apr 1 – Aug 4, 2026 · Total {sc['gitlab_merges']} merges",
        [
            f"{sc['gitlab_merges']} merges to main; July peak at 38 merges during MSC API sprint",
            "Repositories: automation (V2), prime-test-automation (V3), api-test-automation (MSC)",
            "Merge count is a delivery-velocity metric — not the same as test case inventory",
        ],
    )
    add_chart_section(
        doc, "Monthly Automation Delivery", "08-monthly-automation-test-cases-added.png",
        f"Source: Jira AMSQUAD Sprints 26.04–26.12 · ~{dc.get('period_delivery_estimate', 1212)} est. cases Apr–Aug",
        [
            "Period delivery estimate — not cumulative nightly inventory",
            "Counts include multi-plan, multi-environment, and positive/negative permutations",
            "Pre-April foundation is in nightly totals; this chart shows when work landed",
        ],
    )

    # ── 2 Jira Sprint Delivery ────────────────────────────────────────────────
    doc.add_heading("2. Jira Sprint Delivery", level=1)
    doc.add_paragraph(
        f"Across AMSQUAD Sprints 26.04–26.12 the squad closed {jt['work_items_in_sprints']} work items "
        f"({jt['story_points']:.0f} story points) and logged {jt['automation_bugs_logged']} automation-discovered defects."
    )
    add_chart_section(doc, "Story Points by Sprint", "09-jira-story-points-by-sprint.png",
        "Source: Jira AMSQUAD export", ["Sustained sprint delivery across reporting window", "Story points track committed work closed"])
    add_chart_section(doc, "Automation Bugs by Sprint", "10-jira-automation-bugs-by-sprint.png",
        "Source: Jira AMSQUAD export", ["Defects found via nightly regression triage", "Fed into automation bug lifecycle standard"])
    doc.add_heading("Sprint detail", level=2)
    add_data_table(doc, ["Sprint", "Work Items", "Stories", "Spikes", "Bugs", "SP"],
        [[s["sprint"].replace("AMSQUAD Sprint ", ""), str(s["work_items"]), str(s["stories"]),
          str(s["spikes"]), str(s["bugs"]), f"{s['story_points']:.0f}"] for s in j["sprints"]])

    # ── 3 V2 ──────────────────────────────────────────────────────────────────
    doc.add_heading("3. V2 Legacy UI Automation", level=1)
    doc.add_paragraph(
        f"V2 Jenkins Stage1 nightly snapshot: {sc['v2_nightly_methods']} test methods across 12 modules "
        f"(Aug 4, 2026). This count reflects the primary Mon–Fri nightly job only — built since Q2 2025, "
        "not delivered in Apr–Aug alone."
    )
    doc.add_paragraph(
        "Not included in 592: Stage 5 smoke, Stage 2 smoke, on-demand fast smoke, and +33 CSR Actions "
        "scenarios (built — pending Jenkins nightly wire). Empower also runs as a separate dedicated nightly job."
    )
    add_chart_section(doc, "V2 Test Methods by Module", "04-v2-regression-by-module.png",
        "Source: Jenkins STAGE1-Daily-Unite-Prime-Regression · Aug 4, 2026 snapshot",
        [
            "592 total methods — Stage1 primary nightly only",
            "Enrollments largest module (scenario × plan matrix)",
            "CSR maintenance modules added Apr–Jul on earlier foundation",
            "+33 CSR Actions scenarios — next expansion (not yet in nightly job)",
        ],
    )

    # ── 4 V3 ──────────────────────────────────────────────────────────────────
    doc.add_heading("4. V3 Universal Platform", level=1)
    doc.add_paragraph(
        f"V3 GitLab Stage1 nightly snapshot: {v3['total_methods']} test methods (Aug 4, 2026). "
        "Framework, CI/CD pipelines, and suites were built since Q2 2025 — Apr–Aug is acceleration, not greenfield."
    )
    doc.add_paragraph(
        "Not included in 442: Stage 5 smoke suites (UE + IDP), integration XML profiles, and Entity suites "
        "still expanding on a separate GitLab track. Universal Enrollment (303) reflects enrollment scenarios "
        "multiplied across plan/traunch permutations — not 303 unique manual scripts."
    )
    add_chart_section(doc, "V3 Test Methods by Module", "11-v3-regression-by-module.png",
        "Source: GitLab scheduled_regression_job · Aug 4, 2026 snapshot",
        [
            "442 total methods across UE, IDP Login, Web Registration, CSR, Contributions, Withdrawals",
            "Entity platform suites expanding on a separate GitLab track",
            "GitLab CI scheduled regression operational",
        ],
    )
    add_data_table(doc, ["Module", "Test Methods", "Notes"],
        [[mod["module"].replace(" Stage1 Environment", ""), str(mod["methods"]),
          "Largest module — multi-plan enrollment matrix" if mod["methods"] > 200 else "Functional coverage"]
         for mod in v3["modules"]], header_fill=DOCX_VIOLET)

    # ── 5 API / MSC ─────────────────────────────────────────────────────────
    doc.add_heading("5. API / Unite MSC", level=1)
    doc.add_paragraph(
        "Mobile 2 API automation is complete at 25/25 endpoints. Mobile 1 is at ~25/29 core endpoints. "
        "Delivered in approximately half the original ETA using AI-assisted migration."
    )
    add_chart_section(doc, "MSC Endpoint Coverage", "05-unite-msc-coverage.png",
        "Source: api-test-automation repo inventory", [f"M2 {sc['msc_m2_endpoints']} · M1 {sc['msc_m1_core']}", "P0: GitLab nightly scheduling (QA-1405)"])
    add_chart_section(doc, "API Module Breakdown", "13-api-regression-by-module.png",
        "Source: api-test-automation repo inventory", ["Auth, profile, biometric, device, bank categories", "M1 master suite in progress"])

    # ── 6 Performance ───────────────────────────────────────────────────────
    doc.add_heading("6. Performance Testing", level=1)
    doc.add_paragraph(
        f"Performance inventory: {perf_base} base transaction flows expand to {sc['perf_test_cases']} test cases "
        "when multiplied across plan permutations (e.g. IDP × 7 plans). Four Jenkins scenarios schedule these runs."
    )
    add_chart_section(doc, "Performance Flows vs Expanded Cases", "12-perf-test-case-inventory.png",
        "Source: Perf inventory model · Jenkins + BlazeMeter",
        [
            f"{perf_base} base flows → {sc['perf_test_cases']} plan-expanded test cases",
            "IDP login: 15 transaction labels × 7 plans = 105 cases alone",
            "Barcode SYN-443 emergency cycle delivered in ~1 week",
        ],
    )

    # ── 7 Portfolio Value ─────────────────────────────────────────────────────
    doc.add_heading("7. Portfolio Value & Investment", level=1)
    for item in [
        "Framework architecture and canonical repo structure (API, perf, UI)",
        "AI-accelerated documentation, Postman collections, and TestNG boilerplate",
        "qTest master suite design and automation bug lifecycle standard",
        "Pipeline/DevOps co-design and cross-team emergency support",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_chart_section(doc, "Investment Allocation", "06-work-allocation-index.png",
        "Squad effort estimate Apr–Jul 2026", ["MSC API 35%", "V2 UI 20%", "V3 UP 15%", "Perf 12%"])
    add_chart_section(doc, "Release Automation Impact", "07-release-automation-impact.png",
        "Release validation model", [f"~{sc['release_automation_pct']}% automated", "17 FTE → 2 FTE equivalent"])

    # ── Appendix A: Data Confidence ───────────────────────────────────────────
    doc.add_heading("Appendix A — Data Confidence", level=1)
    doc.add_paragraph(dc.get("key_distinction", ""))
    for point in dc.get("leadership_talking_points", []):
        doc.add_paragraph(point, style="List Bullet")

    # ── Appendix B: Roadmap & Asks ────────────────────────────────────────────
    doc.add_heading("Appendix B — Roadmap & Leadership Asks", level=1)
    for p in ["Mobile 2 GitLab nightly (QA-1405)", "Mobile 1 master suite completion",
              "MSC enrollment API automation", "CSR Actions nightly expansion", "Automated monthly dashboard"]:
        doc.add_paragraph(p, style="List Bullet")
    doc.add_paragraph(
        "Leadership asks: (1) engage AM Squad at SDLC start for roadmap visibility; "
        "(2) administrative capacity to free technical lead for architecture and AI tooling."
    )

    # ── Appendix C: References ────────────────────────────────────────────────
    doc.add_heading("Appendix C — References", level=1)
    for label, url in REFS:
        p = doc.add_paragraph(style="List Bullet")
        styled_run(p, f"{label}: ", bold=True, size=10)
        styled_run(p, url, size=10, color=TEAL)

    styled_run(doc.add_paragraph(), f"Generated {datetime.now().strftime('%B %d, %Y')}.", italic=True, size=9, color=GRAY)
    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


# ── PPTX (modern design — updates original deck only) ────────────────────────

def build_pptx(m: dict) -> None:
    sc = m["scorecard"]
    j = m["jira"]["totals"]
    dc = m.get("data_confidence", {})
    perf_base = sum(a["labels"] for a in m["perf_inventory"]["areas"])
    prs = new_presentation()
    n = 0

    slide_hero(prs, "QA Automation — AM Squad", f"Leadership Update · {SUBTITLE}", "Portfolio briefing")
    slide_kpi_dashboard(prs, n := n + 1, m)

    slide_bullets_modern(prs, n := n + 1, "Executive Headline", [
        f"{sc['gitlab_merges']} merged changes to main across 3 automation repositories",
        f"{j['story_points']:.0f} Jira story points · {j['work_items_in_sprints']} work items · {j['automation_bugs_logged']} bugs found",
        f"~{sc['release_automation_pct']}% of monthly release validations automated (17 FTE → 2 FTE)",
        "Unite MSC rescued — Mobile 2 at 25/25 endpoints, ~50% ETA savings",
        "Six parallel tracks: V2 · V3 · API/MSC · Perf · Pipeline · Standards",
    ], "Executive summary", "Overview")

    slide_chart_insight(prs, n := n + 1, "GitLab Delivery Velocity", "01-gitlab-mrs-by-month.png",
        "What this shows", [f"{sc['gitlab_merges']} total merges", "July peak: 38 merges during MSC sprint", "Separate from test case counts"],
        "GitLab MR export · Apr 1 – Aug 4, 2026", "GitLab")
    slide_chart_insight(prs, n := n + 1, "Monthly Automation Delivery", "08-monthly-automation-test-cases-added.png",
        "Period delivery", dc.get("leadership_talking_points", [])[:4],
        "Jira AMSQUAD · period delivery estimate", "Automation")
    slide_chart_insight(prs, n := n + 1, "Jira Sprint Delivery", "09-jira-story-points-by-sprint.png",
        "Sprint outcomes", [f"{j['work_items_in_sprints']} items · {j['story_points']:.0f} SP", f"{j['automation_bugs_logged']} bugs logged"],
        "Jira AMSQUAD Sprints 26.04–26.12", "Jira")
    slide_chart_insight(prs, n := n + 1, "Automation Defects Discovered", "10-jira-automation-bugs-by-sprint.png",
        "Quality signal", ["Defects found via nightly triage", "Fed into automation bug lifecycle"],
        "Jira AMSQUAD", "Quality")

    slide_bullets_modern(prs, n := n + 1, "Beyond the Metrics", [
        "Framework architecture & canonical repo design",
        "AI-accelerated docs, Postman, TestNG boilerplate",
        "qTest master suite & automation bug lifecycle standard",
        "Pipeline/DevOps co-design & module switches",
        "Cross-team emergency support & release triage",
    ], "Additional value not captured in merge or test-count metrics", "Value", note="These investments are not fully captured in delivery charts.")

    slide_section_modern(prs, "V2 Legacy UI Automation", "Jenkins Stage1 nightly · built since Q2 2025")
    slide_chart_insight(prs, n := n + 1, "V2 Stage1 Nightly Snapshot", "04-v2-regression-by-module.png",
        "What this counts", [
            f"{sc['v2_nightly_methods']} methods — Stage1 primary nightly only (Aug 4)",
            "Built since Q2 2025 — not Apr–Aug alone",
            "Excludes smoke, Stage 2/5, +33 CSR Actions pending wire",
            "CSR maintenance added Apr–Jul on earlier foundation",
        ],
        "Jenkins STAGE1-Daily-Unite-Prime-Regression · Aug 4", "V2",
        footnote="Additional V2 coverage: Stage 5 smoke, Stage 2 smoke, on-demand fast smoke, +33 CSR Actions.")
    slide_bullets_modern(prs, n := n + 1, "V2 Business Value", [
        "Core member journeys covered nightly before release",
        "CSR maintenance gaps closed — high-risk flows automated",
        "Enrollments/login triage separates env vs product defects",
    ], "V2 value", "V2")

    slide_section_modern(prs, "V3 Universal Platform", "GitLab CI nightly · built since Q2 2025")
    slide_chart_insight(prs, n := n + 1, "V3 Stage1 Nightly Snapshot", "11-v3-regression-by-module.png",
        "What this counts", [
            f"{sc['v3_nightly_methods']} methods — GitLab Stage1 nightly only (Aug 4)",
            "Framework + CI/CD accumulated over ~1 year",
            "UE 303 = scenarios × plan/traunch matrix",
            "Entity suites expanding — not all in Aug 4 log",
        ],
        "GitLab scheduled_regression_job · Aug 4", "V3",
        footnote="Additional V3 coverage: Stage 5 smoke (UE + IDP), integration profiles, Entity track — not in 442.")
    slide_bullets_modern(prs, n := n + 1, "V3 Delivery Highlights", [
        "Entity registration/login suites expanded",
        "IDP open-account and member withdrawal regression",
        "Flaky-test stabilization and web registration flows",
        "Stage 5 smoke suites for UE + IDP",
    ], "V3 highlights", "V3")

    slide_section_modern(prs, "API / Unite MSC", "Rescued · accelerated")
    slide_chart_insight(prs, n := n + 1, "MSC Endpoint Coverage", "05-unite-msc-coverage.png",
        "MSC status", [f"M2 {sc['msc_m2_endpoints']}", f"M1 {sc['msc_m1_core']}", "~50% ETA savings"],
        "api-test-automation", "MSC")
    slide_chart_insight(prs, n := n + 1, "API Module Breakdown", "13-api-regression-by-module.png",
        "M1 categories", ["Auth, profile, biometric, device, bank", "Master suite in progress"],
        "api-test-automation", "API")

    slide_section_modern(prs, "Performance Testing", "Labels × plans model")
    slide_chart_insight(prs, n := n + 1, "Perf Test Case Inventory", "12-perf-test-case-inventory.png",
        "Why the numbers look large", [
            f"{perf_base} base transaction flows → {sc['perf_test_cases']} plan-expanded cases",
            "Bars = flows; line = cases after × plan matrix",
            "4 Jenkins scenarios schedule these permutations",
            "IDP alone: 15 labels × 7 plans = 105 cases",
        ],
        "Perf inventory", "Perf")
    slide_bullets_modern(prs, n := n + 1, "Performance Value", [
        "Repeatable baselines — no manual re-run each release",
        "Evidence for platform team pre/post patch comparison",
        "Department perf DoD and BlazeMeter reporting standard",
    ], "Perf value", "Perf")

    slide_chart_insight(prs, n := n + 1, "Investment Allocation", "06-work-allocation-index.png",
        "Effort mix", ["MSC API 35%", "V2 20%", "V3 15%", "Perf 12%"],
        "Squad estimate Apr–Jul", "Portfolio")
    slide_chart_insight(prs, n := n + 1, "Release Automation Impact", "07-release-automation-impact.png",
        "Business value", [f"~{sc['release_automation_pct']}% automated", "17 FTE → 2 FTE equivalent"],
        "Release validation", "Value")

    slide_bullets_modern(prs, n := n + 1, "AI Acceleration", [
        "MSC migration agents — Postman, docs, TestNG boilerplate",
        "Automation bug lifecycle — Cursor skill + standardized reporting",
        "Chart & metrics generators — reproducible leadership packs",
        "Coverage intelligence foundation for monthly dashboard",
    ], "AI tooling", "AI")
    slide_bullets_modern(prs, n := n + 1, "Roadmap Q3–Q4 2026", [
        "P0: MSC GitLab nightly · M1 master suite",
        "P1: MSC enrollment API · CSR Actions nightly expansion",
        "P2: Entity V3 nightly · automated monthly dashboard",
    ], "Roadmap", "Roadmap")
    slide_bullets_modern(prs, n := n + 1, "Leadership Asks", [
        "Roadmap visibility — AM Squad at SDLC start, not sign-off deadline",
        "Administrative capacity — free lead for architecture & AI tooling",
        "Recommended: 30-minute live walkthrough",
    ], "Asks", "Asks")

    slide_close(prs, "QA Automation AM Squad", "Delivering across six automation channels · Team since Q2 2025")
    prs.save(PPTX_OUT)
    print(f"Wrote {PPTX_OUT}")


def main() -> None:
    for path in (DOCX_OUT, PPTX_OUT):
        if path.exists():
            path.unlink()
    m = load_metrics()
    build_pptx(m)
    build_docx(m)


if __name__ == "__main__":
    main()
