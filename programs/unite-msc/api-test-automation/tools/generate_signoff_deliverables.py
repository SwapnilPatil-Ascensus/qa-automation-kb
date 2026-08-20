#!/usr/bin/env python3
"""Generate Mobile 1 & Mobile 2 sign-off CSV, markdown, metrics, and DOCX deliverables."""

from __future__ import annotations

import csv
import re
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_SIGNOFF = ROOT / "docs" / "06-coverage" / "signoff"
OUT_MAPPINGS = ROOT / "mappings"
OUT_METRICS = ROOT / "docs" / "06-coverage"
ASSETS = ROOT / "docs" / "06-coverage" / "signoff" / "_assets"
AUTOMATION = Path(r"C:\Workspace\GitLab\api-test-automation")

SIGNOFF_DATE = date.today().strftime("%B %d, %Y")
SIGNOFF_DATE_ISO = date.today().isoformat()
BASELINE_COMMIT = "main @ August 2026"

NAVY = RGBColor(0x00, 0x30, 0x57)
TEAL = RGBColor(0x00, 0x7A, 0x8C)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xE6, 0x51, 0x00)
RED = RGBColor(0xC6, 0x28, 0x28)
GRAY = RGBColor(0x61, 0x61, 0x61)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "E8EEF4"
GREEN_BG = "E8F5E9"
AMBER_BG = "FFF3E0"
RED_BG = "FFEBEE"

# Mobile 1 — canonical register (verified against api-test-automation main)
M1_ENDPOINTS = [
    ("M1-01", "POST", "/mobile1api/v1/mobilemembersession", "Authentication", "Mobile1AuthenticationTest", "getValidMemberSession", "integration+regression", "okdirect,nmdirect", "Legacy: unite-mobile1 session POST; Postman: Mobile1 Member Session", "migrated", "L1-L4 lean", "Y", "N", "Auth foundation — 3 test legs (OKD session, OKD username, NMD session)"),
    ("M1-02", "GET", "/mobile1api/v1/mobilememberusername", "Authentication", "MobileMemberUsernameGetRequestTest", "getMobileMemberUsername_returnsUsername", "auth-regression", "okdirect", "Password-change prerequisite; Postman MSC", "migrated", "L1-L4", "Y", "N", "Runs in mobile1-auth-regression"),
    ("M1-03", "GET", "/mobile1api/v1/mobileowner", "Profile", "MobileOwnerRequestTest", "getMobileOwner_returnsOwner", "profileowner-*", "okdirect,nmdirect", "Legacy profile GET; Postman mobileowner", "migrated", "L1-L4 POJO", "Y", "N", ""),
    ("M1-04", "GET", "/mobile1api/v1/mobileOwnerMenu", "Profile", "MobileOwnerMenuRequestTest", "getMobileOwnerMenu_returnsOwnerMenu", "profileowner-*", "okdirect,nmdirect", "Legacy owner menu", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-05", "GET", "/mobile1api/v1/mobileprofilemenu", "Profile", "MobileProfileMenuRequestTest", "getMobileProfileMenu_returnsProfileMenu", "profileowner-*", "okdirect,nmdirect", "Legacy profile menu", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-06", "PUT", "/mobile1api/v1/mobileowner", "Profile", "MobileOwnerPutRequestTest", "putMobileOwner_updatesOwnerProfile_returnsOk", "mobile1-smoke", "okdirect", "Mutating owner update — smoke only", "migrated", "L1-L4", "N", "Y", "Destructive — smoke suite"),
    ("M1-07", "GET", "/mobile1api/v1/mobilebeneficiaryByExt/{ext}", "Beneficiary", "MobileBeneficiaryByExtRequestTest", "getMobileBeneficiaryByExt_returnsBeneficiary", "beneficiary-*", "okdirect", "Legacy beneficiary lookup", "migrated", "L1-L4", "Y", "N", "OKD only in suites"),
    ("M1-08", "POST", "/mobile1api/v1/mobilecloseaccount/{ext}", "Beneficiary", "MobileCloseAccountPostRequestTest", "postMobileCloseAccount_preClosureCheck", "beneficiary-regression", "okdirect", "Pre-closure eligibility check", "migrated", "L1-L4", "Y", "N", "preClosureCheck=true"),
    ("M1-09", "POST", "/mobile1api/v1/mobilecloseaccount/{ext}", "Beneficiary", "MobileActualCloseAccountPostRequestTest", "postMobileActualCloseAccount_closesAccount", "mobile1-smoke", "nmdirect", "Actual account close — destructive", "migrated", "L1-L4", "N", "Y", "Smoke — NMD only"),
    ("M1-10", "GET", "/mobile1api/v1/mobilebankinfobyroutingnum/{routingNum}", "Bank Info", "MobileBankInfoByRoutingNumRequestTest", "getMobileBankInfoByRoutingNum_returnsBankInfo", "bankinfo-*", "okdirect", "Aligned with Mobile2 banks routing 011000138", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-11", "POST", "/mobile1api/v1/mobilememberbiometric", "Biometric", "MobileMemberBiometricPostRequestTest", "postMobileMemberBiometric_enrollsToken", "memberbiometric-*", "okdirect", "Legacy biometric enroll", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-12", "GET", "/mobile1api/v1/mobilememberbiometric", "Biometric", "MobileMemberBiometricGetRequestTest", "getMobileMemberBiometric_returnsBiometricToken", "memberbiometric-*", "okdirect", "GET after POST enroll", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-13", "DELETE", "/mobile1api/v1/mobilememberbiometric", "Biometric", "MobileMemberBiometricDeleteRequestTest", "deleteMobileMemberBiometric_removesBiometricToken", "mobile1-smoke", "okdirect", "Destructive cleanup", "migrated", "L1-L4", "N", "Y", "Smoke only"),
    ("M1-14", "POST", "/mobile1api/v1/requestPhoneNumberAuthentication", "Phone Auth", "MobileRequestPhoneNumberAuthenticationPostRequestTest", "postRequestPhoneNumberAuthentication_returnsOwnerPhoneDetails", "phoneauthentication-*", "okdirect", "Postman: requestPhoneNumberAuthentication", "migrated", "L1-L4", "Y", "N", "May trigger SMS outside QC4"),
    ("M1-15", "POST", "/mobile1api/v1/mobilememberdevices", "Device", "MobileMemberDeviceRequestTest", "postMobileMemberDevice_registersDevice", "memberdevice-*", "okdirect", "Device registration — no body", "migrated", "L1-L4", "Y", "N", "5 tests in class"),
    ("M1-16", "POST", "/mobile1api/v1/mobilememberpushnotificationtokens", "Device", "MobileMemberDeviceRequestTest", "postMobileMemberPushNotificationToken_registersToken", "memberdevice-*", "okdirect", "Push token POST", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-17", "PUT", "/mobile1api/v1/mobilememberpushnotificationtokens", "Device", "MobileMemberDeviceRequestTest", "putMobileMemberPushNotificationToken_updatesToken", "memberdevice-*", "okdirect", "Push token PUT", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-18", "GET", "/mobile1api/v1/mobilememberpushnotificationtokens/deviceuuid/{deviceUuid}", "Device", "MobileMemberDeviceRequestTest", "getMobileMemberPushNotificationToken_returnsRegisteredToken", "memberdevice-*", "okdirect", "Push token GET by device UUID", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-19", "GET", "/mobile1api/v1/mobilememberdevices", "Device", "MobileMemberDeviceRequestTest", "getMobileMemberDevice_returnsRegisteredDevice", "memberdevice-*", "okdirect", "Device GET after POST", "migrated", "L1-L4", "Y", "N", ""),
    ("M1-20", "PATCH", "/mobile1api/v1/mobilemembers", "Password", "MobileChangePasswordRequestTest", "patchMobileMembers_changesPasswordAndRelogin", "mobile1-smoke", "okdirect", "Password rotation + re-login", "migrated", "L1-L4", "N", "Y", "Uses JSON user id 2"),
    ("M1-21", "POST", "/mobile1api/v1/mobilecsrasmembersession", "CSR", "MobileCsrAsMemberSessionRequestTest", "postMobileCsrAsMemberSession_returnsSession", "csrasmember-*", "okdirect", "Public endpoint; optional CSR JWT header", "migrated", "L1-L4", "Y", "N", "INVALID_CREDENTIALS without real token"),
    ("M1-22", "POST", "/mobile1api/v1/idptokenexchange", "IDP", "MobileIdpTokenExchangeRequestTest", "postIdpTokenExchange_returnsAccessToken", "memberidptoken-*", "nmdirect", "PKCE → IDP token exchange", "migrated", "L1-L4", "Y", "N", "NMD IDP branding"),
    ("M1-23", "POST", "/mobile1api/v1/mobilememberidptoken", "IDP", "MobileMemberIdpTokenRequestTest", "postMobileMemberIdpToken_returnsMemberSession", "memberidptoken-*", "nmdirect", "IDP JWT → member session", "migrated", "L1-L4", "Y", "N", "QC4 often 401 on automation JWT"),
    ("M1-24", "GET", "/mobile1api/v1/mobilemembersession/{id}", "Session", "MobileMemberSessionByIdRequestTest", "getMobileMemberSessionById_returnsSession", "membersession-smoke", "okdirect", "Session GET by id from login", "migrated", "L1-L4", "N", "Y", "Smoke — SQL user 1"),
    ("M1-25", "POST", "/mobile1api/v1/mobilemembersession/validateBiometricToken", "Session", "MobileMemberSessionValidateBiometricTokenRequestTest", "postValidateBiometricToken_returnsMemberSession", "membersession-smoke", "okdirect", "Biometric validate flow", "migrated", "L1-L4", "N", "Y", "Functional group"),
    ("M1-26", "POST", "/mobile1api/v1/mobilemembersessionpin", "Session", "MobileMemberSessionPinRequestTest", "postMobileMemberSessionPin_returnsSessionPin", "membersessionpin-*", "okdirect", "1-factor member JWT; POST only", "migrated", "L1-L4", "Y", "N", "No GET in legacy API"),
]

M2_ENDPOINTS = [
    ("M2-01", "GET", "/mobile2api/v1/mobileactivity/{ext}", "Activity", "MobileActivityRequestTest", "getMobileActivity_returnsActivitySummary", "activity-*", "okdirect,newyork", "Cucumber mobileactivity feature", "migrated-simplified", "L1-L4 lean", "Y", "N", "Master regression"),
    ("M2-02", "GET", "/mobile2api/v1/mobiletransactionhistory/{ext}", "Transactions", "MobileTransactionHistoryRequestTest", "getMobileTransactionHistory_returnsTransactions", "transactionhistory-*", "okdirect,newyork", "Legacy transaction history", "migrated-simplified", "L1-L4", "Y", "N", ""),
    ("M2-03", "GET", "/mobile2api/v1/investments/{ext}", "Investment", "MobileInvestmentRequestTest", "getMobileInvestments_returnsInvestments", "investment-*", "okdirect,newyork", "Legacy investments", "migrated-simplified", "L1-L4", "Y", "N", ""),
    ("M2-04", "GET", "/mobile2api/v1/mobilebanks", "Banks", "MobileBanksRequestTest", "getMobileBanks_filterDomesticBanks_returnsBanks", "banks-*", "okdirect", "Postman MSC banks list", "migrated", "L1-L4", "Y", "N", "OKD only"),
    ("M2-05", "GET", "/mobile2api/v1/mobilebanks/{id}", "Banks", "MobileBanksRequestTest", "getMobileBankById_returnsBank", "banks-*", "okdirect", "QA-1386", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-06", "POST", "/mobile2api/v1/mobilebanks", "Banks", "MobileBanksRequestTest", "postMobileBanks_addsDomesticBank_returnsBanks", "banks-*", "okdirect", "Domestic bank add", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-07", "PUT", "/mobile2api/v1/mobilebanks", "Banks", "MobileBanksRequestTest", "putMobileBanks_updatesDomesticBank_returnsBanks", "mobile2-smoke", "okdirect", "Destructive update", "migrated", "L1-L4", "N", "Y", "Excluded from master"),
    ("M2-08", "DELETE", "/mobile2api/v1/mobilebanks", "Banks", "MobileBanksRequestTest", "deleteMobileBanks_deletesDomesticBank_returnsBanks", "mobile2-smoke", "okdirect", "Destructive delete", "migrated", "L1-L4", "N", "Y", "Excluded from master"),
    ("M2-09", "GET", "/mobile2api/v1/content", "Content", "MobileContentRequestTest", "getContent_commonSavingTips_returnsContent", "content-*", "okdirect,newyork", "CMS content gateway", "migrated-simplified", "L1-L4", "Y", "N", ""),
    ("M2-10", "GET", "/mobile2api/v1/plans", "Plans", "MobilePlansRequestTest", "getMobilePlans_returnsPlans", "plans-*", "okdirect,newyork", "Plan list", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-11", "GET", "/mobile2api/v1/plans/{id}", "Plans", "MobilePlansRequestTest", "getMobilePlanById_returnsPlan", "plans-*", "okdirect,newyork", "Plan by branding id", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-12", "GET", "/mobile2api/v1/mobilecontribution", "Contribution", "MobileContributionRequestTest", "getMobileContribution_returnsContributionOptions", "contribution-*", "okdirect,newyork", "Contribution options", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-13", "GET", "/mobile2api/v1/mobilecontributioncheck", "Contribution", "MobileContributionCheckRequestTest", "getMobileContributionCheck_returnsShowContributionFlag", "contribution-*", "okdirect,newyork", "Show contribution flag", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-14", "GET", "/mobile2api/v1/mobilecontribution/{ext}/{id}", "Contribution", "MobileContributionDetailRequestTest", "getMobileContributionById_returnsRecurringContribution", "contribution-*", "okdirect,newyork", "Dynamic SQL fixture id 472560", "migrated", "L1-L4", "Y", "N", "Stage1 fixture caveat"),
    ("M2-15", "POST", "/mobile2api/v1/mobilecontribution", "Contribution", "MobileContributionPostRequestTest", "postMobileContribution_createsRecurringContribution", "contribution-*", "okdirect,newyork", "Create recurring contribution", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-16", "PUT", "/mobile2api/v1/mobilecontribution/{ext}/{id}", "Contribution", "MobileContributionPutRequestTest", "putMobileContributionById_updatesRecurringContribution", "contribution-*", "okdirect,newyork", "Update recurring", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-17", "DELETE", "/mobile2api/v1/mobilecontribution/{ext}/{id}", "Contribution", "MobileContributionDeleteRequestTest", "deleteMobileContributionById_removesAutomationOwnedContribution", "contribution-regression", "okdirect", "Automation-owned DELETE", "migrated", "L1-L4", "N", "N", "OKD module only; not master"),
    ("M2-18", "GET", "/mobile2api/v1/mobiledashboard", "Dashboard", "MobileDashboardRequestTest", "getMobileDashboard", "dashboard-*", "okdirect,newyork", "First vertical slice; 8-test→1 lean", "migrated-simplified", "L1-L4 lean", "Y", "N", "Reference case study"),
    ("M2-19", "GET", "/mobile2api/v1/mobileytdsummary/{ext}", "Dashboard", "MobileYtdSummaryRequestTest", "getMobileYtdSummary_returnsYtdContributionSummary", "dashboard-*,mobile2-smoke", "okdirect,newyork", "YTD contribution summary", "migrated", "L1-L4", "Y", "Y", "Master + smoke"),
    ("M2-20", "GET", "/mobile2api/v1/mobilemembers/{planId}/{username}", "Harness", "MobileMembersRequestTest", "getMobileMembers_returnsMemberForHarness", "mobile2-smoke", "okdirect", "Acceptance harness — 401 with member JWT by design", "excluded", "L1 only", "N", "Y", "Out of business numerator"),
    ("M2-21", "GET", "/mobile2api/v1/mobilebalancetrend/{ext}", "Balance", "MobileBalanceTrendRequestTest", "getMobileBalanceTrend_returnsBalanceTrend", "balancetrend-*", "okdirect,newyork", "Balance trend chart data", "migrated-simplified", "L1-L4", "Y", "N", ""),
    ("M2-22", "GET", "/mobile2api/v1/mobileperformance/{ext}", "Performance", "MobilePerformanceRequestTest", "getMobilePerformance_returnsPerformance", "balancetrend-*", "okdirect,newyork", "Performance metrics", "migrated-simplified", "L1-L4", "Y", "N", ""),
    ("M2-23", "GET", "/mobile2api/v1/mobilestackup/{planId}", "Stackup", "MobileStackupRequestTest", "getMobileStackup_returnsStackup", "balancetrend-*,mobile2-smoke", "okdirect,newyork,nmdirect", "Strict stackup in smoke (NMD)", "migrated-simplified", "L1-L4", "Y", "Y", "Duplicate class in 2 packages"),
    ("M2-24", "GET", "/mobile2api/v1/mobileugift", "UGift", "MobileUgiftRequestTest", "getMobileUgift_returnsUgiftPage", "ugift-*", "okdirect,newyork", "UGift page GET", "migrated", "L1-L4", "Y", "N", ""),
    ("M2-25", "PATCH", "/mobile2api/v1/mobileugift/{ext}", "UGift", "MobileUgiftRequestTest", "patchMobileUgift_assignsUgiftId", "ugift-*", "okdirect,newyork", "Idempotent ugift assign", "migrated", "L1-L4", "Y", "N", ""),
]

REGISTER_HEADERS = [
    "endpoint_id", "http_method", "path", "feature_area", "test_class", "test_method",
    "suite_profiles", "branding_support", "legacy_source", "migration_status", "validation_layers",
    "module_regression", "smoke_suite", "notes",
]


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 9) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def style_table_header(row) -> None:
    for cell in row.cells:
        shade_cell(cell, "003057")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"


def add_page_number_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup_doc(doc: Document, header_text: str, footer_left: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(header_text)
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = fp.add_run(footer_left)
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY
    r2.font.name = "Calibri"
    # page number right-aligned in same footer row via tab — use second paragraph
    pp = section.footer.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr = pp.add_run(f"Page ")
    pr.font.size = Pt(8)
    pr.font.color.rgb = GRAY
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    pr2 = pp.add_run()
    pr2._r.append(fld_char1)
    pr2._r.append(instr)
    pr2._r.append(fld_char2)


def add_heading(doc: Document, title: str, level: int = 1) -> None:
    h = doc.add_heading(title, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"


def add_callout(doc: Document, text: str, bg: str = GREEN_BG, bold: bool = True) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    set_cell_text(cell, text, size=11, bold=bold, color=NAVY)


def write_register_csv(path: Path, module: str, rows: list) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(REGISTER_HEADERS)
        for row in rows:
            w.writerow([module] + list(row))


def write_combined_register() -> None:
    path = OUT_MAPPINGS / "endpoint-signoff-register.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(REGISTER_HEADERS)
        for row in M1_ENDPOINTS:
            w.writerow(["mobile1"] + list(row))
        for row in M2_ENDPOINTS:
            w.writerow(["mobile2"] + list(row))


def count_tests_java(module: str) -> int:
    base = AUTOMATION / "mobile" / module / "src" / "test" / "java"
    if not base.exists():
        return 0
    return len(re.findall(r"@Test\s*\(", "\n".join(p.read_text(encoding="utf-8", errors="ignore") for p in base.rglob("*.java"))))


def chart_bar(path: Path, title: str, labels: list, values: list, colors: list) -> None:
    fig, ax = plt.subplots(figsize=(8, 3.5))
    bars = ax.barh(labels, values, color=colors, height=0.5)
    ax.set_title(title, fontsize=13, fontweight="bold", color="#003057", pad=12)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height() / 2, str(val), va="center", fontsize=10, color="#334155")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.patch.set_facecolor("white")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_mobile_docx(module: str, title: str, endpoints: list, metrics: dict, chart_path: Path, out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    setup_doc(doc, f"{title}  |  Internal Sign-Off", f"QA Automation — AMSQUAD  |  {SIGNOFF_DATE}")

    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Formal Sign-Off & Handover Package")
    r2.font.size = Pt(20)
    r2.font.color.rgb = TEAL
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = st.add_run("STATUS: COMPLETE")
    r3.bold = True
    r3.font.size = Pt(18)
    r3.font.color.rgb = GREEN
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [f"Sign-off date: {SIGNOFF_DATE}", "Prepared by: QA Automation (AMSQUAD)", "Classification: Internal — Handover to Engineering & Support"]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    doc.add_page_break()

    add_heading(doc, "Document Control", level=1)
    ctrl = doc.add_table(rows=8, cols=2)
    ctrl.style = "Table Grid"
    items = [
        ("Document title", f"{title} — Sign-Off & Handover"),
        ("Version", "2.0 — Final"),
        ("Sign-off date", SIGNOFF_DATE),
        ("Evidence baseline", BASELINE_COMMIT),
        ("Automation repo", "ascensus-gs/products/depot/qa-automation/api-test-automation"),
        ("Knowledge base", "qa-automation-kb/programs/unite-msc/api-test-automation"),
        ("Endpoints automated", str(metrics["automated"])),
        ("Sign-off determination", "COMPLETE"),
    ]
    for i, (k, v) in enumerate(items):
        shade_cell(ctrl.rows[i].cells[0], LIGHT_BG)
        set_cell_text(ctrl.rows[i].cells[0], k, bold=True, color=NAVY)
        set_cell_text(ctrl.rows[i].cells[1], v, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", level=1)
    add_callout(doc, metrics["summary_callout"])
    doc.add_paragraph(metrics["summary_body"])
    if chart_path.exists():
        doc.add_picture(str(chart_path), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    kpi = doc.add_table(rows=6, cols=2)
    kpi.style = "Table Grid"
    for i, (k, v) in enumerate(metrics["kpis"]):
        shade_cell(kpi.rows[i].cells[0], LIGHT_BG)
        set_cell_text(kpi.rows[i].cells[0], k, bold=True, color=NAVY)
        bg = GREEN_BG if i == len(metrics["kpis"]) - 1 else "FFFFFF"
        shade_cell(kpi.rows[i].cells[1], bg)
        set_cell_text(kpi.rows[i].cells[1], v, bold=(i == len(metrics["kpis"]) - 1), color=GREEN if i == len(metrics["kpis"]) - 1 else GRAY)
    doc.add_page_break()

    add_heading(doc, "2. Migration Summary", level=1)
    for item in metrics["migration_bullets"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    add_heading(doc, "3. Code Coverage & Test Metrics", level=1)
    doc.add_paragraph(
        "Note: JaCoCo line coverage applies to framework Java sources, not API endpoint coverage. "
        "Endpoint coverage is measured as automated business endpoints / in-scope denominator.")
    met = doc.add_table(rows=1 + len(metrics["coverage_rows"]), cols=3)
    met.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value", "Notes"]):
        set_cell_text(met.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(met.rows[0])
    for ri, row in enumerate(metrics["coverage_rows"], start=1):
        for ci, val in enumerate(row):
            shade_cell(met.rows[ri].cells[ci], LIGHT_BG if ci == 0 else "FFFFFF")
            set_cell_text(met.rows[ri].cells[ci], val, bold=(ci == 0), color=NAVY if ci == 0 else GRAY)
    doc.add_page_break()

    add_heading(doc, "4. Endpoint Register (Full Mapping)", level=1)
    reg = doc.add_table(rows=1 + len(endpoints), cols=8)
    reg.style = "Table Grid"
    headers = ["ID", "Method", "Path", "Area", "Test Class", "Migration", "Master", "Notes"]
    for i, h in enumerate(headers):
        set_cell_text(reg.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(reg.rows[0])
    status_bg = {"migrated": GREEN_BG, "migrated-simplified": "E3F2FD", "excluded": AMBER_BG}
    for ri, row in enumerate(endpoints, start=1):
        eid, method, path, area, cls, method_name, suite, branding, legacy, mig, layers, mod_reg, smoke, notes = row
        master = "Y" if "master" in suite else ("N" if smoke == "Y" else mod_reg)
        vals = [eid, method, path, area, cls.split(".")[-1] if "." not in cls else cls, mig, master, notes[:80]]
        for ci, val in enumerate(vals):
            bg = status_bg.get(mig, LIGHT_BG if ci == 0 else "FFFFFF")
            if ci == 5:
                bg = status_bg.get(mig, "FFFFFF")
            shade_cell(reg.rows[ri].cells[ci], bg)
            set_cell_text(reg.rows[ri].cells[ci], val, size=7, color=NAVY if ci in (0, 5) else GRAY)
    doc.add_page_break()

    add_heading(doc, "5. How to Run & Troubleshoot", level=1)
    for item in metrics["run_bullets"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    add_heading(doc, "6. Known Limitations", level=1)
    for item in metrics["limitations"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    add_heading(doc, "7. Sign-Off Approval", level=1)
    add_callout(doc, f"{title} API automation is COMPLETE for the defined scope effective {SIGNOFF_DATE}.")
    sig = doc.add_table(rows=5, cols=4)
    sig.style = "Table Grid"
    for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
        set_cell_text(sig.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(sig.rows[0])
    for ri, role in enumerate(["QA Automation Lead", "Program Lead / SME", "Engineering Lead", "Support / Operations"], start=1):
        shade_cell(sig.rows[ri].cells[0], LIGHT_BG)
        set_cell_text(sig.rows[ri].cells[0], role, bold=True, color=NAVY)
        for ci in range(1, 4):
            set_cell_text(sig.rows[ri].cells[ci], "")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Created: {out_path}")


def write_markdown_summary(path_obj: Path, module: str, title: str, endpoints: list, metrics: dict) -> None:
    path_obj.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        f"# {title} — Sign-Off Summary",
        "",
        f"**Sign-off date:** {SIGNOFF_DATE_ISO}  ",
        f"**Status:** COMPLETE  ",
        f"**KB path:** `programs/unite-msc/api-test-automation/docs/06-coverage/signoff/`",
        "",
        "## Executive metrics",
        "",
        "| Metric | Value |",
        "|--------|-------|",
    ]
    for k, v in metrics["kpis"]:
        lines.append(f"| {k} | {v} |")
    lines += ["", "## Endpoint register", "", "| ID | Method | Path | Test | Master |", "|----|--------|------|------|--------|"]
    for row in endpoints:
        eid, method, path, area, cls, tmethod, suite, branding, legacy, mig, layers, mod_reg, smoke, notes = row
        master = "Y" if "master" in suite else ("N" if smoke == "Y" else mod_reg)
        lines.append(f"| {eid} | {method} | `{path}` | `{cls}.{tmethod}` | {master} |")
    lines += [
        "",
        "## DOCX deliverable",
        "",
        f"- [`{path_obj.stem.replace('-summary', '')}.docx`](./{path_obj.stem.replace('-summary', '')}.docx)",
        "",
        "## Evidence",
        "",
        "- Regression logs: `evidence/regression-runs/`",
        "- Mapping CSV: `mappings/endpoint-signoff-register.csv`",
        "",
    ]
    path_obj.write_text("\n".join(lines), encoding="utf-8")
    print(f"Created: {path_obj}")


def main() -> None:
    OUT_SIGNOFF.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)

    write_register_csv(OUT_MAPPINGS / "mobile1-endpoint-current-state.csv", "mobile1", M1_ENDPOINTS)
    write_register_csv(OUT_MAPPINGS / "mobile2-endpoint-current-state.csv", "mobile2", M2_ENDPOINTS)
    write_combined_register()

    m1_tests = count_tests_java("mobile1")
    m2_tests = count_tests_java("mobile2")
    m1_automated = len(M1_ENDPOINTS)
    m2_business = len([e for e in M2_ENDPOINTS if e[9] != "excluded"])
    m2_total = len(M2_ENDPOINTS)
    m2_pct = round(100 * m2_business / 25, 1)  # 25 documented business denominator

    chart_bar(ASSETS / "mobile1_coverage.png", "Mobile 1 — Automated Endpoint Operations",
                ["Automated operations", "@Test methods"], [m1_automated, m1_tests], ["#007A8C", "#003057"])
    chart_bar(ASSETS / "mobile2_coverage.png", f"Mobile 2 — Business Endpoint Coverage ({m2_business}/25 = {m2_pct}%)",
                ["In-scope automated", "Excluded harness"], [m2_business, 1], ["#2E7D32", "#CBD5E1"])

    m1_metrics = {
        "automated": m1_automated,
        "summary_callout": f"Mobile 1 API automation is COMPLETE: {m1_automated} endpoint operations automated with {m1_tests} TestNG test methods across module regression and smoke suites.",
        "summary_body": (
            "Canonical TestNG automation replaces legacy unite-mobile1 Cucumber/WAR tests. "
            "Coverage spans authentication, profile, beneficiary, bank info, biometric, device/push tokens, "
            "phone authentication, password change, CSR-as-member, IDP token flows, and member session operations. "
            "Shared auth uses MobileServerClient via MobileBaseRequestTest and SQL fixture get.mobile.auth.user."
        ),
        "kpis": [
            ("Endpoint operations automated", str(m1_automated)),
            ("TestNG @Test methods", str(m1_tests)),
            ("Module regression suite XMLs", "15+"),
            ("Maven profiles", "30+"),
            ("Validation layers delivered", "L1–L4 (lean)"),
            ("Sign-off determination", "COMPLETE"),
        ],
        "migration_bullets": [
            "Legacy source: unite-mobile1 Cucumber features + Postman MSC Mobile Endpoints collection",
            "New location: api-test-automation/mobile/mobile1/",
            "Pattern: MobileBaseRequestTest + POJO convertToPOJO + module/integration/regression suite XML",
            "Auth consolidated to jsonapi-auth MobileServerClient (INFI-8078)",
            "Destructive flows isolated to mobile1-smoke (password, owner PUT, close account, biometric DELETE)",
            "IDP flows (nmdirect): idptokenexchange + mobilememberidptoken — QC4 may return 401 on automation JWT",
        ],
        "coverage_rows": [
            ("Endpoint operations", str(m1_automated), "Distinct HTTP method + path pairs with canonical tests"),
            ("Test methods", str(m1_tests), "Includes multi-branding legs in auth regression"),
            ("Module suites passing (QC4)", "See evidence/regression-runs/", "Per-profile CSV Jul 2026"),
            ("SQL validation (L5)", "Deferred", "No guessed DB queries — api-validation KB for future"),
            ("JaCoCo (framework code)", "N/A this package", "Endpoint coverage is primary sign-off metric"),
        ],
        "run_bullets": [
            "Parent: mvn -f mobile/pom.xml clean install -DskipTests",
            "Auth regression: mvn -f mobile/mobile1/pom.xml clean test \"-Pacceptance-qc4,mobile1-auth-regression\" \"-Denvironment.properties=qc4.properties\"",
            "Module example: mvn -f mobile/mobile1/pom.xml clean test \"-Pmobile1-memberdevice-regression\" \"-Denvironment.properties=qc4.properties\"",
            "Smoke (destructive): mvn -f mobile/mobile1/pom.xml clean test \"-Pmobile1-smoke\" \"-Denvironment.properties=qc4.properties\"",
            "Report: mobile/mobile1/target/mobile-ms-report/index.html (module suites with listener)",
            "Troubleshooting: qa-automation-kb/.../docs/02-daily-usage/04-troubleshooting.md",
        ],
        "limitations": [
            "CSR-as-member requires optional -Dcsr-as-member-jwt-token for HTTP 200 with real session",
            "IDP mobilememberidptoken: use -Didp-jwt-token= for live 200 on QC4",
            "Phone authentication may trigger SMS — pin verification out of scope",
            "L5 SQL API–DB reconciliation deferred to api-validation program",
        ],
    }

    m2_metrics = {
        "automated": m2_business,
        "summary_callout": f"Mobile 2 API automation is COMPLETE: {m2_business} of 25 documented business endpoints ({m2_pct}%). One harness endpoint excluded by design.",
        "summary_body": (
            "Migration from unite-mobile2 Cucumber WAR to lean TestNG suites. Master regression runs OKD (okdirect) "
            "and NYD (newyork) brandings. Dashboard intentionally simplified from 8-test regression to lean baseline. "
            f"{m2_tests} TestNG methods across 19+ RequestTest classes."
        ),
        "kpis": [
            ("Documented business endpoints", "25"),
            ("Automated (in-scope)", str(m2_business)),
            ("Coverage percentage", f"{m2_pct}%"),
            ("Master regression tests", "~40 (OKD + NYD legs)"),
            ("Validation layers", "L1–L4 lean"),
            ("Sign-off determination", "COMPLETE"),
        ],
        "migration_bullets": [
            "Legacy: unite-mobile2 Cucumber + Postman unite-mobile2 collection",
            "Canonical: api-test-automation/mobile/mobile2/",
            "Prerequisite: mvn -f mobile/mobile1/pom.xml install -DskipTests before mobile2 runs",
            "Master suite: mobile-ms-master-regression — all stable endpoints OKD + NYD",
            "Destructive: Banks PUT/DELETE and Contribution DELETE excluded from master by design",
            "Excluded: GET mobilemembers/{planId}/{username} — acceptance harness only",
        ],
        "coverage_rows": [
            ("Business endpoints (denominator)", "25", "Dinesh Mobile2.xlsx / program scope"),
            ("Automated in-scope", str(m2_business), "Excludes M2-20 harness"),
            ("QC4 module profiles pass", "20/22", "Contribution profiles had 5 failures Jul 2026 — see evidence logs"),
            ("Master regression", "40 tests expected", "okdirect + newyork branding legs"),
            ("L5 SQL validation", "Deferred", "api-validation/mappings YAML exists for future"),
        ],
        "run_bullets": [
            "mvn -f mobile/mobile1/pom.xml install -DskipTests  # always first",
            "mvn -f mobile/mobile2/pom.xml test \"-Pmobile-ms-master-regression,acceptance-qc4\" \"-Dmobile.ms.report.environment=QC4\"",
            "Stage1: add acceptance-stage1 + <COMPUTERNAME>.properties host file",
            "Batch QC4: programs/unite-msc/api-test-automation/scripts/run-qc4-all-suites.ps1",
            "Nexus/GitHub: docs/04-pipelines/02-github-actions-nexus-pipeline.md",
        ],
        "limitations": [
            "Contribution module: 5 failing legs in Jul 2026 QC4 batch — investigate fixture/data",
            "Stage1 contribution detail/PUT may 401 without fixture refresh",
            "Banks module: okdirect only — not run under NMD/NYD in master",
            "Dashboard lean test defers 6 of 8 legacy assertion groups — documented in dashboard coverage matrix",
        ],
    }

    build_mobile_docx(
        "mobile1", "Mobile 1 API Automation", M1_ENDPOINTS, m1_metrics,
        ASSETS / "mobile1_coverage.png",
        OUT_SIGNOFF / "Mobile-1-API-Automation-Sign-Off.docx",
    )
    build_mobile_docx(
        "mobile2", "Mobile 2 API Automation", M2_ENDPOINTS, m2_metrics,
        ASSETS / "mobile2_coverage.png",
        OUT_SIGNOFF / "Mobile-2-API-Automation-Sign-Off.docx",
    )

    write_markdown_summary(OUT_SIGNOFF / "mobile1-signoff-summary.md", "mobile1", "Mobile 1 API Automation", M1_ENDPOINTS, m1_metrics)
    write_markdown_summary(OUT_SIGNOFF / "mobile2-signoff-summary.md", "mobile2", "Mobile 2 API Automation", M2_ENDPOINTS, m2_metrics)

    metrics_md = OUT_METRICS / "05-code-coverage-metrics.md"
    metrics_md.write_text(
        f"""# Code Coverage & Test Metrics — Mobile 1 & Mobile 2

**As of:** {SIGNOFF_DATE_ISO}

## Summary

| Module | Endpoint ops / business endpoints | @Test methods | Sign-off |
|--------|-----------------------------------|---------------|----------|
| Mobile 1 | {m1_automated} operations | {m1_tests} | COMPLETE |
| Mobile 2 | {m2_business}/25 business ({m2_pct}%) | {m2_tests} | COMPLETE |

## QC4 execution snapshot (Jul 2026)

See `evidence/regression-runs/qc4-module-suites-results.csv`:
- Mobile 2: 20/22 module profiles PASS
- Contribution integration/regression: 5 failures (investigate)

## Deliverables

| Artifact | Path |
|----------|------|
| Combined endpoint register | `mappings/endpoint-signoff-register.csv` |
| Mobile 1 sign-off DOCX | `docs/06-coverage/signoff/Mobile-1-API-Automation-Sign-Off.docx` |
| Mobile 2 sign-off DOCX | `docs/06-coverage/signoff/Mobile-2-API-Automation-Sign-Off.docx` |

## Definition

- **Endpoint coverage** = canonical TestNG test exists for HTTP method + path
- **JaCoCo** = Java line coverage on framework code (separate CI gate program in government-savings-assessment)
""",
        encoding="utf-8",
    )
    print(f"Created: {metrics_md}")
    print("Done.")


if __name__ == "__main__":
    main()
