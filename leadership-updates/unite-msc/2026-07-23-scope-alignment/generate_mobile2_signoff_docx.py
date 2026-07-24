#!/usr/bin/env python3
"""Generate Mobile 2 API Automation formal sign-off document (standalone — no KB paths)."""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent
ASSETS = OUT_DIR / "_assets"
DOCX_OUT = OUT_DIR / "_output" / "Mobile-2-API-Automation-Sign-Off-2026-07-23.docx"

SIGNOFF_DATE = "July 23, 2026"
BASELINE_COMMIT = "cee0de9"
PRIOR_SIGNOFF_DATE = "July 14, 2026"
PRIOR_SIGNOFF_COMMIT = "7ccaf46"

M2_BUSINESS_TOTAL = 25
M2_IMPLEMENTED = 24
M2_PCT = 96.0
EXCLUDED_PATH = "GET /mobile2api/v1/mobilemembers/{planId}/{username}"

NAVY = RGBColor(0x00, 0x30, 0x57)
TEAL = RGBColor(0x00, 0x7A, 0x8C)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xE6, 0x51, 0x00)
GRAY = RGBColor(0x61, 0x61, 0x61)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "E8EEF4"
GREEN_BG = "E8F5E9"
AMBER_BG = "FFF3E0"

# Verified endpoint register (Jul 23, 2026)
ENDPOINTS = [
    ("M2-01", "GET", "/mobile2api/v1/mobileactivity/{ext}", "Activity", "Complete", "Master regression"),
    ("M2-02", "GET", "/mobile2api/v1/mobiletransactionhistory/{ext}", "Transactions", "Complete", "Master regression"),
    ("M2-03", "GET", "/mobile2api/v1/investments/{ext}", "Investment", "Complete", "Master regression"),
    ("M2-04", "GET", "/mobile2api/v1/mobilebanks", "Banks", "Complete", "Master regression"),
    ("M2-05", "GET", "/mobile2api/v1/mobilebanks/{id}", "Banks", "Complete", "Master regression"),
    ("M2-06", "POST", "/mobile2api/v1/mobilebanks", "Banks", "Complete", "Master regression"),
    ("M2-07", "PUT", "/mobile2api/v1/mobilebanks", "Banks", "Complete", "Smoke / module only (destructive)"),
    ("M2-08", "DELETE", "/mobile2api/v1/mobilebanks", "Banks", "Complete", "Smoke / module only (destructive)"),
    ("M2-09", "GET", "/mobile2api/v1/content", "Content", "Complete", "Master regression"),
    ("M2-10", "GET", "/mobile2api/v1/plans", "Plans", "Complete", "Master regression"),
    ("M2-11", "GET", "/mobile2api/v1/plans/{id}", "Plans", "Complete", "Master regression"),
    ("M2-12", "GET", "/mobile2api/v1/mobilecontribution", "Contribution", "Complete", "Master regression"),
    ("M2-13", "GET", "/mobile2api/v1/mobilecontributioncheck", "Contribution", "Complete", "Master regression"),
    ("M2-14", "GET", "/mobile2api/v1/mobilecontribution/{ext}/{id}", "Contribution", "Complete", "Master regression; Stage 1 fixture note"),
    ("M2-15", "POST", "/mobile2api/v1/mobilecontribution", "Contribution", "Complete", "Master regression"),
    ("M2-16", "PUT", "/mobile2api/v1/mobilecontribution/{ext}/{id}", "Contribution", "Complete", "Master regression; Stage 1 fixture note"),
    ("M2-17", "DELETE", "/mobile2api/v1/mobilecontribution/{ext}/{id}", "Contribution", "Complete", "Module only (destructive)"),
    ("M2-18", "GET", "/mobile2api/v1/mobiledashboard", "Dashboard", "Complete", "Master regression"),
    ("M2-19", "GET", "/mobile2api/v1/mobileytdsummary/{ext}", "Dashboard", "Complete", "Master + smoke regression"),
    ("M2-20", "GET", "/mobile2api/v1/mobilemembers/{planId}/{username}", "Harness", "Excluded", "Acceptance harness — out of business scope"),
    ("M2-21", "GET", "/mobile2api/v1/mobilebalancetrend/{ext}", "Balance", "Complete", "Master regression"),
    ("M2-22", "GET", "/mobile2api/v1/mobileperformance/{ext}", "Performance", "Complete", "Master regression"),
    ("M2-23", "GET", "/mobile2api/v1/mobilestackup/{planId}", "Stackup", "Complete", "Master regression"),
    ("M2-24", "GET", "/mobile2api/v1/mobileugift", "UGift", "Complete", "Master regression"),
    ("M2-25", "PATCH", "/mobile2api/v1/mobileugift/{ext}", "UGift", "Complete", "Master regression"),
]


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 9) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def style_table_header(row) -> None:
    for cell in row.cells:
        shade_cell(cell, "003057")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"


def add_heading(doc: Document, title: str, level: int = 1) -> None:
    h = doc.add_heading(title, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"


def add_callout(doc: Document, text: str, bg: str = GREEN_BG) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    set_cell_text(cell, text, size=11, bold=True, color=NAVY)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Mobile 2 API Automation Sign-Off  |  Internal")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"QA Automation — AMSQUAD  |  {SIGNOFF_DATE}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"


def chart_coverage(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8, 3.2))
    ax.barh(["Business endpoints"], [M2_IMPLEMENTED], color="#16A34A", height=0.45, label=f"Implemented ({M2_IMPLEMENTED})")
    ax.barh(["Business endpoints"], [1], left=[M2_IMPLEMENTED], color="#CBD5E1", height=0.45, label="Excluded (1)")
    ax.set_xlim(0, M2_BUSINESS_TOTAL)
    ax.set_xlabel("Endpoint count (documented business scope = 25)", fontsize=10, color="#334155")
    ax.set_title(f"Mobile 2 API Automation — {M2_IMPLEMENTED}/{M2_BUSINESS_TOTAL} Complete ({M2_PCT}%)",
                 fontsize=13, fontweight="bold", color="#003057", pad=10)
    ax.legend(loc="lower right", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.text(0.5, 0.01, f"Excluded: {EXCLUDED_PATH} — acceptance harness only",
             ha="center", fontsize=8.5, color="#64748B")
    fig.patch.set_facecolor("white")
    fig.tight_layout(rect=[0, 0.04, 1, 1])
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build() -> None:
    chart_path = ASSETS / "chart_m2_signoff_coverage.png"
    chart_coverage(chart_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    setup_header_footer(doc)

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MOBILE 2 API AUTOMATION")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Formal Sign-Off Certificate")
    r2.font.size = Pt(22)
    r2.font.color.rgb = TEAL
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = st.add_run("STATUS: COMPLETE")
    r3.bold = True
    r3.font.size = Pt(18)
    r3.font.color.rgb = GREEN
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Unite MSC — Mobile 2 API Test Automation",
        SIGNOFF_DATE,
        "Prepared by: QA Automation (AMSQUAD)",
        "Classification: Internal",
    ]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
    doc.add_page_break()

    # Document control
    add_heading(doc, "Document Control", level=1)
    ctrl = doc.add_table(rows=7, cols=2)
    ctrl.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Document title", "Mobile 2 API Automation — Sign-Off Certificate"),
        ("Version", "1.0 — Final"),
        ("Sign-off date", SIGNOFF_DATE),
        ("Program", "Unite MSC Mobile 2 API Automation"),
        ("Evidence baseline", f"Commit {BASELINE_COMMIT} on main branch"),
        ("Prior baseline", f"{PRIOR_SIGNOFF_DATE} — 22/25 endpoints (88%) @ {PRIOR_SIGNOFF_COMMIT}"),
        ("Status", "COMPLETE — 24/25 business endpoints (96%)"),
    ]):
        shade_cell(ctrl.rows[i].cells[0], LIGHT_BG)
        set_cell_text(ctrl.rows[i].cells[0], k, bold=True, color=NAVY)
        set_cell_text(ctrl.rows[i].cells[1], v, color=GRAY)
    doc.add_page_break()

    # 1 Executive summary
    add_heading(doc, "1. Executive Summary", level=1)
    add_callout(doc,
        f"Mobile 2 API automation is COMPLETE for the defined business scope: "
        f"{M2_IMPLEMENTED} of {M2_BUSINESS_TOTAL} documented business endpoints ({M2_PCT}%) "
        f"have canonical automated tests on the main branch.")
    doc.add_paragraph(
        "This sign-off certifies delivery of the Mobile 2 canonical TestNG automation framework, "
        "endpoint test coverage, master and module regression suites, OKD/NMD branding support, "
        "and API validation through Layer 4 (business data assertions). "
        "One harness-only endpoint is intentionally excluded from the business automation numerator.")
    doc.add_picture(str(chart_path), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    kpi = doc.add_table(rows=5, cols=2)
    kpi.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Documented business endpoints", "25"),
        ("Automated (business scope)", f"{M2_IMPLEMENTED}"),
        ("Coverage percentage", f"{M2_PCT}%"),
        ("Excluded (by design)", "1 — acceptance harness endpoint"),
        ("Sign-off determination", "COMPLETE"),
    ]):
        shade_cell(kpi.rows[i].cells[0], LIGHT_BG)
        set_cell_text(kpi.rows[i].cells[0], k, bold=True, color=NAVY)
        bg = GREEN_BG if i == 4 else "FFFFFF"
        shade_cell(kpi.rows[i].cells[1], bg)
        set_cell_text(kpi.rows[i].cells[1], v, bold=(i == 4), color=GREEN if i == 4 else GRAY)
    doc.add_page_break()

    # 2 Scope
    add_heading(doc, "2. Scope Definition", level=1)
    for item in [
        "In scope: All documented Mobile 2 business API endpoints under /mobile2api/v1 (denominator = 25).",
        "Automation type: Canonical TestNG API tests with HTTP, schema, and business assertions (L1–L4).",
        "Branding: OKD (okdirect) and NMD (nmdirect / newyork) session paths supported.",
        "Environments: QC4 and Stage 1 via Maven acceptance profiles.",
        "Out of scope for this sign-off: SQL API–DB reconciliation (L5), Enrollment API, Mobile 1 business endpoints.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Excluded endpoint (intentional)", level=2)
    add_callout(doc,
        f"{EXCLUDED_PATH} — Acceptance/harness helper used for member lookup in test flows. "
        "Smoke test exists but endpoint is excluded from the business automation numerator by approved scope definition.",
        AMBER_BG)
    doc.add_page_break()

    # 3 Sign-off criteria
    add_heading(doc, "3. Sign-Off Criteria — All Met", level=1)
    criteria = doc.add_table(rows=9, cols=3)
    criteria.style = "Table Grid"
    for i, h in enumerate(["#", "Criterion", "Status"]):
        set_cell_text(criteria.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(criteria.rows[0])
    rows = [
        ("1", "Canonical TestNG framework on main branch", "MET"),
        ("2", "24/25 business endpoints have automated tests", "MET"),
        ("3", "Master regression suite wires stable endpoints (OKD + NMD)", "MET"),
        ("4", "L1–L4 API validation (HTTP, contract, schema, business assertions)", "MET"),
        ("5", "Module regression suites per functional area", "MET"),
        ("6", "HTML test reporting and Maven profile execution", "MET"),
        ("7", "Dashboard + YTD summary endpoints automated", "MET"),
        ("8", "Documented exclusions and destructive-test placement", "MET"),
    ]
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            bg = GREEN_BG if ci == 2 else (LIGHT_BG if ci == 0 else "FFFFFF")
            shade_cell(criteria.rows[ri].cells[ci], bg)
            set_cell_text(criteria.rows[ri].cells[ci], val, bold=(ci == 2), color=GREEN if ci == 2 else (NAVY if ci == 0 else GRAY))
    doc.add_page_break()

    # 4 Endpoint register
    add_heading(doc, "4. Endpoint Register", level=1)
    reg = doc.add_table(rows=1 + len(ENDPOINTS), cols=6)
    reg.style = "Table Grid"
    for i, h in enumerate(["ID", "Method", "Path", "Area", "Status", "Regression placement"]):
        set_cell_text(reg.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(reg.rows[0])
    status_bg = {"Complete": GREEN_BG, "Excluded": AMBER_BG}
    for ri, row in enumerate(ENDPOINTS, start=1):
        for ci, val in enumerate(row):
            bg = status_bg.get(row[4], LIGHT_BG if ci == 0 else "FFFFFF")
            if ci == 4:
                bg = status_bg.get(val, "FFFFFF")
            shade_cell(reg.rows[ri].cells[ci], bg)
            set_cell_text(reg.rows[ri].cells[ci], val, bold=(ci == 4), color=NAVY if ci in (0, 4) else GRAY, size=8)
    doc.add_page_break()

    # 5 Validation
    add_heading(doc, "5. Validation Standard Delivered", level=1)
    val = doc.add_table(rows=6, cols=3)
    val.style = "Table Grid"
    for i, h in enumerate(["Layer", "Description", "Delivered"]):
        set_cell_text(val.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(val.rows[0])
    for ri, row in enumerate([
        ("L1 — HTTP / Auth", "Status codes, session tokens, transport and auth errors", "Yes"),
        ("L2 — Response contract", "Content-type, JSON parsing, error payloads", "Yes"),
        ("L3 — Schema / structure", "JSON schema validation, required fields, types", "Yes"),
        ("L4 — Business assertions", "Field values, business rules, branding-specific checks", "Yes"),
        ("L5 — SQL API–DB", "JDBC reconciliation against API JSON", "Out of scope"),
    ], start=1):
        for ci, v in enumerate(row):
            bg = GREEN_BG if v == "Yes" else (AMBER_BG if v == "Out of scope" else "FFFFFF")
            shade_cell(val.rows[ri].cells[ci], bg if ci == 2 else (LIGHT_BG if ci == 0 else "FFFFFF"))
            set_cell_text(val.rows[ri].cells[ci], v, bold=(ci == 0), color=NAVY if ci == 0 else GRAY)
    doc.add_page_break()

    # 6 Framework deliverables
    add_heading(doc, "6. Framework & Program Deliverables", level=1)
    for item in [
        "Migration from legacy Cucumber to canonical TestNG framework",
        "Shared authentication module with dynamic SQL credential loading (OKD + NMD)",
        "19 Mobile 2 RequestTest classes covering all business functional areas",
        "25 TestNG suite XML definitions and 30+ Maven execution profiles",
        "Master regression suite (OKD + NMD branding runs)",
        "Module regression suites per area (activity, banks, contribution, dashboard, etc.)",
        "Smoke suite for destructive operations (PUT/DELETE banks) and harness endpoints",
        "HTML reporting integration for test execution evidence",
        "Dashboard automation including YTD summary endpoint",
        "Custom reporting design for leadership and program metrics",
        "GitHub Actions pipeline — Dashboard vertical slice validated",
        "Performance regression foundation on Stage 1 (separate track)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    # 7 Evidence
    add_heading(doc, "7. Environment & Execution Evidence", level=1)
    add_callout(doc,
        "Hybrid environment approach: Stage 1 is the primary execution environment for sign-off evidence "
        "while QC4 has team dependencies. Tests pass when authentication path is healthy.",
        LIGHT_BG)
    env = doc.add_table(rows=5, cols=3)
    env.style = "Table Grid"
    for i, h in enumerate(["Environment", "Status", "Notes"]):
        set_cell_text(env.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(env.rows[0])
    for ri, row in enumerate([
        ("Stage 1", "Primary evidence", "Execution verified July 22, 2026 — NMD paths confirmed"),
        ("QC4", "Secondary", "Re-run when environment stable; some OKD failures are env-specific"),
        ("OKD branding", "Supported", "okdirect session path in master suite"),
        ("NMD branding", "Supported", "nmdirect / newyork session path in master suite"),
    ], start=1):
        for ci, val in enumerate(row):
            shade_cell(env.rows[ri].cells[ci], LIGHT_BG if ci == 0 else "FFFFFF")
            set_cell_text(env.rows[ri].cells[ci], val, color=NAVY if ci == 0 else GRAY)

    add_heading(doc, "Prior baseline progression", level=2)
    doc.add_paragraph(
        f"July 14, 2026 baseline: 22 of 25 endpoints (88%) verified at commit {PRIOR_SIGNOFF_COMMIT}. "
        f"Subsequent delivery added YTD summary and Banks GET-by-id endpoints, bringing coverage to "
        f"{M2_IMPLEMENTED}/{M2_BUSINESS_TOTAL} ({M2_PCT}%) at commit {BASELINE_COMMIT}.")
    doc.add_page_break()

    # 8 Known limitations
    add_heading(doc, "8. Known Limitations (Documented — Not Sign-Off Blockers)", level=1)
    lim = doc.add_table(rows=5, cols=2)
    lim.style = "Table Grid"
    for i, h in enumerate(["Item", "Description"]):
        set_cell_text(lim.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(lim.rows[0])
    for ri, row in enumerate([
        ("Destructive tests", "PUT/DELETE banks and DELETE contribution run in smoke/module suites only — excluded from master by design"),
        ("Contribution detail/PUT", "Stage 1 env-specific test-data fixture may return 401; QC4 fixture available"),
        ("Harness endpoint", "mobilemembers excluded from business numerator — smoke coverage only"),
        ("GitLab nightly job", "Recurring scheduled regression — follow-up DevOps item (QA-1405)"),
    ], start=1):
        shade_cell(lim.rows[ri].cells[0], LIGHT_BG)
        set_cell_text(lim.rows[ri].cells[0], row[0], bold=True, color=NAVY)
        set_cell_text(lim.rows[ri].cells[1], row[1], color=GRAY)
    doc.add_page_break()

    # 9 Out of scope
    add_heading(doc, "9. Explicitly Out of Scope", level=1)
    for item in [
        "L5 SQL API–DB reconciliation (separate enhancement program if leadership approves)",
        "Mobile 1 API business endpoint migration (active program — separate sign-off)",
        "Enrollment API automation (future multi-sprint program)",
        "Full GitLab nightly Mobile 2 regression scheduling (follow-up item)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 10 Sign-off
    add_heading(doc, "10. Sign-Off Approval", level=1)
    add_callout(doc,
        "Mobile 2 API automation is hereby marked COMPLETE for the defined business scope "
        f"({M2_IMPLEMENTED}/{M2_BUSINESS_TOTAL} endpoints, {M2_PCT}%) effective {SIGNOFF_DATE}.")
    doc.add_paragraph()
    sig = doc.add_table(rows=5, cols=4)
    sig.style = "Table Grid"
    for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
        set_cell_text(sig.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(sig.rows[0])
    for ri, role in enumerate(["QA Automation Lead", "Program Lead / SME", "Engineering Lead", "Leadership Approval"], start=1):
        shade_cell(sig.rows[ri].cells[0], LIGHT_BG)
        set_cell_text(sig.rows[ri].cells[0], role, bold=True, color=NAVY)
        for ci in range(1, 4):
            set_cell_text(sig.rows[ri].cells[ci], "", color=GRAY)
            sig.rows[ri].cells[ci].height = Cm(1.2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Certification statement: ").bold = True
    p.add_run(
        "The undersigned confirm that Mobile 2 API automation meets the sign-off criteria defined in "
        "Section 3, with documented exclusions in Sections 2 and 8, and validation through Layer 4 "
        "as the approved completion standard."
    )

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(f"Created: {DOCX_OUT}")


if __name__ == "__main__":
    build()
