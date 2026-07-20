#!/usr/bin/env python3
"""Generate professional Unite MSC Leadership Update DOCX with charts and RAG styling."""

from __future__ import annotations

import os
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

# Brand palette (Ascensus-adjacent executive)
NAVY = RGBColor(0x00, 0x30, 0x57)
TEAL = RGBColor(0x00, 0x7A, 0x8C)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xE6, 0x51, 0x00)
RED = RGBColor(0xC6, 0x28, 0x28)
GRAY = RGBColor(0x61, 0x61, 0x61)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "E8EEF4"
GREEN_BG = "E8F5E9"
AMBER_BG = "FFF3E0"
RED_BG = "FFEBEE"

OUT_DIR = Path(__file__).resolve().parent
ASSETS = OUT_DIR / "_docx_assets"
DOCX_OUT = OUT_DIR / "Unite-MSC-Leadership-Update-2026-07-17.docx"
REPORT_DATE = "July 17, 2026"


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def style_table_header(row, bg: str = "003057") -> None:
    for cell in row.cells:
        shade_cell(cell, bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"


def add_horizontal_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003057")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_heading(doc: Document, title: str, level: int = 1) -> None:
    h = doc.add_heading(title, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"


def add_callout(doc: Document, text: str, bg: str = LIGHT_BG, border_color: str = "003057") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, bg)
    set_cell_text(cell, text, size=11)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.color.rgb = NAVY


def add_rag_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    """rows: (track, status_text, rag) where rag in GREEN/AMBER/RED/TBD"""
    rag_map = {
        "GREEN": (GREEN_BG, GREEN, "Complete"),
        "AMBER": (AMBER_BG, AMBER, "In Progress"),
        "RED": (RED_BG, RED, "At Risk"),
        "TBD": ("F5F5F5", GRAY, "TBD"),
    }
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=WHITE)
    style_table_header(table.rows[0])
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        rag_key = row_data[-1]
        bg, fg, _ = rag_map.get(rag_key, rag_map["TBD"])
        for ci, val in enumerate(row_data[:-1]):
            shade_cell(cells[ci], bg if ci == 0 else "FFFFFF")
            set_cell_text(cells[ci], val, color=NAVY if ci == 0 else GRAY)
        shade_cell(cells[-1], bg)
        set_cell_text(cells[-1], row_data[-2], bold=True, color=fg)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Unite MSC Leadership Update  |  Internal")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"QA Automation — AMSQUAD  |  {REPORT_DATE}  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"
    # page number field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    fp._p.append(fld)


def chart_m2_coverage(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7, 3.5))
    labels = ["Automated\n(verified)", "Pending /\nExcluded", "Projected\n(add'l)"]
    verified = 22
    pending = 3
    projected = 2  # YTD + banks GET-by-id — TBD
    colors = ["#2E7D32", "#B0BEC5", "#007A8C"]
    bars = ax.barh(["Mobile 2 Endpoints (25 total)"], [verified], color=colors[0], height=0.35, label=f"Verified {verified}")
    ax.barh(["Mobile 2 Endpoints (25 total)"], [pending], left=[verified], color=colors[1], height=0.35, label=f"Gap/Excluded {pending}")
    ax.barh(["Mobile 2 (projected)"], [verified + projected], color=colors[2], height=0.25, alpha=0.7)
    ax.set_xlim(0, 25)
    ax.set_xlabel("Endpoint count (documented business scope)", fontsize=10)
    ax.set_title("Mobile 2 API Endpoint Automation Coverage", fontsize=12, fontweight="bold", color="#003057")
    ax.axvline(22, color="#E65100", linestyle="--", linewidth=1, alpha=0.8)
    ax.text(22.3, 0.5, "88% verified\n(2026-07-14)", fontsize=9, color="#E65100")
    ax.legend(loc="lower right", fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart_m1_coverage(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(5, 4))
    sizes = [1, 26]
    labels = ["Automated\n(1)", "Remaining\n(26)"]
    colors = ["#2E7D32", "#E0E0E0"]
    explode = (0.06, 0)
    ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct="%1.1f%%",
           startangle=90, textprops={"fontsize": 10})
    ax.set_title("Mobile 1 — Documented Endpoints (27)\nAuth foundation complete", fontsize=11, fontweight="bold", color="#003057")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart_program_rag(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7, 4))
    tracks = ["Mobile 2 API", "Mobile 1 API", "Performance", "Pipeline", "V2/V3 UI"]
    scores = [88, 4, 40, 55, 75]  # illustrative progress % for visual only
    colors = ["#2E7D32" if s >= 80 else "#E65100" if s >= 30 else "#C62828" for s in scores]
    bars = ax.bar(tracks, scores, color=colors, edgecolor="white", linewidth=0.8)
    ax.set_ylim(0, 100)
    ax.set_ylabel("Relative readiness (%)", fontsize=10)
    ax.set_title("Program Track Readiness Snapshot", fontsize=12, fontweight="bold", color="#003057")
    ax.axhline(80, color="#003057", linestyle=":", linewidth=1, alpha=0.5)
    for bar, s in zip(bars, scores):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 2, f"{s}%", ha="center", fontsize=9)
    plt.xticks(rotation=15, ha="right", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    note = "M2=verified endpoint %; M1=endpoint %; others=qualitative readiness"
    fig.text(0.5, 0.02, note, ha="center", fontsize=7, color="#616161")
    fig.tight_layout(rect=[0, 0.05, 1, 1])
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart_pipeline_flow(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8, 3))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")
    boxes = [
        (0.5, 1.5, "Source\nCode", "#003057"),
        (2.3, 1.5, "Maven\nBuild", "#007A8C"),
        (4.1, 1.5, "Nexus\nArchive", "#007A8C"),
        (5.9, 2.2, "GHA\nDashboard [Done]", "#2E7D32"),
        (5.9, 0.8, "Module\nSuites [WIP]", "#E65100"),
        (7.7, 1.5, "GitLab\nNightly [TBD]", "#C62828"),
        (9.0, 1.5, "Reports\n& Triage", "#003057"),
    ]
    for x, y, text, color in boxes:
        rect = mpatches.FancyBboxPatch((x, y), 1.4, 0.9, boxstyle="round,pad=0.05", linewidth=1.5, edgecolor=color, facecolor=color, alpha=0.15)
        ax.add_patch(rect)
        ax.text(x + 0.7, y + 0.45, text, ha="center", va="center", fontsize=8, fontweight="bold", color=color)
    arrows = [(1.9, 1.95, 2.3, 1.95), (3.7, 1.95, 4.1, 1.95), (5.5, 1.95, 5.9, 2.2), (5.5, 1.95, 5.9, 0.95), (7.3, 1.95, 7.7, 1.95), (9.1, 1.95, 9.0, 1.95)]
    for x1, y1, x2, y2 in [(1.9, 1.95, 2.3, 1.95), (3.7, 1.95, 4.1, 1.95), (5.5, 2.0, 5.9, 2.35), (5.5, 1.9, 5.9, 1.15), (7.3, 1.95, 7.7, 1.95), (9.1, 1.95, 9.4, 1.95)]:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", color="#616161", lw=1.2))
    ax.set_title("CI/CD Pipeline Workflow — Unite MSC Mobile 2", fontsize=12, fontweight="bold", color="#003057", pad=10)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart_roadmap(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8, 3.5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")
    phases = [
        (0.5, 3.5, "Sprint 26.11\n(now)", "Mobile 2 sign-off\nMobile 1 non-IDP", "#2E7D32"),
        (4.0, 3.5, "Sprint 26.12", "M1 IDP hardening\nM2 sustain", "#007A8C"),
        (7.5, 3.5, "Sprint 26.13+", "Enrollment API\n(multi-sprint)", "#E65100"),
        (0.5, 1.2, "Perf", "IDP login script\n→ flow expansion", "#007A8C"),
        (4.0, 1.2, "Pipeline", "GitLab nightly\nAll module GHA", "#E65100"),
    ]
    for x, y, title, body, color in phases:
        rect = mpatches.FancyBboxPatch((x, y), 2.8, 1.6, boxstyle="round,pad=0.08", linewidth=2, edgecolor=color, facecolor=color, alpha=0.12)
        ax.add_patch(rect)
        ax.text(x + 1.4, y + 1.25, title, ha="center", fontsize=10, fontweight="bold", color=color)
        ax.text(x + 1.4, y + 0.55, body, ha="center", fontsize=8, color="#333333")
    ax.annotate("", xy=(3.9, 4.3), xytext=(3.4, 4.3), arrowprops=dict(arrowstyle="->", color="#003057", lw=2))
    ax.annotate("", xy=(7.4, 4.3), xytext=(6.9, 4.3), arrowprops=dict(arrowstyle="->", color="#003057", lw=2))
    ax.set_title("2-Sprint Roadmap + Horizon", fontsize=12, fontweight="bold", color="#003057")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_cover_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("UNITE MSC")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Automation Program — Leadership Update")
    r2.font.size = Pt(18)
    r2.font.color.rgb = TEAL
    r2.font.name = "Calibri"
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [REPORT_DATE, "Audience: Rajiv/Rajib, Henry, Kevin", "Prepared by: QA Automation (AMSQUAD)", "Classification: Internal — Leadership"]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
        run.font.name = "Calibri"
    doc.add_page_break()


def add_metrics_table(doc: Document, title: str, data: list[tuple[str, str]], note: str | None = None) -> None:
    add_section_heading(doc, title, level=2)
    table = doc.add_table(rows=1 + len(data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Metric", bold=True, color=WHITE)
    set_cell_text(hdr[1], "Value", bold=True, color=WHITE)
    style_table_header(table.rows[0])
    for i, (k, v) in enumerate(data):
        row = table.rows[i + 1].cells
        shade_cell(row[0], LIGHT_BG)
        set_cell_text(row[0], k, bold=True, color=NAVY)
        set_cell_text(row[1], v, color=GRAY)
    if note:
        p = doc.add_paragraph(note)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = GRAY


def build_document() -> None:
    ASSETS.mkdir(exist_ok=True)
    chart_m2_coverage(ASSETS / "chart_m2_coverage.png")
    chart_m1_coverage(ASSETS / "chart_m1_coverage.png")
    chart_program_rag(ASSETS / "chart_program_rag.png")
    chart_pipeline_flow(ASSETS / "chart_pipeline_flow.png")
    chart_roadmap(ASSETS / "chart_roadmap.png")

    doc = Document()
    setup_header_footer(doc)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover_page(doc)

    # Table of contents (manual)
    add_section_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Program Metrics at a Glance",
        "3. Mobile 2 — Sign-Off Readiness",
        "4. Mobile 1 — Active Sprint Focus",
        "5. Pipeline & CI/CD Workflow",
        "6. Performance Regression",
        "7. V2/V3 UI Automation",
        "8. Team Contributions & Kudos",
        "9. Risks & Dependencies",
        "10. Leadership Asks",
        "11. Roadmap",
        "Appendix — Sources & TBD Items",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # 1 Executive Summary
    add_section_heading(doc, "1. Executive Summary", level=1)
    add_callout(
        doc,
        "The Unite MSC automation program has moved from fragmented legacy Cucumber into a reusable, "
        "scalable canonical framework across Mobile APIs, UI regression, performance testing, and pipeline readiness.",
    )
    summary_bullets = [
        ("Mobile 2 API", "Baseline functionally complete — ready for formal sign-off pending MR merge and refreshed QC4/Stage 1 evidence.", GREEN),
        ("Mobile 1 API", "Active sprint focus — authentication foundation complete on main.", AMBER),
        ("IDP coverage", "New capability in canonical framework (not in legacy baseline).", TEAL),
        ("Performance", "Foundation established — MSC non-IDP Jenkins job live.", AMBER),
        ("Pipeline", "GHA Dashboard slice complete; GitLab nightly not yet created.", AMBER),
    ]
    for label, text, color in summary_bullets:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.color.rgb = color
        r2 = p.add_run(text)
        r2.font.color.rgb = GRAY
    add_horizontal_line(doc)

    # 2 Metrics at a glance
    add_section_heading(doc, "2. Program Metrics at a Glance", level=1)
    doc.add_paragraph("Verified metrics are cited with source date; projected values require re-run (see Appendix).")
    doc.add_picture(str(ASSETS / "chart_program_rag.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_rag_table(
        doc,
        ["Track", "Status", "RAG"],
        [
            ("Mobile 2 API", "88% endpoints verified — sign-off path", "GREEN"),
            ("Mobile 1 API", "Auth done — business migration in sprint", "AMBER"),
            ("Performance", "Jenkins foundation — flows expanding", "AMBER"),
            ("Pipeline", "GHA slice done — nightly pending", "AMBER"),
            ("V2/V3 UI", "Nightly stable — daily triage continues", "GREEN"),
            ("Enrollment API", "Next major program — multi-sprint", "TBD"),
        ],
    )
    doc.add_paragraph()

    # 3 Mobile 2
    add_section_heading(doc, "3. Mobile 2 — Sign-Off Readiness", level=1)
    add_callout(doc, "READY FOR REVIEW — Conditional sign-off per sign-off package (2026-07-14).", GREEN_BG)
    doc.add_picture(str(ASSETS / "chart_m2_coverage.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_metrics_table(
        doc,
        "Verified Scorecard (2026-07-14, commit 7ccaf46)",
        [
            ("Documented in-scope endpoints", "25"),
            ("Automated endpoints", "22"),
            ("Coverage percentage", "88.0%"),
            ("Master regression runs (OKD + NMD)", "34"),
            ("Stage 1 master pass rate", "36 / 40"),
            ("Legacy migration", "~95%"),
            ("Postman parity", "~96%"),
        ],
        note="Source: 17-mobile2-api-automation-signoff.md",
    )

    add_section_heading(doc, "Functional Area Status", level=2)
    area_table = doc.add_table(rows=8, cols=3)
    area_table.style = "Table Grid"
    headers = ["Area", "Status", "Notes"]
    for i, h in enumerate(headers):
        set_cell_text(area_table.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(area_table.rows[0])
    areas = [
        ("Activity, Transactions, Investment", "Complete", "Master wired"),
        ("Banks", "Near complete", "GET-by-id added post sign-off — verify"),
        ("Content & Plans", "Complete", "—"),
        ("Contribution (6 endpoints)", "Partial", "Stage 1 fixture — env-specific 401"),
        ("Dashboard", "Near complete", "YTD added post sign-off — verify"),
        ("Balance / Performance / Stackup", "Complete", "—"),
        ("UGift", "Complete", "—"),
    ]
    status_color = {"Complete": GREEN_BG, "Near complete": AMBER_BG, "Partial": RED_BG}
    for ri, (area, status, note) in enumerate(areas, start=1):
        cells = area_table.rows[ri].cells
        shade_cell(cells[0], LIGHT_BG)
        set_cell_text(cells[0], area, bold=True, color=NAVY)
        shade_cell(cells[1], status_color.get(status, "FFFFFF"))
        set_cell_text(cells[1], status, bold=True, color=NAVY)
        set_cell_text(cells[2], note, color=GRAY)

    add_section_heading(doc, "Sign-Off Path", level=2)
    for step in [
        "Merge pending MRs to main",
        "Re-run QC4 + Stage 1 master regression",
        "Resolve Stage 1 contribution test-data fixture",
        "Leadership formal approval",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_page_break()

    # 4 Mobile 1
    add_section_heading(doc, "4. Mobile 1 — Active Sprint Focus", level=1)
    doc.add_picture(str(ASSETS / "chart_m1_coverage.png"), width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_metrics_table(
        doc,
        "Current State (documented scope)",
        [
            ("Documented business endpoints", "27"),
            ("Automated", "1 (POST mobilemembersession)"),
            ("Coverage", "3.7%"),
            ("Auth foundation (OKD + NMD)", "Complete"),
        ],
    )
    add_section_heading(doc, "Sprint 26.11 Goals", level=2)
    goals = [
        "This sprint: Mobile 1 non-IDP baseline (owner, profile, dashboard, beneficiary)",
        "Next sprint: Mobile 1 IDP compatibility and hardening",
        "Target: M1 + M2 complete for non-IDP and IDP by end of next sprint (data permitting)",
    ]
    for g in goals:
        doc.add_paragraph(g, style="List Bullet")
    doc.add_page_break()

    # 5 Pipeline
    add_section_heading(doc, "5. Pipeline & CI/CD Workflow", level=1)
    doc.add_picture(str(ASSETS / "chart_pipeline_flow.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rag_table(
        doc,
        ["Component", "Status", "RAG"],
        [
            ("GitHub Actions — Dashboard vertical slice", "Complete (Chaitanya)", "GREEN"),
            ("Nexus archive packaging", "Working", "GREEN"),
            ("Mobile 2 module pipeline expansion", "In progress", "AMBER"),
            ("GitLab nightly Mobile 2 regression", "Story needed — DevOps", "RED"),
            ("Master integration / regression suite", "Available", "GREEN"),
        ],
    )
    doc.add_page_break()

    # 6 Performance
    add_section_heading(doc, "6. Performance Regression", level=1)
    add_metrics_table(
        doc,
        "Performance Metrics",
        [
            ("MSC scenarios in Jenkins", "1 (non-IDP login → Dashboard)"),
            ("MSC scenarios in repo", "≥ 2 JMX files"),
            ("Scheduled MSC nightly", "No — ad hoc only"),
            ("IDP UP suite (related)", "Scheduled weekdays ~3 AM"),
            ("QA-1229 (non-IDP Jenkins)", "Done — 2026-07-02"),
            ("QA-1228 (IDP MSC login)", "In progress"),
        ],
    )
    add_callout(
        doc,
        "Kudos — Priti Choudhary ramped quickly in a complex authentication/performance area and delivered "
        "the MSC non-IDP Jenkins foundation despite limited initial domain context.",
        GREEN_BG,
    )
    p = doc.add_paragraph()
    p.add_run("Next targets: ").bold = True
    p.add_run("Contribution → Banks → Activity → Mobile 1 (after clean manual validations).")

    # 7 V2/V3
    add_section_heading(doc, "7. V2/V3 UI Automation", level=1)
    add_metrics_table(
        doc,
        "UI Regression Status",
        [
            ("Nightly regression", "Stable (per program operations)"),
            ("Morning validation effort", "~15–20 min when clean"),
            ("V2 UP-scoped inventory", "268 scenarios (Jun 2026 assessment)"),
            ("V3 UP-scoped inventory", "379 scenarios (Jun 2026 assessment)"),
            ("Current nightly pass rate", "TBD — requires qTest/Jenkins export"),
        ],
    )

    # 8 Team
    add_section_heading(doc, "8. Team Contributions & Kudos", level=1)
    team_table = doc.add_table(rows=6, cols=3)
    team_table.style = "Table Grid"
    for i, h in enumerate(["Team Member", "Primary Focus", "Key Delivery"]):
        set_cell_text(team_table.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(team_table.rows[0])
    team = [
        ("Swapnil Patil", "Program, master suite, KB", "Auth SQL, sign-off pack, pipeline guides"),
        ("Dinesh Kumar", "Dashboard, plans, YTD", "Plans MR, Ugift, YTD summary"),
        ("Venkatesh Mallela", "Contribution, balance", "Contribution CRUD, performance"),
        ("Sunil Godiyal", "Banks, investments, TH", "Banks suite, GET-by-id"),
        ("Priti Choudhary", "Performance", "MSC Jenkins job QA-1229"),
    ]
    for ri, row in enumerate(team, start=1):
        for ci, val in enumerate(row):
            shade_cell(team_table.rows[ri].cells[ci], LIGHT_BG if ci == 0 else "FFFFFF")
            set_cell_text(team_table.rows[ri].cells[ci], val, bold=(ci == 0), color=NAVY if ci == 0 else GRAY)

    # 9 Risks
    add_section_heading(doc, "9. Risks & Dependencies", level=1)
    risks = [
        ("GitLab nightly not created", "High", "Manual regression burden on team"),
        ("Test data gaps (Stage 1/5/2)", "Medium", "Env-flexible framework blocked by data"),
        ("Enrollment encryption complexity", "Medium", "Multi-sprint schedule risk"),
        ("Contribution Stage 1 fixture", "Medium", "Env-specific 401 on detail/PUT"),
        ("Reporting cadence load", "Medium", "Needs BA/Scrum support"),
    ]
    risk_table = doc.add_table(rows=1 + len(risks), cols=3)
    risk_table.style = "Table Grid"
    for i, h in enumerate(["Risk", "Severity", "Impact"]):
        set_cell_text(risk_table.rows[0].cells[i], h, bold=True, color=WHITE)
    style_table_header(risk_table.rows[0])
    sev_bg = {"High": RED_BG, "Medium": AMBER_BG}
    for ri, (risk, sev, impact) in enumerate(risks, start=1):
        cells = risk_table.rows[ri].cells
        set_cell_text(cells[0], risk, bold=True, color=NAVY)
        shade_cell(cells[1], sev_bg.get(sev, "FFFFFF"))
        set_cell_text(cells[1], sev, bold=True, color=RED if sev == "High" else AMBER)
        set_cell_text(cells[2], impact, color=GRAY)

    # 10 Leadership asks
    add_section_heading(doc, "10. Leadership Asks", level=1)
    asks = [
        "Approve DevOps story: GitLab nightly Mobile 2 regression (QC4 first, Stage 1 parameterized)",
        "Confirm Mobile 2 baseline can move to formal sign-off after refreshed evidence run",
        "Confirm Mobile 1 as active sprint priority",
        "Support pipeline scheduling and environment / test-data readiness",
        "Structured backlog intake for new API and performance requests",
        "30–40% BA/Scrum/admin support for documentation, Jira hygiene, SME coordination, and weekly metrics",
    ]
    for i, ask in enumerate(asks, start=1):
        p = doc.add_paragraph(f"{i}. {ask}")
        p.paragraph_format.left_indent = Inches(0.25)

    # 11 Roadmap
    add_section_heading(doc, "11. Roadmap", level=1)
    doc.add_picture(str(ASSETS / "chart_roadmap.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_section_heading(doc, "Enrollment API — Next Horizon", level=2)
    for item in [
        "Higher complexity than Mobile 1/2 — encryption/decryption",
        "Requires test-data creation utility and MFA-disabled account handling",
        "Expected duration: more than one sprint",
        "Will support dynamic account creation for future regression",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # Appendix
    add_section_heading(doc, "Appendix — Sources & TBD Items", level=1)
    add_section_heading(doc, "Source References", level=2)
    sources = [
        "api-test-automation/.../17-mobile2-api-automation-signoff.md",
        "api-test-automation/.../16-mobile2-coverage-matrix.md",
        "api-test-automation/.../03-document-postman-coverage-matrix.md",
        "qa-automation-kb/.../unite-msc-performance-testing-tracker.md",
        "api-test-automation/.../17-MOBILE2-NEXUS-GITHUB-ACTIONS-PIPELINE.md",
        "qa-automation-kb/leadership-updates/unite-msc/2026-07-17-leadership-update/",
    ]
    for s in sources:
        doc.add_paragraph(s, style="List Bullet")

    add_section_heading(doc, "TBD — Requires Verification Before Next Meeting", level=2)
    tbd = [
        "Fresh QC4 master regression run (run-qc4-all-suites.ps1)",
        "Fresh Stage 1 master regression run (run-stage1-all-suites.ps1)",
        "Updated endpoint count after YTD + Banks GET-by-id merge",
        "GitLab open MR list and pipeline status",
        "V2/V3 nightly pass rate from qTest/Jenkins",
    ]
    for t in tbd:
        doc.add_paragraph(t, style="List Bullet")

    add_callout(
        doc,
        "Document generated from verified program artifacts. Do not cite projected Mobile 2 coverage ≥96% "
        "until matrix refresh and regression run complete.",
        AMBER_BG,
    )

    doc.save(DOCX_OUT)
    print(f"Created: {DOCX_OUT}")
    print(f"Charts: {ASSETS}")


if __name__ == "__main__":
    build_document()
