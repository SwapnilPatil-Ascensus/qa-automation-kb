#!/usr/bin/env python3
"""Build final Universal Platform Test Automation Coverage Assessment (leadership)."""
from __future__ import annotations

import os
import subprocess
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "02-deliverables"
DOCX_OUT = DELIVERABLES / "Universal-Platform-Test-Automation-Coverage-Assessment-Final.docx"
PDF_OUT = DELIVERABLES / "Universal-Platform-Test-Automation-Coverage-Assessment-Final.pdf"

PUBLISHED = "July 1, 2026"
ASSESSMENT_AS_OF = "June 29, 2026"
VERSION = "1.0"
STATUS = "Final Current-State Assessment"

# --- Final SME-validated numbers ---
V2_TOTAL = 744
V2_SCOPED = 268
V2_OUT = 476
V2_PCT = 36.0
V3_TOTAL = 436
V3_SCOPED = 379
V3_OTHER = 57
V3_PCT = 86.9

NAVY = "1F4E78"
BLUE = "5B9BD5"
LIGHT = "D9EAF7"
GREY = "666666"
BODY_FONT = "Aptos"
TITLE_FONT = "Aptos Display"


def ensure_docx():
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])


def rgb(hexcolor: str):
    from docx.shared import RGBColor
    h = hexcolor.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def build():
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, Twips

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- Styles ---
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 16), (2, 13), (3, 11)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = BODY_FONT
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = rgb(NAVY)
        h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True

    def set_cell_shading(cell, fill: str):
        tc = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc.append(shd)

    def set_repeat_header(row):
        tr_pr = row._tr.get_or_add_trPr()
        el = OxmlElement("w:tblHeader")
        el.set(qn("w:val"), "true")
        tr_pr.append(el)

    def set_row_cant_split(row):
        tr_pr = row._tr.get_or_add_trPr()
        el = OxmlElement("w:cantSplit")
        el.set(qn("w:val"), "true")
        tr_pr.append(el)

    def add_field(paragraph, instr: str):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = instr
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_el)
        run._r.append(fld_sep)
        run._r.append(fld_end)

    def ensure_continuous_page_numbers(section):
        """Prevent section from restarting page numbering."""
        sect_pr = section._sectPr
        for tag in ("w:pgNumType",):
            el = sect_pr.find(qn(tag))
            if el is not None:
                sect_pr.remove(el)

    def ensure_no_title_page(section):
        """Remove title-page flag so section-start pages use standard header/footer."""
        sect_pr = section._sectPr
        el = sect_pr.find(qn("w:titlePg"))
        if el is not None:
            sect_pr.remove(el)

    def ensure_title_page(section):
        sect_pr = section._sectPr
        if sect_pr.find(qn("w:titlePg")) is None:
            sect_pr.append(OxmlElement("w:titlePg"))

    def setup_header_footer(section, landscape=False, cover_section=False):
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        ensure_continuous_page_numbers(section)

        if cover_section:
            section.different_first_page_header_footer = True
            ensure_title_page(section)
        else:
            section.different_first_page_header_footer = False
            ensure_no_title_page(section)

        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header = section.header
        footer = section.footer
        while len(header.paragraphs) > 1:
            p = header.paragraphs[-1]._element
            p.getparent().remove(p)
        while len(footer.paragraphs) > 1:
            p = footer.paragraphs[-1]._element
            p.getparent().remove(p)

        def write_header(paragraph):
            paragraph.clear()
            hr = paragraph.add_run(
                f"Universal Platform Test Automation Coverage Assessment | Final | Version {VERSION}"
            )
            hr.font.name = BODY_FONT
            hr.font.size = Pt(9)
            hr.font.color.rgb = rgb(GREY)

        def write_footer(paragraph):
            paragraph.clear()
            tab_pos = Inches(7.2 if not landscape else 10.0)
            paragraph.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
            fl = paragraph.add_run("QA Automation | Internal Use")
            fl.font.name = BODY_FONT
            fl.font.size = Pt(9)
            fl.font.color.rgb = rgb(GREY)
            paragraph.add_run("\t")
            paragraph.add_run("Page ").font.name = BODY_FONT
            paragraph.runs[2].font.size = Pt(9)
            paragraph.runs[2].font.color.rgb = rgb(GREY)
            add_field(paragraph, " PAGE ")
            paragraph.add_run(" of ").font.name = BODY_FONT
            paragraph.runs[4].font.size = Pt(9)
            paragraph.runs[4].font.color.rgb = rgb(GREY)
            add_field(paragraph, " NUMPAGES ")

        if cover_section:
            section.first_page_header.paragraphs[0].clear()
            section.first_page_footer.paragraphs[0].clear()
            write_header(header.paragraphs[0])
            write_footer(footer.paragraphs[0])
        else:
            write_header(header.paragraphs[0])
            write_footer(footer.paragraphs[0])

    setup_header_footer(doc.sections[0], cover_section=True)

    def para(text="", *, bold=False, italic=False, size=11, align=None, color=GREY, space_after=6):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.font.name = BODY_FONT
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
            if color:
                r.font.color.rgb = rgb(color) if isinstance(color, str) else color
        return p

    def bullet(text, bold_lead=None):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if bold_lead:
            rb = p.add_run(bold_lead)
            rb.bold = True
            rb.font.name = BODY_FONT
            rb.font.size = Pt(11)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        return p

    def add_table(headers, rows, widths=None, font_size=10):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.autofit = False
        hdr = t.rows[0]
        set_repeat_header(hdr)
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(font_size)
            r.font.color.rgb = rgb("FFFFFF")
            set_cell_shading(cell, NAVY)
            p.paragraph_format.keep_with_next = True
        for row_data in rows:
            row = t.add_row()
            set_row_cant_split(row)
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                r.font.name = BODY_FONT
                r.font.size = Pt(font_size)
                if i == 0:
                    set_cell_shading(cell, LIGHT)
        if widths:
            for i, w in enumerate(widths):
                for cell in t.columns[i].cells:
                    cell.width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    # ========== COVER PAGE (no header/footer on cover) ==========

    for _ in range(4):
        doc.add_paragraph()
    tp = para("", align=WD_ALIGN_PARAGRAPH.CENTER)
    r = tp.add_run("Universal Platform Test Automation")
    r.bold = True
    r.font.name = TITLE_FONT if TITLE_FONT else BODY_FONT
    r.font.size = Pt(26)
    r.font.color.rgb = rgb(NAVY)
    sp = para("", align=WD_ALIGN_PARAGRAPH.CENTER)
    rs = sp.add_run("Coverage Assessment")
    rs.font.name = TITLE_FONT if TITLE_FONT else BODY_FONT
    rs.font.size = Pt(20)
    rs.font.color.rgb = rgb(NAVY)
    para("", space_after=12)
    for line in [
        "Final Current-State Assessment",
        "",
        f"Prepared by: QA Automation",
        f"Published: {PUBLISHED}",
        f"Assessment data as of: {ASSESSMENT_AS_OF}",
        f"Version: {VERSION}",
        "",
        "Internal Use",
    ]:
        if line == "":
            doc.add_paragraph()
            continue
        pp = para(line, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=NAVY if "Final" in line else GREY)

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Executive Summary",
        "Scope and Counting Methodology",
        "Universal Platform Share of Existing UI Regression Inventory",
        "Leadership Dashboard",
        "Workstream Detail",
        "Current Automation Strengths",
        "Current Gaps and Improvement Opportunities",
        "Interpretation and Measurement Boundaries",
        "Conclusion",
    ]
    # Automatic TOC field
    p = doc.add_paragraph()
    add_field(p, r' TOC \o "1-3" \h \z \u ')
    p.paragraph_format.space_after = Pt(12)
    doc.add_page_break()

    # ========== 1. EXECUTIVE SUMMARY ==========
    doc.add_heading("Executive Summary", level=1)
    para(
        "This assessment inventories automated test assets for five Aha workstreams that define "
        "Universal Platform automation scope. It reports what is established, scheduled, and "
        "available for continued expansion — not requirement-level coverage percentages."
    )
    para("Workstreams assessed:", bold=True, color=NAVY, space_after=4)
    for ws in [
        "Universal Enrollment — ACS-I-2679",
        "IDP / Authentication — ACS-I-2680",
        "ABLE — ACS-I-2681",
        "Angular — ACS-I-2682",
        "Withdrawal — ACS-I-2690",
    ]:
        bullet(ws)

    para(
        "The automation team has established and maintains a substantial foundation across UI, API, "
        "and performance automation. Scheduled regression exists on both V2 and V3 frameworks, with "
        "additional assets implemented and ready for pipeline alignment."
    )

    doc.add_heading("V2 daily regression", level=2)
    bullet(f"{V2_TOTAL} total scheduled regression test cases in the qTest daily PRIME cycle.", bold_lead="Source inventory: ")
    bullet(
        f"{V2_SCOPED} cases mapped within the five-workstream Universal Platform scope, "
        "covering Enrollment, IDP/Web Login, LA ABLE, and Withdrawal.",
        bold_lead="UP-scoped inventory: ",
    )
    bullet(f"The remaining {V2_OUT} cases support legacy, non-UP, or other out-of-scope areas and are intentionally excluded from the scoped subtotal.", bold_lead="Out of scope: ")

    doc.add_heading("V3 regression", level=2)
    bullet(f"{V3_TOTAL} total TestNG methods in the source regression reports (Universal Enrollment 303 + Unite suites 133).", bold_lead="Source inventory: ")
    bullet(f"{V3_SCOPED} methods directly mapped to Universal Enrollment, IDP Login, and Member Withdrawal for this assessment.", bold_lead="UP-scoped inventory: ")
    bullet(f"Additional V3 automation ({V3_OTHER} methods in contributions, CSR account maintenance, and IDP web registration) exists outside the directly scoped subtotal.", bold_lead="Adjacent inventory: ")

    doc.add_heading("Other inventory (reported separately)", level=2)
    bullet("6 ABLE Entity Platform scenarios (V3, implemented; not in scoped V3 subtotal).")
    bullet("22 Angular lib-ui component test definitions (separate from end-to-end automation).")
    bullet("11 unique in-scope API operations across Enrollment, IDP, and Withdrawal.")
    bullet("15 in-scope performance business journeys across Enrollment, IDP, Withdrawal, and ABLE Entity.")

    para(
        "Requirement-level coverage percentages are not claimed. The Aha work items define "
        "high-level outcomes but do not yet provide an enumerated requirements baseline.",
        italic=True,
        size=10,
    )

    doc.add_heading("Current Automation Position", level=2)
    bullet("Substantial scheduled UI regression on V2 and V3 across enrollment, authentication, withdrawal, and ABLE foundations.")
    bullet("API automation implemented for enrollment, authentication, and withdrawal microservices.")
    bullet("Performance journey assets available for enrollment, IDP, withdrawal, and ABLE entity flows.")
    bullet("V2 Unite ABLE/CSR automation maintained across nine ABLE plans plus a dedicated LA ABLE suite.")
    bullet("V3 ABLE Entity Platform and Angular lib-ui component automation implemented and available for pipeline enablement.")
    bullet("Inventory-share percentages (36.0% V2, 86.9% V3) measure allocation of existing regression inventory to the five workstreams — not requirement coverage.")

    doc.add_page_break()

    # ========== 2. SCOPE AND COUNTING METHODOLOGY ==========
    doc.add_heading("Scope and Counting Methodology", level=1)
    para(
        "This assessment distinguishes source regression populations from UP-scoped inventory mapped "
        "to the five Aha workstreams. Source populations describe what the daily regression suites "
        "execute; scoped inventory describes what maps to Universal Platform scope for leadership reporting."
    )
    add_table(
        ["Population", "V2", "V3"],
        [
            ["Total source inventory", str(V2_TOTAL), str(V3_TOTAL)],
            ["UP-scoped inventory", str(V2_SCOPED), str(V3_SCOPED)],
            ["Non-UP / legacy / out-of-scope or adjacent", str(V2_OUT), str(V3_OTHER)],
        ],
        widths=[2.5, 1.5, 1.5],
    )
    doc.add_heading("Counting units", level=2)
    add_table(
        ["Layer", "Counting unit"],
        [
            ["V2 UI", "Executed qTest test cases"],
            ["V3 UI", "Executed TestNG methods"],
            ["API", "Unique HTTP operations"],
            ["Performance", "Unique business journeys"],
            ["Angular", "Component/library test definitions"],
        ],
        widths=[2.0, 4.5],
    )
    para(
        "Counting units are reported separately and are not combined into a single overall total. "
        "Source populations are not presented as coverage totals."
    )
    doc.add_page_break()

    # ========== 3. INVENTORY SHARE ==========
    doc.add_heading("Universal Platform Share of Existing UI Regression Inventory", level=1)
    add_table(
        ["Platform", "Total regression inventory", "UP-scoped inventory", "UP share of inventory"],
        [
            ["V2 daily regression", str(V2_TOTAL), str(V2_SCOPED), f"{V2_PCT}%"],
            ["V3 regression", str(V3_TOTAL), str(V3_SCOPED), f"{V3_PCT}%"],
        ],
        widths=[1.8, 1.6, 1.6, 1.5],
    )
    para(
        "These percentages represent the share of the existing UI regression inventory mapped to "
        "the five Universal Platform workstreams. They are not requirement-level coverage percentages. "
        "Requirement-level coverage cannot be calculated until the Aha work items contain an enumerated "
        "requirements baseline.",
        italic=True,
        size=10,
    )
    doc.add_page_break()

    # ========== 4. LEADERSHIP DASHBOARD (landscape) ==========
    dash_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_header_footer(dash_section, landscape=True)

    doc.add_heading("Leadership Dashboard", level=1)
    para(
        "Scoped inventory by workstream. V2 and V3 UI counts reflect UP-scoped inventory only. "
        "API counts are unique operations. Performance counts are unique business journeys.",
        italic=True,
        size=10,
    )
    add_table(
        ["Workstream", "V2 UI inventory", "V3 UI inventory", "API", "Performance", "Current position", "Gap / opportunity"],
        [
            [
                "Universal Enrollment — ACS-I-2679",
                "151", "303", "5 operations", "7 journeys",
                "Established scheduled automation across V2 and V3",
                "Consolidation onto the target platform and continued business-flow expansion",
            ],
            [
                "IDP / Authentication — ACS-I-2680",
                "27", "56", "4 operations", "6 journeys",
                "Established UI automation with API and performance foundations",
                "API pipeline alignment and continued authentication-flow expansion",
            ],
            [
                "ABLE — Unite ABLE / CSR — ACS-I-2681",
                "17 dedicated cases plus nine-plan parameterized coverage",
                "Not applicable", "None identified", "Not applicable",
                "Established V2 foundation",
                "ABLE-specific API automation and suite scheduling alignment",
            ],
            [
                "ABLE — Entity Platform — ACS-I-2681",
                "Not applicable", "6 scenarios", "None identified", "1 journey",
                "V3 foundation implemented",
                "Pipeline enablement and plan expansion",
            ],
            [
                "Angular — ACS-I-2682",
                "Not applicable", "22 component tests (reported separately from E2E)",
                "Not applicable", "Not applicable",
                "Shared component automation implemented",
                "Pipeline scheduling and future business-flow alignment",
            ],
            [
                "Withdrawal — ACS-I-2690",
                "73", "20", "2 operations", "1 journey",
                "Established V2 and V3 UI automation with API/performance assets",
                "API gate alignment and targeted expansion",
            ],
        ],
        widths=[1.35, 1.05, 1.05, 0.75, 0.75, 1.35, 1.35],
        font_size=9.5,
    )

    # ========== 5. WORKSTREAM DETAIL (landscape section, same headers/footers) ==========
    doc.add_heading("Workstream Detail", level=1)

    doc.add_heading("Universal Enrollment — ACS-I-2679", level=2)
    para(
        "The 303 V3 Universal Enrollment methods are mapped at the suite level to ACS-I-2679. "
        "The V2 enrollment inventory contributes 151 additional UP-scoped cases. Together, enrollment "
        "automation spans standard, continue, multi-beneficiary, Upromise, PAGSP, matching-grant, "
        "and negative enrollment flows across multiple plan variants."
    )
    para("API: 5 unique enrollment/account operations. Performance: 7 enrollment business journeys.")

    doc.add_heading("IDP / Authentication — ACS-I-2680", level=2)
    para(
        "The scoped inventory includes 56 V3 IDP Login methods and 27 V2 Web Login cases. Other "
        "registration suites remain outside the scoped totals used in this assessment. Authentication "
        "flows include login, lockout, forgot username, and forgot password."
    )
    para("API: 4 unique authentication operations. Performance: 6 IDP-related business journeys.")

    doc.add_heading("ABLE — ACS-I-2681", level=2)
    doc.add_heading("Part A — Unite ABLE / CSR (V2)", level=3)
    para(
        "Established V2 automation includes 17 dedicated LA ABLE cases, nine ABLE plans "
        "(AKB, COB, ILB, MIB, NHB, NYB, PAB, RIB, TNB), and parameterized coverage embedded in "
        "the V2 regression framework. CSR-Able enrollment, profile maintenance, contributions, "
        "withdrawals, fee, P&E, and share adjustment flows are represented."
    )
    doc.add_heading("Part B — Entity Platform (V3)", level=3)
    para(
        "Six Entity Platform scenarios cover entity registration and entity dashboard login for MIB, "
        "NEB, and VAB. This foundation is implemented in V3 and available for continued regression "
        "expansion and pipeline alignment."
    )
    para("No ABLE-specific API automation was identified in the reviewed repositories.")

    doc.add_heading("Angular — ACS-I-2682", level=2)
    para(
        "The current Angular inventory consists of 22 lib-ui dynamic-forms component test definitions. "
        "These are component/library tests and are not presented as end-to-end business-flow coverage."
    )

    doc.add_heading("Withdrawal — ACS-I-2690", level=2)
    para(
        "The scoped inventory includes 73 V2 withdrawal cases, 20 V3 member-withdrawal methods, "
        "two unique API operations, and one performance journey. Coverage spans member and CSR "
        "distributions and systematic withdrawal flows."
    )

    # Return to portrait for remaining sections
    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_header_footer(portrait, landscape=False)

    # ========== 6. STRENGTHS ==========
    doc.add_heading("Current Automation Strengths", level=1)
    bullet("744-case V2 daily regression provides broad legacy and modern plan coverage maintained on a scheduled cadence.")
    bullet("436-method V3 regression suite represents the strategic modern platform with 379 methods directly scoped to UP workstreams.")
    bullet("Dual-framework coverage for enrollment, authentication, and withdrawal enables continuity during platform transition.")
    bullet("Nine ABLE plans plus dedicated LA ABLE suite on V2; Entity Platform foundation on V3.")
    bullet("Eleven unique API operations and fifteen performance journeys provide multi-layer automation beyond UI regression.")
    bullet("Twenty-two Angular lib-ui component tests support shared UI quality for dynamic forms.")

    # ========== 7. GAPS ==========
    doc.add_heading("Current Gaps and Improvement Opportunities", level=1)
    add_table(
        ["Area", "Current state", "Gap / improvement opportunity"],
        [
            [
                "ABLE-specific API automation",
                "No ABLE-specific API operations identified",
                "Net-new API automation is required",
            ],
            [
                "ABLE Entity Platform",
                "Six V3 scenarios implemented",
                "Pipeline enablement and plan expansion",
            ],
            [
                "Angular lib-ui",
                "Twenty-two component tests implemented",
                "Scheduled pipeline enablement and ongoing component expansion",
            ],
            [
                "API automation",
                "Eleven unique in-scope operations implemented",
                "Align automated operations with standard scheduled API gates",
            ],
            [
                "Performance automation",
                "Fifteen in-scope business journeys available",
                "Confirm recurring scheduling, workload standards, and SLA governance",
            ],
            [
                "V2/V3 alignment",
                "Automation exists on both frameworks",
                "Continue consolidation toward the strategic target platform while preserving regression coverage",
            ],
            [
                "Requirement traceability",
                "Aha work items define high-level outcomes",
                "Enumerated requirements would enable requirement-level coverage measurement",
            ],
        ],
        widths=[1.6, 2.2, 2.7],
    )

    # ========== 8. INTERPRETATION ==========
    doc.add_heading("Interpretation and Measurement Boundaries", level=1)
    bullet("V2 contains broader legacy and non-UP regression that is intentionally excluded from the 268-case scoped subtotal.")
    bullet("V2 and V3 are separate platform inventories and may contain functional overlap; they are not summed.")
    bullet("UI, API, performance, and component tests use different counting units and are reported separately.")
    bullet(f"The published {V2_PCT}% and {V3_PCT}% values measure inventory allocation, not requirement coverage.")
    bullet("Requirement-level coverage percentages require an enumerated requirement baseline from the Aha work items.")
    bullet("ABLE Entity Platform (6 scenarios) and Angular lib-ui (22 component tests) are reported separately from the V3 scoped UI subtotal of 379.")

    # ========== 9. CONCLUSION ==========
    doc.add_heading("Conclusion", level=1)
    para(
        "The assessment confirms a substantial and actively maintained automation foundation across "
        "the five Universal Platform workstreams. The current inventory provides a measurable baseline "
        "for future expansion, consolidation, and requirement-level traceability once an enumerated "
        "requirements baseline is available."
    )

    # Metadata
    core = doc.core_properties
    core.title = "Universal Platform Test Automation Coverage Assessment"
    core.author = "QA Automation"
    core.subject = "Universal Platform automated test inventory and business-flow coverage"
    core.category = VERSION
    core.comments = STATUS

    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    return DOCX_OUT


def update_word_fields(docx_path: Path) -> dict:
    """Update TOC and page fields via Word COM; return TOC page map for validation."""
    result = {"updated": False, "toc_entries": [], "page_count": 0}
    try:
        import win32com.client  # type: ignore
    except ImportError:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pywin32", "-q"])
            import win32com.client  # type: ignore
        except Exception:
            return result
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        # Two-pass update ensures TOC reflects final pagination
        for _ in range(2):
            doc.Fields.Update()
            for toc in doc.TablesOfContents:
                toc.Update()
        result["page_count"] = doc.ComputeStatistics(2)  # wdStatisticPages
        for toc in doc.TablesOfContents:
            for i in range(1, toc.Range.Paragraphs.Count + 1):
                text = toc.Range.Paragraphs(i).Range.Text.strip()
                if text:
                    result["toc_entries"].append(text)
        doc.Save()
        doc.Close()
        result["updated"] = True
        return result
    except Exception as exc:
        print("Word field update failed:", exc)
        return result
    finally:
        if word is not None:
            word.Quit()


def build_pdf(docx_path: Path) -> Path:
    try:
        import docx2pdf
        docx2pdf.convert(str(docx_path), str(PDF_OUT))
        if PDF_OUT.exists():
            return PDF_OUT
    except Exception as exc:
        print("docx2pdf failed:", exc)
    raise RuntimeError("PDF generation failed")


def main():
    ensure_docx()
    docx = build()
    print("WROTE", docx)
    field_result = update_word_fields(docx)
    print("TOC/fields updated via Word:", field_result.get("updated"))
    print("Word page count:", field_result.get("page_count"))
    for entry in field_result.get("toc_entries", []):
        print("  TOC:", entry)
    pdf = build_pdf(docx)
    print("WROTE", pdf)
    return docx, pdf, field_result


if __name__ == "__main__":
    main()
