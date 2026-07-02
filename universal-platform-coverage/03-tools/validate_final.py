#!/usr/bin/env python3
"""Final validation for leadership assessment deliverables."""
from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "02-deliverables/Universal-Platform-Test-Automation-Coverage-Assessment-Final.docx"
PDF = ROOT / "02-deliverables/Universal-Platform-Test-Automation-Coverage-Assessment-Final.pdf"
EMAIL = ROOT / "02-deliverables/Kevin-Rajib-Final-Update-Email.md"

PROHIBITED = [
    "pending sme", "near-final", "provisional", "validation required",
    "nick checklist", "recommended ownership", "pass/fail", " failed",
    " passed", " red ", " green ", "evidence id", "cursor", "\\\\",
    "prepared for kevin", "prepared for rajib", "sme factual review",
    "safe after", "complete nick",
]

REQUIRED = ["744", "268", "476", "436", "379", "57", "36.0%", "86.9%", "11", "15"]
REQUIRED_DOCX = ["Final Current-State Assessment"]
REQUIRED_PDF_COVER = ["Final Current", "State Assessment"]


def docx_text(path: Path) -> str:
    from docx import Document
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            parts.append(" | ".join(c.text for c in r.cells))
    return "\n".join(parts)


def pdf_text(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        from pypdf import PdfReader
    return "\n".join((p.extract_text() or "") for p in PdfReader(str(path)).pages)


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def main():
    errors = []
    # Arithmetic
    assert 744 - 268 == 476
    assert round(268 / 744 * 100, 1) == 36.0
    assert 436 - 379 == 57
    assert round(379 / 436 * 100, 1) == 86.9
    assert 151 + 27 + 17 + 73 == 268
    assert 303 + 56 + 20 == 379
    assert 5 + 4 + 2 == 11
    assert 7 + 6 + 1 + 1 == 15
    print("Arithmetic: OK")

    if not DOCX.exists():
        errors.append("DOCX missing")
    if not PDF.exists():
        errors.append("PDF missing")

    dx = docx_text(DOCX) if DOCX.exists() else ""
    pd = pdf_text(PDF) if PDF.exists() else ""

    for label, text in [("DOCX", dx), ("PDF", pd), ("EMAIL", EMAIL.read_text(encoding="utf-8") if EMAIL.exists() else "")]:
        low = text.lower()
        for term in PROHIBITED:
            if term in low:
                errors.append(f"{label}: prohibited '{term.strip()}'")
        for req in REQUIRED:
            if req not in text and label != "EMAIL":
                errors.append(f"{label}: missing '{req}'")
        if label == "DOCX":
            for req in REQUIRED_DOCX:
                if req not in text:
                    errors.append(f"{label}: missing '{req}'")
        if label == "PDF":
            for req in REQUIRED_PDF_COVER:
                if req not in text:
                    errors.append(f"{label}: missing cover '{req}'")

    if EMAIL.exists() and word_count(EMAIL.read_text(encoding="utf-8")) > 200:
        errors.append(f"Email word count {word_count(EMAIL.read_text(encoding='utf-8'))} exceeds 200")

    try:
        from PyPDF2 import PdfReader
    except ImportError:
        from pypdf import PdfReader
    pdf_pages = len(PdfReader(str(PDF)).pages) if PDF.exists() else 0
    print(f"PDF pages: {pdf_pages}")

    if errors:
        print("VALIDATION FAILED:")
        for e in errors:
            print(" -", e)
        sys.exit(1)
    print("VALIDATION: PASS")
    return pdf_pages


if __name__ == "__main__":
    main()
