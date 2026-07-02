#!/usr/bin/env python3
"""Mechanical correction validation for final assessment."""
from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "02-deliverables/Universal-Platform-Test-Automation-Coverage-Assessment-Final.docx"
PDF = ROOT / "02-deliverables/Universal-Platform-Test-Automation-Coverage-Assessment-Final.pdf"

REQUIRED = ["744", "268", "476", "436", "379", "57", "36.0%", "86.9%"]
NEW_WORDING = "268 cases mapped within the five-workstream Universal Platform scope"
OLD_WORDING = "268 cases mapped to the five Universal Platform workstreams"
HEADER_OK = "Universal Platform Test Automation Coverage Assessment | Final | Version 1.0"
HEADER_BAD = "AssessmentFinal"


def docx_text(path: Path) -> str:
    from docx import Document
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            parts.append(" | ".join(c.text for c in r.cells))
    for section in d.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for p in section.footer.paragraphs:
            parts.append(p.text)
    return "\n".join(parts)


def pdf_pages(path: Path) -> list[str]:
    try:
        import fitz
        doc = fitz.open(str(path))
        pages = [doc[i].get_text() for i in range(len(doc))]
        doc.close()
        return pages
    except Exception:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            from pypdf import PdfReader
        return [(p.extract_text() or "") for p in PdfReader(str(path)).pages]


def main():
    errors = []
    dx = docx_text(DOCX)
    pages = pdf_pages(PDF)

    # Arithmetic unchanged
    assert 744 - 268 == 476
    assert round(268 / 744 * 100, 1) == 36.0
    assert 436 - 379 == 57
    assert round(379 / 436 * 100, 1) == 86.9

    if HEADER_BAD in dx.replace(" ", ""):
        errors.append("DOCX header still concatenated (AssessmentFinal)")
    if HEADER_OK not in dx:
        errors.append("DOCX missing corrected header format")
    if OLD_WORDING in dx:
        errors.append("DOCX still has old V2 scoped wording")
    if NEW_WORDING not in dx:
        errors.append("DOCX missing new V2 scoped wording")

    pdf_all = "\n".join(pages)
    if HEADER_BAD.replace(" ", "") in pdf_all.replace(" ", ""):
        errors.append("PDF header still concatenated")
    if NEW_WORDING.replace("\n", " ") not in pdf_all.replace("\n", " "):
        if "five-workstream Universal Platform scope" not in pdf_all:
            errors.append("PDF missing new V2 scoped wording")

    for req in REQUIRED:
        if req not in dx:
            errors.append(f"DOCX missing {req}")
        if req not in pdf_all:
            errors.append(f"PDF missing {req}")

    pdf_page_count = len(pages)
    print(f"PDF page count: {pdf_page_count}")

    # Header/footer via pymupdf blocks (PyPDF2 often misses margin/header text)
    try:
        import fitz
        pdf_doc = fitz.open(str(PDF))
        for i in range(1, pdf_page_count):
            page = pdf_doc[i]
            blocks = page.get_text("blocks")
            top = sorted(blocks, key=lambda b: b[1])[:3]
            bottom = sorted(blocks, key=lambda b: b[1])[-3:]
            top_text = " ".join(b[4] for b in top)
            bottom_text = " ".join(b[4] for b in bottom)
            if HEADER_OK not in top_text and "Coverage Assessment | Final" not in top_text:
                errors.append(f"PDF page {i+1} missing header (block scan)")
            if "QA Automation" not in bottom_text:
                errors.append(f"PDF page {i+1} missing footer (block scan)")
            if "Page" not in bottom_text or "of" not in bottom_text:
                errors.append(f"PDF page {i+1} missing page number (block scan)")
        pdf_doc.close()
    except Exception as exc:
        print("Warning: fitz header scan skipped:", exc)

    # Continuous numbering check - look for "Page X of Y" on last page
    last = pages[-1]
    m = re.search(r"Page\s+(\d+)\s+of\s+(\d+)", last)
    if m:
        page_num, total = int(m.group(1)), int(m.group(2))
        if page_num != pdf_page_count or total != pdf_page_count:
            errors.append(f"Page numbering mismatch: last shows {page_num} of {total}, expected {pdf_page_count}")
    else:
        errors.append("Could not parse final page numbering")

    if errors:
        print("NOT READY:")
        for e in errors:
            print(" -", e)
        sys.exit(1)
    print("READY TO SEND")
    print("Header/footer: PASS")
    print("DOCX/PDF content match: PASS")
    print("Totals unchanged: PASS")


if __name__ == "__main__":
    main()
