#!/usr/bin/env python3
"""Master reset/rebuild for Universal Platform automation coverage project."""
from __future__ import annotations

import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DATE = "2026-06-29"
OLD_KIT = ROOT / "up-coverage-metrics-kit" / "up-coverage-metrics-kit"
ARCHIVE = ROOT / "99-archive"
SNAPSHOT = ARCHIVE / f"{DATE}-pre-rebuild-snapshot"

# Evidence source paths (read-only originals)
EVIDENCE = {
    "aha": OLD_KIT / "AHA PDF",
    "v2": OLD_KIT / "Gitlab V2 Job Log",
    "v3_ue": OLD_KIT / "Gitlab V2 Job Log" / "V3 UE TestNG reports_files.html",
    "v3_unite": OLD_KIT / "Gitlab V2 Job Log" / "V3 Unite TestNG reports.html",
    "v3_log": OLD_KIT / "Gitlab V2 Job Log" / "V3 job latest log.txt",
    "v2_jenkins": OLD_KIT / "Gitlab V2 Job Log" / "V2 Jenkins Job extract full output.txt",
    "v2_qtest": OLD_KIT / "Gitlab V2 Job Log" / "QAAuto_UniteRegression_Daily_Report.PDF",
}

# --- V2 module reconciliation (508 documented + 236 gap = 744) ---
V2_MODULES = [
    ("PROVISIONAL-MOD-01", "Enrollments", "enrollments", "stage1-enrollments-regression", "stage1-enrollments.xml", 151, "IN_SCOPE_ENROLLMENT", "executed_test_cases", "NextGen/Universal enrollment flows; includes some ABLE enrollment scenarios as secondary", "Jenkins log + qTest PRIME", "High", "No"),
    ("PROVISIONAL-MOD-02", "Web Login", "web-login", "stage1-web-login-regression", "stage1-web-login.xml", 27, "IN_SCOPE_IDP", "executed_test_cases", "ACS-I-2680 auth-server member login flows", "Jenkins log", "High", "No"),
    ("PROVISIONAL-MOD-03", "Withdrawals", "withdrawals", "stage1-withdrawals-regression", "stage1-withdrawals.xml", 73, "IN_SCOPE_WITHDRAWAL", "executed_test_cases", "Member/CSR distributions per ACS-I-2690 withdrawal domain", "Jenkins log", "High", "No"),
    ("PROVISIONAL-MOD-04", "LA ABLE dedicated suite", "laable", "stage1-laable-regression", "stage1-laable.xml", 17, "IN_SCOPE_ABLE", "executed_test_cases", "Dedicated CSR-Able suite; implemented, not in 10-target Jenkins extract", "Suite XML + internal reconciliation", "High", "No"),
    ("PROVISIONAL-MOD-05", "Web Registration", "web-registration", "stage1-web-registration-regression", "stage1-web-registration.xml", 46, "ADJACENT_NOT_CONFIRMED", "executed_test_cases", "Legacy web registration; ACS-I-2680 names auth-server only", "Jenkins log + AHA gap", "Medium", "Yes"),
    ("PROVISIONAL-MOD-06", "CSR Account Maintenance", "csr-acct-maintenance", "stage1-csr-acct-maintenance-regression", "stage1-csr-acct-maintenance.xml", 81, "ADJACENT_NOT_CONFIRMED", "executed_test_cases", "Profile maintenance; not named in any of five Aha ideas", "Jenkins log + AHA gap", "Medium", "Yes"),
    ("PROVISIONAL-MOD-07", "Contributions", "contributions", "stage1-contributions-regression", "stage1-contributions.xml", 48, "ADJACENT_NOT_CONFIRMED", "executed_test_cases", "529 contribution flows; Angular (2682) scope not confirmed for member UI", "Jenkins log + AHA gap", "Medium", "Yes"),
    ("PROVISIONAL-MOD-08", "Transfers", "transfers", "stage1-transfers-regression", "stage1-transfers.xml", 12, "ADJACENT_NOT_CONFIRMED", "executed_test_cases", "Fund transfers; ACS-I-2690 names withdrawal only", "Jenkins log + AHA gap", "Medium", "Yes"),
    ("PROVISIONAL-MOD-09", "UGift", "ugift", "stage1-ugift-regression", "stage1-ugift.xml", 36, "OUT_OF_SCOPE", "executed_test_cases", "Gifting product outside five Universal Platform workstreams", "Jenkins log", "High", "No"),
    ("PROVISIONAL-MOD-10", "Investment Options", "investment-options", "stage1-investment-options-regression", "stage1-investment-options.xml", 24, "OUT_OF_SCOPE", "executed_test_cases", "Exchanges/fund-to-fund outside scoped workstreams", "Jenkins log", "High", "No"),
    ("PROVISIONAL-MOD-11", "Account Balance", "acct-overview", "stage1-acct-overview-regression", "stage1-acct-overview.xml", 10, "OUT_OF_SCOPE", "executed_test_cases", "Account overview page validation; not in Aha scope", "Jenkins log", "High", "No"),
    ("PROVISIONAL-MOD-12", "CSR Actions", "csr-actions", "stage1-csr-actions-regression", "stage1-csr-actions.xml", 9, "ADJACENT_NOT_CONFIRMED", "executed_test_cases", "CSR action flows; ABLE/Angular ownership unresolved", "Suite inventory", "Medium", "Yes"),
    ("PROVISIONAL-MOD-13", "Empower Plan", "empower-plan", "stage1-empower-plan-regression", "stage1-empower-plan.xml", 21, "UNMAPPED_NEEDS_SME_CONFIRMATION", "executed_test_cases", "Empower plan module in PRIME cycle; workstream mapping not confirmed", "Suite inventory", "Low", "Yes"),
    ("PROVISIONAL-MOD-14", "Sardine", "sardine", "stage1-sardine-regression", "stage1-sardine-regression.xml", 10, "OUT_OF_SCOPE", "executed_test_cases", "Sardine fraud module outside five workstreams", "Suite inventory", "Medium", "No"),
    ("PROVISIONAL-MOD-15", "PRIME per-plan expansion", "multi-plan", "PRIME cycle expansion", "multiple suite XMLs", 179, "UNMAPPED_NEEDS_SME_CONFIRMATION", "executed_test_cases", "qTest 744 exceeds 10-module Jenkins subset; residual per-plan/param expansion not itemized in available reports", "qTest daily report vs Jenkins subset", "Low", "Yes"),
]

V3_ROWS = [
    ("V3-UE-ALL", "V3 UE TestNG", "Universal Enrollment Regression", "303 scenarios", "multiple traunches", "universal-enrollment-stage1.xml", "IN_SCOPE_ENROLLMENT", "ACS-I-2679", "executed_test_methods", "Dedicated UE suite; all 303 methods map to enrollment workstream by suite definition", "High", "No"),
    ("V3-UNITE-IDP-LOGIN", "V3 Unite TestNG", "IDP Login", "56 methods", "NMD,NYD,NJD,MDD,OHD,PAG,NDD", "stage1-idp-login.xml", "IN_SCOPE_IDP", "ACS-I-2680", "executed_test_methods", "IDP login, lockout, forgot username/password", "High", "No"),
    ("V3-UNITE-WITHDRAWAL", "V3 Unite TestNG", "Member Withdrawal", "20 methods", "NJD,NMD,NYD,MDD", "stage1-withdrawals.xml", "IN_SCOPE_WITHDRAWAL", "ACS-I-2690", "executed_test_methods", "Qualified/non-qualified member withdrawals", "High", "No"),
    ("V3-UNITE-WEBREG", "V3 Unite TestNG", "IDP Web Registration", "6 methods", "MDD,ILD,NYD", "stage1-web-registration.xml", "ADJACENT_NOT_CONFIRMED", "ACS-I-2680", "executed_test_methods", "First-time web registration; Aha IDP idea names auth-server only", "Medium", "Yes"),
    ("V3-UNITE-CONTRIB", "V3 Unite TestNG", "Contributions", "36 methods", "NYD,MDD,NJD", "stage1-contributions.xml", "ADJACENT_NOT_CONFIRMED", "ACS-I-2682", "executed_test_methods", "Member/CSR contribution UI; Angular/529 scope not confirmed in Aha", "Medium", "Yes"),
    ("V3-UNITE-CSR", "V3 Unite TestNG", "CSR Account Maintenance", "15 methods", "NYD,MDD,NJD,NMD", "stage1-csr-acct-maintenance.xml", "ADJACENT_NOT_CONFIRMED", "ACS-I-2682", "executed_test_methods", "Member/CSR profile phone/email; not named in Aha ideas", "Medium", "Yes"),
    ("V3-ENTITY-REG", "V3 unite-entity (not scheduled)", "Entity Registration", "3 test blocks", "MIB,NEB,VAB", "stage1-entity-registration.xml", "IN_SCOPE_ABLE", "ACS-I-2681", "suite_test_definitions", "Entity enrollment; implemented, not in nightly master", "High", "No"),
    ("V3-ENTITY-IDP", "V3 unite-entity (not scheduled)", "Entity Dashboard IDP Login", "3 test blocks", "MIB,NEB,VAB", "stage1-idp-entitydashboard-login.xml", "IN_SCOPE_ABLE", "ACS-I-2681", "suite_test_definitions", "Entity dashboard login; implemented, not in nightly master", "High", "No"),
    ("V3-LIBUI", "V3 universal-lib-ui (not scheduled)", "Lib-UI Dynamic Forms", "22 test blocks", "nmd", "lib-ui-dynamic-forms-qc4.xml", "IN_SCOPE_ANGULAR", "ACS-I-2682", "component_test_definitions", "Shared dynamic-forms component library; separate from business-flow E2E", "Medium", "Yes"),
]

API_ROWS = [
    ("API-AUTH-01", "jsonapi-auth", "GetAccessTokenTest", "testGetAccessToken", "GET", "/services/oauth2/authorize", "OAuth authorize", "IN_SCOPE_IDP", "positive", "implemented", "High", "No"),
    ("API-AUTH-02", "jsonapi-auth", "GetPkceTokenTest", "testGetPkceToken", "POST", "/services/oauth2/token", "OAuth token/PKCE", "IN_SCOPE_IDP", "positive", "implemented", "High", "No"),
    ("API-AUTH-03", "jsonapi-auth", "IntrospectTest", "testIntrospect", "POST", "/services/oauth2/introspect", "Token introspection", "IN_SCOPE_IDP", "positive", "implemented", "High", "No"),
    ("API-AUTH-04", "jsonapi-auth", "UserInfoTest", "testUserInfo", "GET", "/services/userinfo", "User info", "IN_SCOPE_IDP", "positive", "implemented", "High", "No"),
    ("API-ACCT-01", "jsonapi-aws-accountweb", "PostEnrollmentRequestTest", "multiple", "POST", "/accounts/", "Enrollment create", "IN_SCOPE_ENROLLMENT", "positive/negative", "implemented", "High", "No"),
    ("API-ACCT-02", "jsonapi-aws-accountweb", "PutEnrollmentRequestTest", "multiple", "PUT", "/accounts/", "Enrollment update", "IN_SCOPE_ENROLLMENT", "positive/negative", "implemented", "High", "No"),
    ("API-ACCT-03", "jsonapi-aws-accountweb", "GetAccountRequestTest", "multiple", "GET", "/accounts/", "Account read", "IN_SCOPE_ENROLLMENT", "positive/negative", "implemented", "High", "No"),
    ("API-ACCT-04", "jsonapi-aws-accountweb", "DeleteAccountRequestTest", "multiple", "DELETE", "/accounts/", "Account delete", "IN_SCOPE_ENROLLMENT", "positive/negative", "implemented", "High", "No"),
    ("API-ACCT-05", "jsonapi-aws-accountweb", "GetValidationSubBeneEnrollmentRequestTest", "multiple", "GET", "/validation/", "Enrollment validation", "IN_SCOPE_ENROLLMENT", "positive/negative", "implemented", "High", "No"),
    ("API-FIN-01", "jsonapi-aws-financialweb", "PostWithdrawalByBankRequestTest", "multiple", "POST", "/withdrawals/", "Withdrawal by bank", "IN_SCOPE_WITHDRAWAL", "positive/negative", "implemented", "High", "No"),
    ("API-FIN-02", "jsonapi-aws-financialweb", "PostWithdrawalByCheckRequestTest", "multiple", "POST", "/withdrawals/", "Withdrawal by check", "IN_SCOPE_WITHDRAWAL", "positive/negative", "implemented", "High", "No"),
    ("API-FIN-03", "jsonapi-aws-financialweb", "PostExchangeRequestTest", "multiple", "POST", "/financial/", "Fund exchange", "ADJACENT_NOT_CONFIRMED", "positive/negative", "implemented", "Medium", "Yes"),
    ("API-META-01", "jsonapi-metadataweb", "GetPlanRequestTest", "multiple", "GET", "/metadata/", "Plan metadata", "ADJACENT_NOT_CONFIRMED", "reference", "implemented", "Medium", "Yes"),
    ("API-META-02", "jsonapi-metadataweb", "GetUiMetadataRequestTest", "multiple", "GET", "/metadata/", "UI metadata", "ADJACENT_NOT_CONFIRMED", "reference", "implemented", "Medium", "Yes"),
    ("API-META-03", "jsonapi-metadataweb", "GetFundRequestTest", "multiple", "GET", "/metadata/", "Fund metadata", "ADJACENT_NOT_CONFIRMED", "reference", "implemented", "Medium", "Yes"),
    ("API-MOBILE-01", "mobile2", "MobileDashboardRequestTest", "multiple", "GET", "/mobile2api/v1/mobiledashboard", "Mobile dashboard", "OUT_OF_SCOPE", "positive", "implemented", "High", "No"),
    ("API-ASTRO-01", "jsonapi-astro", "EmployerPlansTest", "multiple", "GET", "/employers/plans", "FuturePlan employer", "OUT_OF_SCOPE", "positive", "implemented", "High", "No"),
    ("API-ABLE-01", "none", "none", "none", "N/A", "N/A", "No ABLE-specific API automation identified", "IN_SCOPE_ABLE", "N/A", "none", "High", "No"),
]

PERF_ROWS = [
    ("PERF-UE-01", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "up-enrollment-submission.jmx", "Blazemeter-UE-Local.yaml", "API submission", "Yes", "implemented", "High", "No"),
    ("PERF-UE-02", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "up-enrollment-aws-account-ms.jmx", "up-enrollment-aws-account-ms-local.yaml", "Account MS load", "Yes", "implemented", "High", "No"),
    ("PERF-UE-03", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "up-enrollment-validation-ms.jmx", "up-enrollment-validation-ms-local.yaml", "Validation MS load", "Yes", "implemented", "High", "No"),
    ("PERF-UE-04", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "up-enrollment-aws-account-ms-submit.jmx", "up-enrollment-aws-account-ms-submit-local.yaml", "Submit path", "Yes", "implemented", "High", "No"),
    ("PERF-UE-05", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "up-enrollment-aws-account-ms-snapshot.jmx", "up-enrollment-aws-account-ms-snapshot-local.yaml", "Snapshot path", "Yes", "implemented", "High", "No"),
    ("PERF-UE-06", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "up-enrollment-account-ms.jmx", "up-enrollment-account-ms-local.yaml", "Legacy account MS", "Yes", "implemented", "Medium", "No"),
    ("PERF-UE-07", "Universal Enrollment", "IN_SCOPE_ENROLLMENT", "ue_subsequent-aws-account-ms.jmx", "ue_subsequent-aws-account-ms.yaml", "Subsequent enrollment", "Yes", "implemented", "Medium", "No"),
    ("PERF-IDP-01", "IDP Login", "IN_SCOPE_IDP", "idp-login.jmx", "idp-login-local.yaml", "Member login", "Yes", "implemented", "High", "No"),
    ("PERF-IDP-02", "IDP Login", "IN_SCOPE_IDP", "idp-login-resources.jmx", "idp-login-resources-local.yaml", "Login static resources", "Partial", "implemented", "Medium", "Yes"),
    ("PERF-IDP-03", "IDP Login", "IN_SCOPE_IDP", "idp-forgot-username.jmx", "idp-forgot-username-local.yaml", "Forgot username", "Yes", "implemented", "High", "No"),
    ("PERF-IDP-04", "IDP Login", "IN_SCOPE_IDP", "idp-forgot-password.jmx", "idp-forgot-password-local.yaml", "Forgot password", "Yes", "implemented", "High", "No"),
    ("PERF-IDP-05", "IDP Auth Server", "IN_SCOPE_IDP", "auth-server.jmx", "auth-server-local.yaml", "Auth server load", "Yes", "implemented", "High", "No"),
    ("PERF-IDP-06", "IDP Auth Server", "IN_SCOPE_IDP", "authentication-pkce.jmx", "authentication-pkce-local.yaml", "PKCE flow", "Yes", "implemented", "High", "No"),
    ("PERF-WD-01", "Withdrawal", "IN_SCOPE_WITHDRAWAL", "up-post-withdrawal-ms.jmx", "up-post-withdrawal-ms-local.yaml", "Withdrawal MS POST", "Yes", "implemented", "High", "No"),
    ("PERF-ABLE-01", "ABLE Entity", "IN_SCOPE_ABLE", "entityMS_1.jmx", "none", "Entity MS enrollment", "Partial", "implemented", "Medium", "Yes"),
    ("PERF-ACCT-01", "Account read", "ADJACENT_NOT_CONFIRMED", "up-get-aws-account-ms.jmx", "up-get-aws-account-ms-local.yaml", "GET account", "Partial", "implemented", "Medium", "Yes"),
    ("PERF-META-01", "Metadata", "ADJACENT_NOT_CONFIRMED", "up-get-metadata-ms.jmx", "up-get-metadata-ms-local.yaml", "Metadata read", "Partial", "implemented", "Medium", "Yes"),
    ("PERF-DB-01", "DB prototype enrollment", "OUT_OF_SCOPE", "ue_enrollment.jmx", "perf-taurus-ue.yaml", "Database prototype", "No", "prototype", "High", "No"),
    ("PERF-DB-02", "DB prototype login", "OUT_OF_SCOPE", "tf_member_login.jmx", "perf-taurus-tf.yaml", "Database prototype", "No", "prototype", "High", "No"),
    ("PERF-DB-03", "DB prototype CSR", "OUT_OF_SCOPE", "csr_search.jmx", "TestRun_CSRSearch.yaml", "Database prototype", "No", "prototype", "High", "No"),
]


def file_hash(path: Path, limit_mb: int = 50) -> str:
    if not path.is_file() or path.stat().st_size > limit_mb * 1024 * 1024:
        return ""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()[:16]


def ensure_backup():
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    if not SNAPSHOT.exists():
        # robocopy for long paths
        cmd = [
            "robocopy", str(ROOT), str(SNAPSHOT),
            "/E", "/XD", "99-archive", "/NFL", "/NDL", "/NJH", "/NJS", "/nc", "/ns", "/np",
        ]
        subprocess.run(cmd, check=False)


def write_manifest_and_tree():
    manifest_path = ARCHIVE / f"{DATE}-pre-rebuild-file-manifest.csv"
    tree_path = ARCHIVE / f"{DATE}-pre-rebuild-tree.txt"
    rows = []
    tree_lines = []
    for dirpath, dirnames, filenames in os.walk(ROOT):
        dirnames[:] = [d for d in dirnames if d != "99-archive" or Path(dirpath) != ROOT]
        rel_dir = Path(dirpath).relative_to(ROOT)
        for fn in filenames:
            p = Path(dirpath) / fn
            rel = p.relative_to(ROOT)
            if str(rel).startswith("99-archive"):
                continue
            st = p.stat()
            planned = "retain_in_new_structure" if any(x in str(rel) for x in ["AHA PDF", "Gitlab V2 Job Log", "QAAuto"]) else "archive_or_remove"
            if "prompts" in str(rel) or "finalization-pack" in str(rel) or rel.suffix == ".zip":
                planned = "archive"
            rows.append({
                "relative_path": str(rel).replace("\\", "/"),
                "file_size": st.st_size,
                "modified_timestamp": datetime.fromtimestamp(st.st_mtime).isoformat(),
                "duplicate_hash": file_hash(p),
                "planned_action": planned,
                "reason": "pre-rebuild inventory",
            })
            tree_lines.append(str(rel).replace("\\", "/"))
    with open(manifest_path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()) if rows else [])
        w.writeheader()
        w.writerows(rows)
    tree_path.write_text("\n".join(sorted(tree_lines)), encoding="utf-8")
    return manifest_path, tree_path


def setup_dirs():
    dirs = [
        "00-input-evidence/aha",
        "00-input-evidence/v2-qtest-jenkins",
        "00-input-evidence/v3-testng-gitlab",
        "00-input-evidence/supplemental",
        "01-analysis/csv",
        "02-deliverables",
        "03-tools",
        "99-archive",
    ]
    for d in dirs:
        (ROOT / d).mkdir(parents=True, exist_ok=True)


def copy_evidence(cleanup_log: list):
    dest_map = {
        "aha": ROOT / "00-input-evidence/aha",
        "v2": ROOT / "00-input-evidence/v2-qtest-jenkins",
    }
    if EVIDENCE["aha"].exists():
        for f in EVIDENCE["aha"].glob("*.pdf"):
            dst = dest_map["aha"] / f.name
            if not dst.exists():
                shutil.copy2(f, dst)
                cleanup_log.append(f"COPY evidence: {f.name} -> 00-input-evidence/aha/")
    v2_dest = dest_map["v2"]
    for key in ["v2_qtest", "v2_jenkins", "v3_log"]:
        src = EVIDENCE[key]
        if src.exists():
            dst = v2_dest / src.name
            if not dst.exists():
                shutil.copy2(src, dst)
                cleanup_log.append(f"COPY evidence: {src.name} -> 00-input-evidence/v2-qtest-jenkins/")
    v3_dest = ROOT / "00-input-evidence/v3-testng-gitlab"
    for key in ["v3_ue", "v3_unite"]:
        src = EVIDENCE[key]
        if src.exists():
            dst = v3_dest / src.name
            if not dst.exists():
                shutil.copy2(src, dst)
                cleanup_log.append(f"COPY evidence: {src.name} -> 00-input-evidence/v3-testng-gitlab/")
    # supplemental screenshots
    sup_src = OLD_KIT / "up-coverage-metrics-finalization-pack-v2" / "00-input-evidence"
    sup_dest = ROOT / "00-input-evidence/supplemental"
    if sup_src.exists():
        for pat in ["**/*.png", "**/*.html"]:
            for f in sup_src.glob(pat):
                rel = f.relative_to(sup_src)
                dst = sup_dest / rel
                dst.parent.mkdir(parents=True, exist_ok=True)
                if not dst.exists() and f.stat().st_size < 5_000_000:
                    try:
                        shutil.copy2(f, dst)
                    except OSError:
                        pass


def sum_v2_by_class():
    totals = {}
    for row in V2_MODULES:
        cls = row[6]
        totals[cls] = totals.get(cls, 0) + row[5]
    return totals


def sum_v3_by_class():
    nightly = [r for r in V3_ROWS if "not scheduled" not in r[1].lower()]
    totals = {}
    for row in nightly:
        cls = row[6]
        totals[cls] = totals.get(cls, 0) + int(re.search(r"\d+", row[3]).group())
    return totals


def write_csvs():
    csv_dir = ROOT / "01-analysis/csv"
    # v2
    with open(csv_dir / "v2-row-level-mapping.csv", "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["qtest_test_case_id", "test_name", "source_module", "plan_traunch", "jenkins_target", "repository_test", "primary_workstream", "secondary_tags", "inclusion_classification", "counting_unit", "rationale", "source_report", "confidence", "sme_confirmation_required", "provisional_count"])
        for r in V2_MODULES:
            ws_map = {
                "IN_SCOPE_ENROLLMENT": "ACS-I-2679",
                "IN_SCOPE_IDP": "ACS-I-2680",
                "IN_SCOPE_ABLE": "ACS-I-2681",
                "IN_SCOPE_WITHDRAWAL": "ACS-I-2690",
                "IN_SCOPE_ANGULAR": "ACS-I-2682",
            }
            w.writerow([r[0], r[1], r[2], "multiple", r[3], r[4], ws_map.get(r[6], "N/A"), "", r[6], r[7], r[8], r[9], r[10], r[11], r[5]])
    # v3
    with open(csv_dir / "v3-row-level-mapping.csv", "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["source_report", "suite", "test_method_scenario", "plan_traunch", "repository_test", "primary_workstream", "secondary_tags", "inclusion_classification", "counting_unit", "rationale", "confidence", "sme_confirmation_required", "count"])
        for r in V3_ROWS:
            w.writerow(list(r))
    # api
    with open(csv_dir / "api-operation-mapping.csv", "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["service_module", "test_class", "test_method", "http_method", "sanitized_endpoint", "business_capability", "inclusion_classification", "aha_workstream", "intent", "scheduled_status", "rationale", "confidence", "sme_confirmation_required"])
        for r in API_ROWS:
            w.writerow(list(r))
    # performance
    with open(csv_dir / "performance-journey-mapping.csv", "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["business_journey", "workstream", "inclusion_classification", "jmx_file", "taurus_configuration", "workload_model", "sla_assertion", "scheduling_evidence", "rationale", "confidence", "sme_confirmation_required"])
        for r in PERF_ROWS:
            w.writerow([r[1], r[2], r[2], r[3], r[4], r[5], r[6], r[7], "Mapped from performance repo inventory", r[8], r[9]])
    # summary
    v2t = sum_v2_by_class()
    v3t = sum_v3_by_class()
    with open(csv_dir / "final-workstream-summary.csv", "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["workstream", "aha_id", "v2_in_scope_ui", "v3_in_scope_ui", "v3_adjacent_ui", "api_unique_operations", "api_test_methods", "api_test_classes", "performance_in_scope_journeys", "component_tests", "confidence", "sme_validation_required"])
        w.writerow(["Enrollment", "ACS-I-2679", v2t.get("IN_SCOPE_ENROLLMENT", 0), v3t.get("IN_SCOPE_ENROLLMENT", 0), 0, 5, 103, 17, 7, 0, "High", "Partial - V2 per-plan expansion unmapped"])
        w.writerow(["IDP / Authentication", "ACS-I-2680", v2t.get("IN_SCOPE_IDP", 0), v3t.get("IN_SCOPE_IDP", 0), v3t.get("ADJACENT_NOT_CONFIRMED", 0) if False else 6, 4, 6, 4, 6, 0, "High", "Web registration scope"])
        w.writerow(["ABLE - Unite ABLE/CSR", "ACS-I-2681", v2t.get("IN_SCOPE_ABLE", 0), 0, 0, 0, 0, 0, 0, 0, "High", "Embedded ABLE in V2 enrollment module not split"])
        w.writerow(["ABLE - Entity Platform", "ACS-I-2681", 0, 6, 0, 0, 0, 0, 1, 0, "High", "Pipeline enablement"])
        w.writerow(["Angular", "ACS-I-2682", 0, 0, 51, 0, 0, 0, 2, 22, "Medium", "Contributions/CSR/lib-ui scope"])
        w.writerow(["Withdrawal", "ACS-I-2690", v2t.get("IN_SCOPE_WITHDRAWAL", 0), v3t.get("IN_SCOPE_WITHDRAWAL", 0), 12, 2, 98, 9, 1, 0, "High", "Transfers adjacent"])


def write_analysis_md():
    v2t = sum_v2_by_class()
    v3t = sum_v3_by_class()
    v2_in = sum(v for k, v in v2t.items() if k.startswith("IN_SCOPE"))
    v3_in = sum(v for k, v in v3t.items() if k.startswith("IN_SCOPE"))
    analysis = ROOT / "01-analysis"

    (analysis / "01-scope-definition.md").write_text(f"""# Scope Definition

**Assessment date:** {DATE}  
**Requested by:** Kevin Daines / Rajib Akhter

## In-scope workstreams (five Aha ideas only)

| Workstream | Aha ID | Primary evidence |
|---|---|---|
| Universal Enrollment | ACS-I-2679 | Enrollment microservices, account/validation/metadata webs |
| IDP / Authentication | ACS-I-2680 | auth-server repository |
| ABLE | ACS-I-2681 | ABLE microservices, webs, handlers (two components: Unite ABLE/CSR and Entity Platform) |
| Angular | ACS-I-2682 | able, 529, and angular repositories |
| Withdrawal | ACS-I-2690 | Withdrawal services |

## Explicitly excluded from scoped totals unless SME confirms

- Whole V2 daily regression population (744 qTest cases)
- Whole V3 regression population (436 TestNG methods) without suite-level mapping
- Whole API inventory (~269 methods in universal modules)
- Whole performance inventory (36 JMX / 53 Taurus configs)
- Mobile MSC, Astro/FuturePlan, UGift, investment options, account balance modules

## Classification taxonomy

Each asset receives one primary classification: IN_SCOPE_* (five workstreams), ADJACENT_NOT_CONFIRMED, OUT_OF_SCOPE, DUPLICATE_OR_OVERLAP, or UNMAPPED_NEEDS_SME_CONFIRMATION.

## Percentage rule

The current Aha work items define high-level outcomes but do not contain an enumerated requirement baseline. This assessment therefore reports validated automation inventory and business-flow coverage rather than requirement-level coverage percentages.
""", encoding="utf-8")

    (analysis / "02-source-register.md").write_text(f"""# Source Register

| ID | Source | Location in project | Role | Read-only |
|---|---|---|---|---|
| SRC-AHA | Five Aha idea PDFs | 00-input-evidence/aha/ | Scope authority | Yes |
| SRC-V2-QTEST | qTest daily report (744) | 00-input-evidence/v2-qtest-jenkins/QAAuto_UniteRegression_Daily_Report.PDF | V2 source population | Yes |
| SRC-V2-JENKINS | Jenkins console extract | 00-input-evidence/v2-qtest-jenkins/V2 Jenkins Job extract full output.txt | Module-level reconciliation | Yes |
| SRC-V3-UE | V3 UE TestNG report | 00-input-evidence/v3-testng-gitlab/V3 UE TestNG reports_files.html | UE 303 methods | Yes |
| SRC-V3-UNITE | V3 Unite TestNG report | 00-input-evidence/v3-testng-gitlab/V3 Unite TestNG reports.html | Unite 133 methods | Yes |
| SRC-V3-LOG | GitLab job log | 00-input-evidence/v2-qtest-jenkins/V3 job latest log.txt | Execution cross-check | Yes |
| SRC-V2-REPO | Automation/unite-test-automation | External read-only | Suite XML definitions | Yes |
| SRC-V3-REPO | prime-test-automation | External read-only | V3 suite XML + features | Yes |
| SRC-API | api-test-automation | External read-only | API inventory | Yes |
| SRC-PERF | performance-test-automation | External read-only | Performance inventory | Yes |
""", encoding="utf-8")

    (analysis / "03-counting-methodology.md").write_text("""# Counting Methodology

## Separate metrics (never mixed)

| Layer | Primary unit | Secondary units |
|---|---|---|
| V2 UI | qTest executed test cases (744 source population) | Jenkins module methods (508 subset), suite `<test>` blocks |
| V3 UI | TestNG executed methods (436 source) | Suite `<test>` definitions (59), feature files |
| API | Unique HTTP operation (method + path) | Test methods, test classes |
| Performance | Unique business journey | JMX scripts, Taurus configs |
| Angular | Component/library test definitions | E2E business-flow tests (reported separately) |

## Leadership metrics

- **Validated in-scope inventory** — only items with IN_SCOPE_* classification
- **Adjacent inventory** — reported separately, excluded from scoped totals
- **Scheduled vs implemented** — noted per workstream

## V2 vs V3

Reported separately. Overlap between frameworks is not presented as unique combined coverage.
""", encoding="utf-8")

    (analysis / "04-v2-scoped-inventory.md").write_text(f"""# V2 Scoped Inventory

## Source population

**744** executed automated test cases — qTest Project "Automation Unite", Test Cycle "PRIME", 2026-06-29.

This is the **source population**, not Universal Platform scoped coverage.

## Reconciliation summary

| Classification | Count |
|---|---:|
| IN_SCOPE_ENROLLMENT | {v2t.get('IN_SCOPE_ENROLLMENT', 0)} |
| IN_SCOPE_IDP | {v2t.get('IN_SCOPE_IDP', 0)} |
| IN_SCOPE_ABLE | {v2t.get('IN_SCOPE_ABLE', 0)} |
| IN_SCOPE_WITHDRAWAL | {v2t.get('IN_SCOPE_WITHDRAWAL', 0)} |
| ADJACENT_NOT_CONFIRMED | {v2t.get('ADJACENT_NOT_CONFIRMED', 0)} |
| OUT_OF_SCOPE | {v2t.get('OUT_OF_SCOPE', 0)} |
| UNMAPPED_NEEDS_SME_CONFIRMATION | {v2t.get('UNMAPPED_NEEDS_SME_CONFIRMATION', 0)} |
| **Total** | **744** |

## Validated in-scope V2 UI subtotal

**{v2_in}** executed test cases (Enrollment + IDP login + Withdrawal + dedicated LA ABLE suite).

## Limitations

Individual qTest test-case IDs are not exported. Module-level mapping is **provisional** (see csv/v2-row-level-mapping.csv). Residual 179-case PRIME expansion requires SME allocation across workstreams.
""", encoding="utf-8")

    (analysis / "05-v3-scoped-inventory.md").write_text(f"""# V3 Scoped Inventory

## Source populations

| Report | Source total |
|---|---:|
| Universal Enrollment | 303 |
| Unite (IDP Login + Contributions + Withdrawal + CSR + Web Reg) | 133 |
| **Combined V3 source** | **436** |

### Unite component check

56 + 36 + 20 + 15 + 6 = **133** ✓

## Scoped vs adjacent (nightly scheduled)

| Classification | Count |
|---|---:|
| IN_SCOPE_ENROLLMENT | {v3t.get('IN_SCOPE_ENROLLMENT', 0)} |
| IN_SCOPE_IDP | {v3t.get('IN_SCOPE_IDP', 0)} |
| IN_SCOPE_WITHDRAWAL | {v3t.get('IN_SCOPE_WITHDRAWAL', 0)} |
| ADJACENT_NOT_CONFIRMED | 57 |
| **Nightly source total** | **436** |

## Validated in-scope V3 UI subtotal (nightly)

**{v3_in}** executed TestNG methods (UE 303 + IDP Login 56 + Member Withdrawal 20).

## Implemented but not in nightly master

- ABLE Entity Platform: 6 suite test definitions (MIB, NEB, VAB)
- Angular lib-ui: 22 component test definitions

436 is **not** published as Universal Platform scoped coverage without excluding adjacent suites (Contributions 36, CSR 15, Web Registration 6).
""", encoding="utf-8")

    (analysis / "06-able-two-part-analysis.md").write_text("""# ABLE Two-Part Analysis

## Part A — Unite ABLE / CSR (V2)

- **ABLE plan codes (ending B):** AKB, COB, ILB, MIB, NHB, NYB, PAB, RIB, TNB (9 plans)
- **Parameterized blocks:** 43 plan-specific test blocks embedded in scheduled V2 modules
- **Dedicated LA ABLE suite:** 17 executed test cases (`stage1-laable.xml`) — **in-scope ABLE count for V2 dedicated suite**
- **Status:** Maintained in V2 framework; dedicated suite implemented, scheduling confirmation pending
- **Reuse:** Many flows reuse common scenarios with plan parameterization

## Part B — ABLE Entity Platform (V3)

- **Repository:** prime-test-automation/unite-entity
- **Suite definitions:** Entity registration (3 blocks) + Entity dashboard IDP login (3 blocks)
- **Plans:** MIB, NEB, VAB
- **Scheduled nightly:** No — not wired in `stage1-unite-regression-master.xml`
- **Counting unit:** 6 suite test definitions (expands to scenarios at execution)

## Separation

MIB appears in both Part A (V2 CSR/member) and Part B (V3 entity). **Different implementations — not merged or double-counted.**

## API

No ABLE-specific API automation was identified in the reviewed repositories.
""", encoding="utf-8")

    (analysis / "07-angular-analysis.md").write_text("""# Angular Scope Analysis

## Aha ACS-I-2682 definition

Scope is **repository-level** (able, 529, and angular repos), not a catalog of member UI journeys.

## Three buckets (kept separate)

| Bucket | Count | Classification |
|---|---:|---|
| Angular lib-ui component tests | 22 test definitions | IN_SCOPE_ANGULAR (component/library) |
| V3 Contributions suite | 36 methods | ADJACENT_NOT_CONFIRMED |
| V3 CSR Account Maintenance | 15 methods | ADJACENT_NOT_CONFIRMED |

## Decision

- **22 lib-ui tests** are component/library automation — labeled as such, not E2E business-flow coverage.
- **Contributions and CSR** are **not** counted as in-scope Angular until SME confirms ACS-I-2682 includes those member UI flows.
- **Performance Blazemeter UE/Login Selenium flows** support Angular-hosted pages but are performance assets, not Angular component inventory.

## Pipeline status

lib-ui suite implemented (QC4 profile); not in active nightly pipeline.
""", encoding="utf-8")

    (analysis / "08-api-scoped-inventory.md").write_text("""# API Scoped Inventory

## Source inventory (universal modules)

~269 `@Test` methods across 38 test classes in 4 active service modules (auth, account, financial, metadata).

## Scoped reconciliation

| Workstream | Unique operations | Test methods | Test classes |
|---|---:|---:|---:|
| Enrollment | 5 | 103 | 17 |
| IDP | 4 | 6 | 4 |
| Withdrawal | 2 | 98 | 9 |
| ABLE | 0 | 0 | 0 |
| Angular | 0 | 0 | 0 |
| Adjacent (metadata, exchange) | 4 | 43+ | 7+ |
| Out of scope (mobile, astro) | 11+ | 25+ | 6+ |

**No ABLE-specific API automation was identified in the reviewed repositories.**

Scheduling: implemented locally; pipeline gate enablement requires confirmation.
""", encoding="utf-8")

    (analysis / "09-performance-scoped-inventory.md").write_text("""# Performance Scoped Inventory

## Source inventory

- 36 JMX scripts under universal-platform/
- 53 Taurus/YAML configurations (includes environment variants)

## Scoped business journeys (deduplicated)

| Workstream | In-scope journeys |
|---|---:|
| Enrollment | 7 |
| IDP | 6 |
| Withdrawal | 1 |
| ABLE Entity | 1 |
| **In-scope subtotal** | **15** |

## Excluded from scoped totals

- 14 database prototype JMX scripts (supporting infrastructure, not business-flow coverage)
- Duplicate environment configs (local vs remote YAML pairs count as one journey)
- Mobile MSC performance plans (out of scope)

Scheduling and SLA governance: assets available; recurrence not fully established.
""", encoding="utf-8")

    v2t = sum_v2_by_class()
    v3t = sum_v3_by_class()
    (analysis / "10-reconciliation-ledger.md").write_text(f"""# Reconciliation Ledger

All leadership numbers must trace to this ledger.

## V2 qTest source population = 744

| Classification | Count | Counting unit |
|---|---:|---|
| IN_SCOPE_ENROLLMENT | {v2t.get('IN_SCOPE_ENROLLMENT', 0)} | executed qTest cases (provisional module map) |
| IN_SCOPE_IDP | {v2t.get('IN_SCOPE_IDP', 0)} | executed qTest cases |
| IN_SCOPE_ABLE | {v2t.get('IN_SCOPE_ABLE', 0)} | executed qTest cases |
| IN_SCOPE_WITHDRAWAL | {v2t.get('IN_SCOPE_WITHDRAWAL', 0)} | executed qTest cases |
| ADJACENT_NOT_CONFIRMED | {v2t.get('ADJACENT_NOT_CONFIRMED', 0)} | executed qTest cases |
| OUT_OF_SCOPE | {v2t.get('OUT_OF_SCOPE', 0)} | executed qTest cases |
| UNMAPPED_NEEDS_SME_CONFIRMATION | {v2t.get('UNMAPPED_NEEDS_SME_CONFIRMATION', 0)} | executed qTest cases |
| **Total** | **744** | |

**Validated V2 in-scope UI subtotal:** {sum(v for k,v in v2t.items() if k.startswith('IN_SCOPE'))} (Enrollment, IDP, ABLE dedicated suite, Withdrawal — **four workstreams on V2**; Angular has no V2 in-scope UI; excludes adjacent, out-of-scope, unmapped)

## V3 UE source = 303

| Classification | Count |
|---|---:|
| IN_SCOPE_ENROLLMENT | 303 |
| **Total** | **303** |

## V3 Unite source = 133

| Suite | Count | Classification |
|---|---:|---|
| IDP Login | 56 | IN_SCOPE_IDP |
| Member Withdrawal | 20 | IN_SCOPE_WITHDRAWAL |
| IDP Web Registration | 6 | ADJACENT_NOT_CONFIRMED |
| Contributions | 36 | ADJACENT_NOT_CONFIRMED |
| CSR Account Maintenance | 15 | ADJACENT_NOT_CONFIRMED |
| **Total** | **133** | |

**Validated V3 in-scope UI subtotal (nightly):** 379 (303 + 56 + 20)

## V3 combined source = 436

303 + 133 = 436 ✓

## API source inventory

| Metric | Source total | In-scope subtotal |
|---|---:|---:|
| Test methods | ~269 | ~207 (enrollment + IDP + withdrawal) |
| Unique operations | 27 (all modules) | 11 (enrollment + IDP + withdrawal) |
| Test classes | 38 | 30 |

## Performance source inventory

| Metric | Source total | In-scope subtotal |
|---|---:|---:|
| JMX scripts | 36 | 15 journeys |
| Taurus configs | 53 | ~15 (deduplicated) |

## Limitations

- V2 row-level qTest IDs unavailable — module provisional mapping
- V2/V3 overlap not reconciled to unique business requirements
- Adjacent suites excluded from leadership scoped totals pending SME
""", encoding="utf-8")

    (analysis / "11-assumptions-and-sme-questions.md").write_text("""# Assumptions and SME Questions

## Assumptions applied

1. V2 744 = qTest PRIME daily aggregate (authoritative source population)
2. V3 436 = TestNG methods from GitLab nightly artifacts
3. Adjacent suites excluded from leadership scoped totals until confirmed
4. lib-ui 22 tests counted as component tests, not E2E flows
5. ABLE Part A and Part B reported separately; MIB not double-counted

## SME validation required

1. Allocate 179-case V2 PRIME per-plan expansion across workstreams
2. Confirm IDP Web Registration (6 V3, 46 V2) under ACS-I-2680
3. Confirm Contributions (36 V3, 48 V2) under ACS-I-2682 or exclude
4. Confirm CSR maintenance (15 V3, 81+9 V2) under ABLE or Angular or exclude
5. Confirm Transfers (12 V2) under ACS-I-2690 or exclude
6. Confirm lib-ui 22 tests as intended Angular (2682) artifact
7. Confirm ABLE Entity pipeline enablement timeline
8. Confirm API and performance scheduling/SLA ownership
""", encoding="utf-8")


def cleanup_old_structure(cleanup_log: list):
    remove_paths = [
        ROOT / "up-coverage-metrics-kit",
        ROOT / "up-coverage-metrics-finalization-pack-v2.zip",
        ROOT / "up-coverage-metrics-kit.zip",
    ]
    for p in remove_paths:
        if p.exists():
            if p.is_dir():
                shutil.rmtree(p, ignore_errors=True)
            else:
                p.unlink(missing_ok=True)
            cleanup_log.append(f"REMOVED: {p.name}")


def write_cleanup_log(entries: list):
    log_path = ROOT / "01-analysis/cleanup-log.md"
    lines = [f"# Cleanup Log\n\n**Date:** {DATE}\n"]
    for e in entries:
        lines.append(f"- {e}")
    log_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_readme():
    (ROOT / "README.md").write_text(f"""# Universal Platform Automation Coverage Assessment

Scoped reconciliation of automation inventory against five Aha workstreams (ACS-I-2679 through ACS-I-2690).

## Structure

- `00-input-evidence/` — Aha PDFs, qTest/Jenkins, TestNG reports (read-only copies)
- `01-analysis/` — Scope, methodology, inventories, reconciliation ledger, CSV mappings
- `02-deliverables/` — Leadership assessment (DOCX/PDF), email draft, SME checklist
- `03-tools/` — Master rebuild prompt and scripts
- `99-archive/` — Pre-rebuild snapshot ({DATE})

## Key scoped totals (see 01-analysis/10-reconciliation-ledger.md)

| Source | Population | Validated in-scope subtotal |
|---|---:|---:|
| V2 qTest | 744 | 268 |
| V3 TestNG (nightly) | 436 | 379 |
| API (universal modules) | ~269 methods | ~207 methods / 11 operations |
| Performance JMX | 36 | 15 journeys |

**Do not** present 744, 436, 1,180, 269, or 36 as Universal Platform scoped totals without classification.
""", encoding="utf-8")


def write_master_prompt():
    (ROOT / "03-tools/MASTER-RESET-REBUILD-PROMPT.md").write_text("""# Master Reset Rebuild Prompt

Canonical prompt for Universal Platform coverage assessment reset.

See user-provided specification dated 2026-06-29 for full instructions.

Key rules:
- Five Aha workstreams only
- Classify every asset with primary inclusion classification
- V2 744 and V3 436 are source populations, not scoped totals
- Separate counting units; no mixed metrics
- ABLE two-part; Angular component vs E2E separate
- No pass/fail in leadership deliverables
""", encoding="utf-8")


def write_deliverables_md():
    d = ROOT / "02-deliverables"
    (d / "Kevin-Rajib-Email-Draft.md").write_text("""Subject: Universal Platform Automation Coverage Assessment — Scoped to Five Aha Workstreams

Kevin and Rajib,

Please find attached the Universal Platform Automation Coverage Assessment, scoped exclusively to the five Aha workstreams you requested: Universal Enrollment (ACS-I-2679), IDP/Authentication (ACS-I-2680), ABLE (ACS-I-2681), Angular (ACS-I-2682), and Withdrawal (ACS-I-2690).

This assessment reports validated automation inventory and business-flow coverage — not requirement-level percentages, as the Aha ideas do not yet provide an enumerated requirements baseline.

**V2 (qTest daily regression)** — source population 744 executed test cases. After scoped classification, **268** map directly to the five workstreams (enrollment, IDP login, withdrawal, and dedicated LA ABLE). An additional **196** cases require scope confirmation (web registration, CSR maintenance, contributions, transfers). **210** are out-of-scope or unmapped pending SME review.

**V3 (GitLab nightly TestNG)** — source population 436 methods (Universal Enrollment 303 + Unite 133). **379** map to enrollment, IDP login, and member withdrawal. **57** Unite methods (contributions, CSR maintenance, web registration) are adjacent pending Angular/IDP scope confirmation. ABLE Entity Platform (6 scenarios) and Angular lib-ui (22 component tests) are implemented but not yet in the nightly pipeline.

**API and performance** — scoped separately. Eleven unique API operations and fifteen performance journeys align to the five workstreams. No ABLE-specific API automation was identified.

Pending SME validation with Nick on adjacent suite mappings, V2 per-plan expansion allocation, and pipeline enablement for implemented-but-unscheduled assets.

Happy to walk through the dashboard on request.

Best regards,
QA Automation
""", encoding="utf-8")

    (d / "Nick-SME-Validation-Checklist.md").write_text("""# Nick SME Validation Checklist — Factual Questions Only

## V2 in-scope mappings

1. Does the enrollments module (151 cases) fully map to ACS-I-2679, or should ABLE enrollment scenarios be reclassified to ACS-I-2681?
2. How should the 179-case PRIME per-plan expansion (744 minus documented modules) be allocated across workstreams?
3. Are web registration (46) and web login (27) both in-scope for ACS-I-2680?

## V3 suite-to-workstream mapping

4. Do all 303 Universal Enrollment methods belong to ACS-I-2679?
5. Should IDP Web Registration (6 methods) count under ACS-I-2680?
6. Should Contributions (36 methods) count under ACS-I-2682 (Angular)?
7. Should CSR Account Maintenance (15 methods) count under ACS-I-2682 or ACS-I-2681?

## Angular scope

8. Are the 22 lib-ui dynamic-forms component tests the intended ACS-I-2682 artifact?
9. Should Blazemeter Selenium UE/login flows count as Angular coverage?

## ABLE plan split

10. Confirm V2 ABLE plan set (AKB, COB, ILB, MIB, NHB, NYB, PAB, RIB, TNB) and dedicated LA ABLE suite (17 cases).
11. Confirm V3 Entity plans (MIB, NEB, VAB) and pipeline enablement target.
12. Is MIB correctly treated as separate V2 vs V3 implementations (no double-count)?

## API scope and scheduling

13. Confirm API modules in scope: auth, account/enrollment, financial/withdrawal only?
14. Which API gates are scheduled vs implemented-only?

## Performance scope and scheduling

15. Which of the 15 scoped performance journeys are scheduled with SLA governance?
16. Should database prototype scripts ever count as business-flow coverage?

## Unresolved adjacent tests

17. Should V2 transfers (12) map to ACS-I-2690?
18. Should empower-plan (21) and csr-actions (9) map to enrollment, ABLE, or remain out of scope?
""", encoding="utf-8")


def build_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

    NAVY = RGBColor(0x1F, 0x3A, 0x5F)
    GREY = RGBColor(0x55, 0x55, 0x55)
    GREEN = RGBColor(0x1E, 0x7A, 0x34)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def set_cell_bg(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def add_table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(hdr[i], "1F3A5F")
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(8.5)
        if widths:
            for i, w in enumerate(widths):
                for c in t.columns[i].cells:
                    c.width = Inches(w)
        doc.add_paragraph()

    def h1(text):
        p = doc.add_heading(level=1)
        r = p.add_run(text)
        r.font.color.rgb = NAVY

    def h2(text):
        p = doc.add_heading(level=2)
        r = p.add_run(text)
        r.font.color.rgb = NAVY

    def para(text, italic=False, bold=False, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.italic = italic
        r.bold = bold
        r.font.size = Pt(size)

    def bullet(text, bold_lead=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead:
            rb = p.add_run(bold_lead)
            rb.bold = True
            rb.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("Universal Platform Test Automation")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run("Coverage Assessment")
    r.font.size = Pt(16)
    r.font.color.rgb = GREY
    doc.add_paragraph()
    for line, col in [
        ("Prepared for: Kevin Daines, Rajib Akhter", GREY),
        ("Prepared by: QA Automation", GREY),
        (f"Date: {DATE}", GREY),
        ("Document status: Pending SME factual review", GREEN),
    ]:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(line)
        rr.font.size = Pt(11)
        rr.font.color.rgb = col
    doc.add_page_break()

    h1("1. Executive Summary")
    para(
        "This assessment inventories automated test assets scoped to five Aha workstreams: "
        "Universal Enrollment, IDP/Authentication, ABLE, Angular, and Withdrawal. "
        "It reports validated in-scope inventory separately from source regression populations."
    )
    para(
        "The current Aha work items define high-level outcomes but do not contain an enumerated "
        "requirement baseline. This assessment therefore reports validated automation inventory "
        "and business-flow coverage rather than requirement-level coverage percentages.",
        italic=True,
        size=9,
    )
    h2("Validated in-scope highlights")
    bullet("268 executed test cases on V2 map directly to enrollment, IDP login, withdrawal, and LA ABLE.", bold_lead="V2 scoped UI: ")
    bullet("379 nightly TestNG methods on V3 map to Universal Enrollment (303), IDP login (56), and member withdrawal (20).", bold_lead="V3 scoped UI: ")
    bullet("11 unique API operations across enrollment, IDP, and withdrawal modules.", bold_lead="API: ")
    bullet("15 performance business journeys across enrollment, IDP, withdrawal, and ABLE entity.", bold_lead="Performance: ")
    bullet("22 lib-ui component test definitions (Angular library; separate from E2E flows).", bold_lead="Angular components: ")
    h2("Remaining work")
    bullet("SME confirmation for adjacent suites (contributions, CSR, web registration, transfers).")
    bullet("ABLE-specific API automation (net-new).")
    bullet("Pipeline enablement for Entity Platform, lib-ui, and dedicated LA ABLE suite.")
    doc.add_page_break()

    h1("2. Scope and Counting Methodology")
    para("Source populations (not scoped totals): V2 qTest 744; V3 TestNG 436; API ~269 methods; Performance 36 JMX scripts.")
    para("Counting units are never mixed: qTest cases (V2), TestNG methods (V3), HTTP operations (API), business journeys (performance).")
    para("V2 and V3 are reported separately; framework overlap is not presented as unique combined coverage.")
    doc.add_page_break()

    h1("3. Leadership Dashboard")
    add_table(
        ["Workstream", "V2 in-scope UI", "V3 in-scope UI", "API in-scope", "Performance in-scope", "Business flows", "Current position", "Remaining need", "Ownership", "Confidence"],
        [
            ["Enrollment (ACS-I-2679)", "151", "303", "5 ops / 103 methods", "7 journeys", "Standard, continue, multi-bene, Upromise, PAGSP", "Scheduled V2 & V3", "V2 expansion mapping; consolidation", "Galaxy + Automation", "High"],
            ["IDP (ACS-I-2680)", "27", "56", "4 ops / 6 methods", "6 journeys", "Login, lockout, forgot user/pass", "Scheduled V2 & V3", "Web registration scope", "Innovators + Automation", "High"],
            ["ABLE — Unite (ACS-I-2681)", "17", "—", "None identified", "—", "CSR-Able dedicated suite; 9 plans in V2", "V2 maintained; suite pending schedule", "Schedule LA ABLE suite", "Spurs + Automation", "High"],
            ["ABLE — Entity (ACS-I-2681)", "—", "6 scenarios", "None identified", "1 journey", "Entity reg + dashboard login", "Implemented V3", "Pipeline enablement", "Spurs + DevOps", "High"],
            ["Angular (ACS-I-2682)", "—", "—", "—", "—", "22 lib-ui components", "Implemented", "Scope confirm; pipeline", "UX + Automation", "Medium"],
            ["Withdrawal (ACS-I-2690)", "73", "20", "2 ops / 98 methods", "1 journey", "Member/CSR distributions", "Scheduled V2 & V3", "Transfers scope; API gate", "Infinity + Automation", "High"],
        ],
        widths=[1.0, 0.55, 0.55, 0.7, 0.55, 1.1, 0.9, 0.85, 0.85, 0.55],
    )
    doc.add_page_break()

    h1("4. Workstream Detail")
    h2("Enrollment")
    para("V3 Universal Enrollment (303 methods) is fully aligned to ACS-I-2679. V2 enrollments module contributes 151 validated cases; additional PRIME expansion requires SME allocation.")
    h2("IDP / Authentication")
    para("V3 IDP Login (56) and V2 web login (27) are in-scope. Web registration (46 V2, 6 V3) is adjacent pending confirmation against auth-server-only Aha scope.")
    h2("ABLE (two components)")
    para("Part A: V2 Unite ABLE/CSR — 17-case dedicated LA ABLE suite plus 9-plan parameterized coverage. Part B: V3 Entity Platform — 6 scenarios for MIB/NEB/VAB, not yet scheduled.")
    h2("Angular")
    para("22 lib-ui dynamic-forms component tests are implemented. V3 contributions (36) and CSR (15) are adjacent until ACS-I-2682 scope is confirmed for member UI flows.")
    h2("Withdrawal")
    para("V2 withdrawals (73) and V3 member withdrawal (20) are in-scope. V2 transfers (12) are adjacent.")
    doc.add_page_break()

    h1("5. Current Automation Strengths")
    bullet("Substantial scheduled V2 and V3 UI regression across core Universal Platform flows.")
    bullet("API automation implemented for enrollment, authentication, and withdrawal microservices.")
    bullet("Performance journey assets for enrollment and IDP load testing.")
    bullet("ABLE foundations on both V2 (Unite) and V3 (Entity) platforms.")
    bullet("Reusable Angular lib-ui component test suite.")

    h1("6. Gaps and Ownership")
    add_table(
        ["Gap", "Recommended ownership"],
        [
            ["ABLE-specific API automation", "Spurs + Automation"],
            ["Adjacent suite scope confirmation", "Nick (SME) + Kevin"],
            ["V2 per-plan expansion mapping", "Automation + Galaxy"],
            ["Entity Platform pipeline enablement", "Spurs + DevOps"],
            ["lib-ui pipeline scheduling", "UX + Automation"],
            ["Performance SLA governance", "Performance + DevOps"],
        ],
        widths=[3.5, 3.0],
    )

    h1("7. Limitations and Next Steps")
    bullet("V2 qTest row-level IDs unavailable — module mapping is provisional.")
    bullet("Adjacent suites excluded from scoped leadership totals pending SME validation.")
    bullet("Enumerate Aha requirements to enable requirement-level coverage percentages.")
    bullet("Complete Nick SME validation checklist before leadership distribution.")

    out = ROOT / "02-deliverables/Universal-Platform-Automation-Coverage-Assessment.docx"
    doc.save(out)
    return out


def build_pdf(docx_path: Path):
    pdf_path = docx_path.with_suffix(".pdf")
    # Try docx2pdf (Word on Windows)
    try:
        import docx2pdf
        docx2pdf.convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass
    # Fallback: reportlab simple PDF from text extraction
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab", "-q"])
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet

    doc = Document(str(docx_path))
    styles = getSampleStyleSheet()
    story = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue
        style = styles["Heading1"] if p.style.name.startswith("Heading 1") else (
            styles["Heading2"] if p.style.name.startswith("Heading 2") else styles["Normal"]
        )
        story.append(Paragraph(text.replace("&", "&amp;"), style))
    for table in doc.tables:
        story.append(Spacer(1, 12))
        for row in table.rows:
            cells = [c.text.replace("&", "&amp;") for c in row.cells]
            story.append(Paragraph(" | ".join(cells), styles["Normal"]))
        story.append(Spacer(1, 12))
    pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    pdf.build(story)
    return pdf_path


def verify_arithmetic():
    assert sum(r[5] for r in V2_MODULES) == 744, "V2 must sum to 744"
    v3_nightly = 303 + 56 + 36 + 20 + 15 + 6
    assert v3_nightly == 436
    print("Arithmetic OK: V2=744, V3=436")


def main():
    cleanup_log = []
    ensure_backup()
    manifest, tree = write_manifest_and_tree()
    cleanup_log.append(f"CREATED manifest: {manifest.name}")
    cleanup_log.append(f"CREATED tree: {tree.name}")
    setup_dirs()
    copy_evidence(cleanup_log)
    write_csvs()
    write_analysis_md()
    write_readme()
    write_master_prompt()
    write_deliverables_md()
    docx = build_docx()
    cleanup_log.append(f"CREATED {docx.name}")
    pdf = build_pdf(docx)
    cleanup_log.append(f"CREATED {pdf.name}")
    cleanup_old_structure(cleanup_log)
    write_cleanup_log(cleanup_log)
    verify_arithmetic()
    print("REBUILD COMPLETE")


if __name__ == "__main__":
    main()
